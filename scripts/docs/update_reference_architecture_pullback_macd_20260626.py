#!/usr/bin/env python3
"""2026-06-26 — Append the 'Pullback / MACD Reversal Screener' section to the canonical Reference
Architecture DOCX. APPEND-ONLY, marker-guarded/idempotent, backup + verify. Follows
docs/DOCX_UPDATE_PROTOCOL.md (python-docx, heading styles pulled from existing paragraphs)."""
import shutil
from datetime import datetime, timezone

from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Pullback / MACD Reversal Screener (2026-06-26)"

ARTIFACTS = [
    ("Artifact", "Type", "Purpose"),
    ("pullback_macd_screener.py", "engine",
     "Daily scan + intraday --monitor; signals, proposals, reconciliation, in-trade adjustments"),
    ("pullback_macd_candidates / _runs", "DB tables", "Ranked candidates (trigger/watch) + run audit"),
    ("trade_plans (generated_by=pullback)", "DB rows", "Authoritative technical levels → clears the trade-plan gate"),
    ("pullback_trade_adjustments", "DB table", "Advisory trail/TP/exit guidance for OPEN positions"),
    ("/api/v2/pullback-macd/{candidates,adjustments,scan,dismiss}", "endpoints", "Tab data + manage actions"),
    ("Command Center 'Pullback/MACD' tab", "UI", "Trigger/watch cards, pullback + VWAP banners, manage controls"),
    ("run_pullback_macd_monitor.sh", "cron 35 9-15", "Hourly intraday monitor, trading-days only"),
]


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def main():
    doc = Document(DOCX_PATH)
    if any(MARKER in (p.text or "") for p in doc.paragraphs):
        print("Pullback/MACD section already present, skip")
        return
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M")
    backup = f"{DOCX_PATH}.bak_pullback_{ts}"
    shutil.copy2(DOCX_PATH, backup)
    print("backup:", backup)

    head(doc, 2, MARKER)
    head(doc, 3, "Strategy")
    doc.add_paragraph(
        "An S&P 500 counter-trend dip-buy screener. It finds names in a standing uptrend (close > SMA200, "
        "SMA50 > SMA200 and rising) that have pulled back into a 12-28% band off their 52-week high, then "
        "confirms the EARLIEST recovery: the MACD histogram has turned up off the pullback (inflection — "
        "still pre-cross) AND price is holding above intraday session VWAP. Requiring the inflection rather "
        "than waiting for the cross catches the move earlier; VWAP is the price-action confirmation. Two "
        "tiers: trigger (both confirmed) and watch (in pullback, not yet confirmed).")
    head(doc, 3, "Authoritative levels & advisory proposals")
    doc.add_paragraph(
        "Levels are technical, never generic R:R geometry: stop = recent swing-low support, target = retrace "
        "toward the 52-week high. Each emitted proposal also writes a trade_plans row, which broker_trade_plan_gate "
        "resolves as authoritative (plan_source=trade_plans) — clearing the system-wide 'no authoritative trade "
        "plan / gambling-blocked' route block. Proposals are advisory (PENDING, manual), trigger-tier only, ranked "
        "by score and hard-capped per scan (max_proposals_per_scan) because each PENDING proposal triggers the "
        "local+cloud LLM oversight fleet.")
    head(doc, 3, "Intraday monitoring & in-trade adjustments")
    doc.add_paragraph(
        "A daily post-close scan sets the pullback universe; an hourly intraday monitor (run_pullback_macd_monitor.sh, "
        "cron 35 9-15 weekdays, market_day_gate so it runs trading days only) keeps everything in sync with the live "
        "tape. It refreshes proposals that still fit the plan, expires those that don't (off-trigger, or price "
        "hit/broke stop/target), and catches new intraday triggers. While IN a trade, _adjust_open_trades writes "
        "advisory guidance to pullback_trade_adjustments for each open position: trail the stop up (swing-low / "
        "breakeven / under-VWAP, raise-only), take-profit at target, or exit on thesis break (lost VWAP or MACD "
        "rolling back down). It is advisory only — it never modifies a live stop; the operator and ATM stop manager "
        "retain control.")
    head(doc, 3, "Monitoring & surfaces")
    doc.add_paragraph(
        "Health collectors cover freshness (pullback_macd_scan_stale, universe size) and a proposal-burst guard "
        "(proposal_creation_burst) that flags bulk proposal creation before it can overload the single-threaded "
        "server. The Command Center 'Pullback/MACD' tab shows trigger/watch cards with pullback and VWAP banners "
        "plus run-scan/dismiss controls.")

    t = doc.add_table(rows=len(ARTIFACTS), cols=3)
    borders(t)
    for ri, row in enumerate(ARTIFACTS):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val

    head(doc, 3, "References")
    doc.add_paragraph("Detail: docs/PULLBACK_MACD_SCREENER.md; docs/CHANGELOG.md (2026-06-26 entries).")

    doc.save(DOCX_PATH)
    chk = Document(DOCX_PATH)
    ok = any(MARKER in (p.text or "") for p in chk.paragraphs)
    print("saved + verified:" if ok else "VERIFY FAILED:", DOCX_PATH, "| paragraphs:", len(chk.paragraphs))


if __name__ == "__main__":
    main()

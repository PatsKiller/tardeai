#!/usr/bin/env python3
"""2026-06-26 — Append 'Options Desk enterprise layer: audit fixes + snapshot retention'
to the canonical Reference Architecture DOCX. APPEND-ONLY, marker-guarded/idempotent.
Backs up before editing, verifies after. Does NOT remove or rewrite existing content.

Follows docs/DOCX_UPDATE_PROTOCOL.md: python-docx append-only, heading styles pulled
from existing paragraphs, table borders via tblPr XML append, backup + verify.
"""
import shutil
from datetime import datetime, timezone

from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Session 2026-06-26 - Options Desk enterprise layer: audit fixes + snapshot retention"

ARTIFACTS = [
    ("Artifact", "Type", "Purpose"),
    ("options_desk_enterprise.py", "module",
     "Enterprise gates, book greeks, vol analytics, approval queue, retention"),
    ("options_approval_queue", "DB table",
     "Operator review/approval gate before any live options submit (fail-closed)"),
    ("options_chain_snapshots", "DB table",
     "Vol-surface snapshots (ATM IV / skew / term slope); retention-bounded"),
    ("/api/v2/options/desk/risk", "endpoint",
     "Book-level greeks + per-symbol concentration preflight"),
    ("/api/v2/options/desk/vol-analytics | vol-history | trends", "endpoint",
     "Term structure, put/call skew, and IV trends from the live Schwab chain"),
    ("/api/v2/options/approval-queue (+ /resolve)", "endpoint",
     "List pending desk approvals; operator approve/reject"),
]


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def head_para(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text
    return p


def main():
    doc = Document(DOCX_PATH)
    if any(MARKER in (p.text or "") for p in doc.paragraphs):
        print("Options Desk 2026-06-26 section already present, skip")
        return

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M")
    backup = f"{DOCX_PATH}.bak_options_desk_{ts}"
    shutil.copy2(DOCX_PATH, backup)
    print("backup:", backup)

    head_para(doc, 2, MARKER)

    # Context: the enterprise desk layer (not previously recorded in this doc)
    head_para(doc, 3, "Enterprise options trade desk")
    doc.add_paragraph(
        "The advisory options engine (scripts/options_engine.py) gained an enterprise trade-desk "
        "layer (scripts/options_desk_enterprise.py) that sits between proposal generation and any "
        "live submit. It adds FMP earnings-blackout checks, hard liquidity gates (open interest, "
        "volume, bid/ask spread), volatility analytics (term structure and put/call skew from the "
        "live Schwab chain), book-level greeks aggregation, a portfolio risk preflight (net-delta "
        "and per-symbol concentration caps), A/B/C desk tiering by edge score, and a database-backed "
        "operator approval queue (options_approval_queue).")
    doc.add_paragraph(
        "Live submit is fail-closed. check_preflight_approval requires a proposal to be present in "
        "the approval queue, operator-approved, and live-eligible; a missing, rejected, or "
        "non-eligible proposal is blocked. Operator approval cannot clear a proposal while "
        "enterprise blocks remain. This sits alongside, not in place of, the existing per-order 2FA "
        "and execution-guard gates.")

    # The audit fixes
    head_para(doc, 3, "Post-merge audit fixes (2026-06-26)")
    doc.add_paragraph(
        "A correctness audit of the desk layer found and fixed six issues. (1) Estimated theta - "
        "used when the chain omits greeks, the common case for short premium - was not sign-flipped "
        "for short legs, so the book's net theta/day reported decay paid rather than collected; the "
        "estimator now returns long-convention (negative) theta and aggregate_book_greeks flips it "
        "for short legs, matching real chain theta. (2) The net-delta concentration check used a "
        "hardcoded assumed share price; it now derives a real dollar-delta (share-equivalent delta "
        "times each leg's actual underlying price) over book market value, consistent with the "
        "no-hardcoded-values rule. (3) The earnings-blackout cache skipped symbols added mid-window; "
        "it now fetches the missing subset so a near-earnings name cannot bypass its blackout. "
        "(4) Vol-surface snapshots stored a byte-sliced full chain that produced malformed JSON and "
        "silently failed to persist on large chains; they now store a small valid summary. "
        "(5) Snapshot retention was added (see below). (6) A proposal with no resolved chain "
        "contract - hence no verifiable fill liquidity - is never marked live-eligible, independent "
        "of the require-chain-for-live override.")

    # Retention
    head_para(doc, 3, "Vol-surface snapshot retention")
    doc.add_paragraph(
        "options_chain_snapshots is bounded two ways. A cheap per-symbol prune runs on each "
        "persist_chain_snapshot insert (active desk names), and a global sweep "
        "(prune_chain_snapshots) runs from the daily IV-snapshot cron "
        "(scripts/options_iv_snapshot.py, 20:16 on weekdays) to catch the tails of symbols that "
        "have gone quiet. Both honor OPTIONS_SNAPSHOT_RETENTION_DAYS (default 45). The retention "
        "logic was verified against the live database.")

    t = doc.add_table(rows=len(ARTIFACTS), cols=3)
    borders(t)
    for ri, row in enumerate(ARTIFACTS):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val

    head_para(doc, 3, "References")
    doc.add_paragraph(
        "Detail: docs/options-module.md; docs/CHANGELOG.md (2026-06-26 entries); "
        "migrations/2026_06_25_options_desk_enterprise.sql. Commits: e01e9847 (audit fixes), "
        "081ec19d (global retention sweep).")

    doc.save(DOCX_PATH)

    # Verify
    chk = Document(DOCX_PATH)
    ok = any(MARKER in (p.text or "") for p in chk.paragraphs)
    print("saved + verified:" if ok else "VERIFY FAILED:", DOCX_PATH, "| paragraphs:", len(chk.paragraphs))


if __name__ == "__main__":
    main()

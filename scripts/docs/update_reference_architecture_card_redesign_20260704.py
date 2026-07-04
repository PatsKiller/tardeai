#!/usr/bin/env python3
"""2026-07-04 — Append 'Decision-surface redesign: shared card grammar (watchlist, proposals, stops)'
to the canonical Reference Architecture DOCX. APPEND-ONLY, marker-guarded/idempotent.
Backs up before editing, verifies after. Follows docs/DOCX_UPDATE_PROTOCOL.md.
"""
import shutil
from datetime import datetime, timezone
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Session 2026-07-03/04 - Decision-surface redesign: shared card grammar (watchlist, proposals, stops)"

ARTIFACTS = [
    ("Artifact", "Type", "Purpose"),
    ("components/primitives/cardPrimitives.tsx", "module",
     "Shared primitives: VerdictBanner, LadderLine, TrustLine, AgeChip — single source for all three decision surfaces"),
    ("lib/watchlistCardAction.ts / lib/proposalCardAction.ts", "module",
     "Priority-ordered decision matrices driving each card's verdict banner (first match wins; red reserved for hard defects)"),
    ("WatchlistCard.tsx (Security Card v3)", "component",
     "LOCKED format: header / tinted banner / 2x2 grid (Trade Plan | Sizing, Conviction | Intelligence) / footer"),
    ("BrokerProposalCard.tsx (rebuilt, PR #80)", "component",
     "LOCKED format: header / banner / 2x2 grid (Order & Levels | Sizing & Fit, Intelligence | Gates & Reviews) / provenance footer + evidence expander"),
    ("StopManagement.tsx expanded row (PR #81)", "component",
     "LOCKED format: verdict banner (NO STOP/REVIEW/EARNINGS/PROTECTED) + 2x2 panel (Protection | Exit Plan, Health | Context) + collapsed-row amber scan chips"),
    ("_broker_enrich_trust_rows (api_v2.py)", "enrichment",
     "Proposal rows: held_shares_in_account/_total, next_earnings_date, catalyst_title/at (both cache paths)"),
    ("stops/management row enrichment (api_v2.py)", "enrichment",
     "Stop rows: rec_at, holdings_llm_at, next_earnings_date, news_title/source/at (batch, failure-tolerant)"),
    ("plan_drift_revalidator.py (cron 17:25 Mon-Fri)", "pipeline",
     "Re-plans incoherent (target<=limit) or drifted (>15%) entry plans via subprocess-batched planner"),
    ("cleanup_stale_proposals.py drift pass", "pipeline",
     "Auto-rejects PENDING proposals >4h old with live price >15% from proposed entry"),
]


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def head_para(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text
    return p


def main():
    doc = Document(DOCX_PATH)
    if any(MARKER in (p.text or "") for p in doc.paragraphs):
        print("Card redesign 2026-07-04 section already present, skip")
        return

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M")
    backup = f"{DOCX_PATH}.bak_card_redesign_{ts}"
    shutil.copy2(DOCX_PATH, backup)
    print("backup:", backup)

    head_para(doc, 2, MARKER)

    head_para(doc, 3, "One decision grammar across the three money surfaces")
    doc.add_paragraph(
        "The Command Center v3 watchlist card, broker-proposal card, and stop-management expanded row "
        "now share one visual and semantic grammar, rendered from a single primitives module "
        "(components/primitives/cardPrimitives.tsx). Each surface answers, in its first two lines, "
        "what the item is, what the operator should do, and why; supporting detail lives in a 2x2 "
        "module grid and evidence lives behind expanders. Color is signal-only: teal = actionable, "
        "amber = caution, red = hard defect, green = price numerals. All three formats are "
        "operator-LOCKED: data wiring and bug fixes only, no unprompted layout changes.")
    doc.add_paragraph(
        "Each card's tinted banner is driven by a priority-ordered decision matrix (first match wins). "
        "Watchlist: private -> no-stop -> incoherent plan (target<=limit) -> price drift >15% -> thin R:R "
        "-> stale/doubt -> CIO-avoid -> READY. Proposals: expired/rejected STALE -> gate/oversight BLOCKED "
        "-> price-drift FIX -> oversized FIX -> earnings CAUTION -> reviews-pending WAIT -> READY. "
        "Stops: NO STOP -> broker-looser REVIEW -> earnings <=7d -> stale-advisory REVIEW -> PROTECTED/MONITORED. "
        "Phase A extracted the primitives with a Playwright pixel-diff regression against the locked "
        "watchlist card (zero differing pixels).")

    head_para(doc, 3, "Cross-surface trust and intelligence context")
    doc.add_paragraph(
        "Decision context that previously lived only on the watchlist card now rides every surface: "
        "held-position awareness (a proposal that adds to an existing position says so, in amber), "
        "earnings proximity (amber within 7 days; on the stop surface it reads as gap risk - a stop "
        "does not protect through an earnings print), catalyst and news headlines, LLM-advisory age "
        "flags (stop advisory >48h and holdings health >72h render stale warnings), and an exit ladder "
        "for held longs (T1 = +1R off the effective stop, T2 = Street mean, T3 = Street high, stop to "
        "breakeven at +1R). Sizing is unified on percent-of-equity position caps: the per-account "
        "policy number (account_automation_policies.max_position_allocation_pct) is served to the UI "
        "and enforced identically by the card sizing table, the propose modal (hard block, no "
        "override), and the backend generator; the generator additionally down-shifts risk on wide "
        "stops (>7% -> x0.75, >10% -> x0.5 via tilt, recorded in sizing_basis).")

    head_para(doc, 3, "Plan and proposal freshness enforcement")
    doc.add_paragraph(
        "Stale levels are now detected and repaired rather than displayed: watchlist cards flag "
        "incoherent plans (target at or below limit) and price drift beyond 15% as FIX states with a "
        "Rebuild Plan action; a nightly re-validator (plan_drift_revalidator.py, weekdays 17:25) "
        "re-plans defective entries in subprocess batches; and the proposal sweeper auto-rejects "
        "PENDING proposals older than 4 hours whose live price has drifted more than 15% from the "
        "proposed entry. Holdings changes trigger narrative re-synthesis and health refreshes "
        "automatically (holdings_change_trigger.py), with the synthesis maturity gates "
        "(final_synthesis_status AND analysis_stage) reset on every re-queue path so re-run agents "
        "always re-synthesize.")

    head_para(doc, 3, "Artifacts")
    t = doc.add_table(rows=len(ARTIFACTS), cols=3)
    borders(t)
    for i, row in enumerate(ARTIFACTS):
        for j, cell in enumerate(row):
            t.rows[i].cells[j].text = cell

    doc.add_paragraph(
        "Delivered 2026-07-03/04 across PRs #51-#81 (all squash-merged, CI-green, live-verified). "
        "Design spec artifact: claude.ai/code/artifact/fe9d8c64-5440-4178-8e7f-332e46396bde. "
        "Advisory boundary unchanged throughout: approve/2FA flows, gate evaluation, and stop "
        "submission machinery were not modified by the redesign.")

    doc.save(DOCX_PATH)

    doc2 = Document(DOCX_PATH)
    assert any(MARKER in (p.text or "") for p in doc2.paragraphs), "marker missing after save"
    print(f"appended + verified: {MARKER}")


if __name__ == "__main__":
    main()

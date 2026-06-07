#!/usr/bin/env python3
"""v1.8 — Append 'Hermes Global Profile Architecture Update — 2026-06-06' to the canonical Reference
Architecture DOCX. APPEND-ONLY, marker-guarded/idempotent. Does NOT remove or rewrite existing content."""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Hermes Global Profile Architecture Update — 2026-06-06"

ROWS = [
    ("Profile", "Path", "Model", "Status", "Safety Boundary"),
    ("default", "~/.hermes", "gemma3:4b", "active/general", "no live/system claims without tools"),
    ("tradeai", "~/.hermes/profiles/tradeai", "gemma3:4b", "stable Trade AI advisory", "no trades/orders/stops/proposals/secrets"),
    ("tradeai12b", "~/.hermes/profiles/tradeai12b", "gemma3:12b-ctx4k", "experimental", "advisory-only; not promoted"),
    ("dev", "~/.hermes/profiles/dev", "unset", "future", "Codex/dev only; human-invoked; no broker secrets"),
    ("serverops", "~/.hermes/profiles/serverops", "unset", "future", "controlled ops only; advisory until configured"),
    ("old sidecar", "hermes_sidecar/.hermes + install", "legacy", "retired/rollback", "not canonical runtime"),
]


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def main():
    doc = Document(DOCX_PATH)
    if any(MARKER in (p.text or "") for p in doc.paragraphs):
        if any(COMPLETION_MARKER in (p.text or "") for p in doc.paragraphs):
            print("v1.8 section + completion addendum already present, skip")
            return
        append_completion(doc)
        doc.save(DOCX_PATH)
        print("completion addendum appended (main v1.8 section already present)")
        return

    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MARKER

    doc.add_paragraph(
        "Hermes has been promoted from a Trade AI sidecar-only install to a global profile-based assistant "
        "layer on ms01-openclaw. The canonical Hermes runtime now lives under the user-level global install "
        "(~/.local/bin/hermes, ~/.local/share/hermes-agent-venv, ~/.hermes; Hermes Agent v0.16.0) and exposes "
        "separate profiles for general use, Trade AI advisory review, experimental 12B review, future "
        "development/Codex work, and future controlled server operations.")
    doc.add_paragraph(
        "Trade AI no longer treats the old hermes_sidecar/.hermes and hermes_sidecar/install directories as "
        "canonical runtime. After operator-approved Stage D rename-retirement (2026-06-06, rename-only, no "
        "deletion) they survive as hermes_sidecar/.hermes.RETIRED_20260606_2140 and install.RETIRED_20260606_2140, "
        "retained only as rollback/audit artifacts; the sidecar wrappers are now retirement stubs. The canonical "
        "Trade AI Hermes entry point is the restricted tradeai profile.")
    doc.add_paragraph(
        "The stable Trade AI profile uses gemma3:4b with tools disabled. The experimental tradeai12b profile "
        "uses gemma3:12b-ctx4k and remains advisory-only and unpromoted. gemma3:12b without the context gate is "
        "not approved as the default Trade AI model. qwen3:14b must not be reintroduced as a Hermes default. "
        "Codex is reserved for the future dev profile as a human-invoked engineering assistant only and is not "
        "approved as autonomous Trade AI runtime.")
    doc.add_paragraph(
        "Hermes profile separation is now part of the safety boundary: Trade AI advisory review, future Codex "
        "development, and future server operations must not share a single unrestricted agent identity. All "
        "local Ollama profiles run with tools disabled; tool-enabled development belongs in the future dev/Codex "
        "path only after explicit operator approval.")

    t = doc.add_table(rows=len(ROWS), cols=5)
    borders(t)
    for ri, row in enumerate(ROWS):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "References"
    doc.add_paragraph(
        "Detail and evidence: docs/hermes/HERMES_GLOBAL_INSTALL_MIGRATION_20260606.md; "
        "docs/hermes/HERMES_PROFILE_MATRIX_20260606.md; docs/hermes/HERMES_MODEL_CANARY_STATUS_20260606.md; "
        "docs/hermes/HERMES_SIDECAR_RETIREMENT_PLAN_20260606.md; "
        "docs/hermes/HERMES_CURATED_MIGRATION_INVENTORY_20260606.md.")

    append_completion(doc)
    doc.save(DOCX_PATH)
    print("v1.8 Hermes Global Profile Architecture section + completion addendum appended + saved")


COMPLETION_MARKER = "Hermes Migration Completion Addendum — 2026-06-06"


def append_completion(doc):
    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = COMPLETION_MARKER
    doc.add_paragraph(
        "The Hermes sidecar-to-global migration is now fully complete. Beyond the rename-retirement recorded "
        "above: the old hermes-gateway.service was found still active/enabled and was stopped and disabled "
        "(operator-approved); the recreated runtime directory was re-retired; previously-tracked sidecar "
        "runtime/state files were removed from Git tracking and gitignored (preserved on disk in the "
        ".RETIRED_* directories and backup tarballs); a launch-path audit confirmed no process, systemd unit, "
        "timer, cron job, shell startup, or PATH command can relaunch the sidecar; and the status scripts "
        "(scripts/api_v2.py, scripts/check_system_versions.sh) were repointed from the retired sidecar to the "
        "global Hermes install (~/.local/share/hermes-agent-venv, ~/.hermes). The disabled hermes-gateway.service "
        "unit file is retained as an inactive audit artifact, not an active launch path. The operator's live "
        "interactive chat sessions were preserved throughout. See HERMES_SIDECAR_RETIREMENT_PLAN_20260606.md "
        "(Stage D, Git Hygiene, Launch Path Audit, Status Path Repoint sections).")


if __name__ == "__main__":
    main()

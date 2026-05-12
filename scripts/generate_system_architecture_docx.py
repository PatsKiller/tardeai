#!/usr/bin/env python3
"""Generate SYSTEM_ARCHITECTURE_COMPLETE.docx from the markdown source.
Uses python-docx to produce a professional Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import re

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__)) + "/.."
MD_PATH = os.path.join(PROJECT_ROOT, "docs/project/SYSTEM_ARCHITECTURE_COMPLETE.md")
DOCX_PATH = os.path.join(PROJECT_ROOT, "docs/project/Trade_AI_v12_System_Architecture.docx")


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))


def shade_cells(row, color="D9E2F3"):
    """Apply background shading to header row cells."""
    from docx.oxml.ns import qn
    from lxml import etree
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = etree.SubElement(tcPr, qn('w:shd'))
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')


def main():
    with open(MD_PATH, 'r') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # Title page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Trade AI v12\nComplete System Architecture")
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    title_run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(
        "Comprehensive Technical Reference for Senior Architects\n\n"
        "Version 12.30  |  2026-05-12  |  Internal Technical Reference\n\n"
        "Covers: Architecture, Components, Interfaces, Trading Workflows,\n"
        "Risk Management, Automation, Integrations, and Observability"
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # Table of Contents placeholder
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_p.add_run("TABLE OF CONTENTS")
    toc_run.bold = True
    toc_run.font.size = Pt(14)
    doc.add_paragraph()

    # Extract section headers for TOC
    toc_entries = []
    for line in lines:
        line = line.rstrip()
        if line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            toc_entries.append(title)

    for i, entry in enumerate(toc_entries):
        p = doc.add_paragraph()
        run = p.add_run(f"{entry}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_page_break()

    # Parse markdown and convert to docx
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    table_alignments = None

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                # Separator row — skip
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            # Check if next line is not a table row
            if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                # Flush table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=1, cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    add_table_borders(tbl)

                    # Header row
                    hdr = tbl.rows[0]
                    shade_cells(hdr)
                    for ci, val in enumerate(table_rows[0]):
                        if ci < num_cols:
                            cell = hdr.cells[ci]
                            cell.text = val
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                                run.font.size = Pt(9)

                    # Data rows
                    for row_data in table_rows[1:]:
                        row = tbl.add_row()
                        for ci, val in enumerate(row_data):
                            if ci < num_cols:
                                row.cells[ci].text = val
                                for run in row.cells[ci].paragraphs[0].runs:
                                    run.font.size = Pt(9)

                    doc.add_paragraph()  # spacing after table
                table_rows = []
                in_table = False
            i += 1
            continue

        # Skip the title line (already on title page)
        if line.startswith('# ') and 'Trade AI v12' in line:
            i += 1
            continue

        # Version/date/classification line
        if line.startswith('**Version:**'):
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Headings
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            i += 1
            continue

        if line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Bold-italic-emphasis lines (like **Output:** or **Note:**)
        # Regular paragraph text
        text = line.strip()

        # Clean up markdown formatting for docx
        p = doc.add_paragraph()

        # Handle bullet points
        if text.startswith('- ') or text.startswith('* '):
            p.style = 'List Bullet'
            text = text[2:]
        elif re.match(r'^\d+\.\s', text):
            p.style = 'List Number'
            text = re.sub(r'^\d+\.\s', '', text)

        # Split text into bold/normal segments
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10.5)
            else:
                # Handle inline code
                code_parts = re.split(r'(`[^`]+`)', part)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`'):
                        run = p.add_run(cp[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x80, 0x30, 0x30)
                    else:
                        run = p.add_run(cp)
                        run.font.size = Pt(10.5)

        i += 1

    # Final page — document info
    doc.add_page_break()
    doc.add_heading("Document Information", level=2)

    info_tbl = doc.add_table(rows=1, cols=2)
    add_table_borders(info_tbl)
    shade_cells(info_tbl.rows[0])
    info_tbl.rows[0].cells[0].text = "Field"
    info_tbl.rows[0].cells[1].text = "Value"
    for run in info_tbl.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
    for run in info_tbl.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True

    info_data = [
        ("Document Title", "Trade AI v12 — Complete System Architecture"),
        ("Version", "12.30"),
        ("Date", "2026-05-12"),
        ("Author", "Trade AI System + Claude Code"),
        ("Classification", "Internal Technical Reference"),
        ("Audience", "Senior Architects, System Designers"),
        ("Source", "docs/project/SYSTEM_ARCHITECTURE_COMPLETE.md"),
        ("DOCX Output", "docs/project/Trade_AI_v12_System_Architecture.docx"),
        ("Supersedes", "All prior partial architecture documentation"),
        ("Covers Automated Trading", "Yes — full pipeline from signal discovery through "
         "autonomous execution, R-multiple trailing stops, critical news auto-close, "
         "phantom detection, and post-trade learning loop"),
    ]
    for field, value in info_data:
        row = info_tbl.add_row()
        row.cells[0].text = field
        row.cells[1].text = value

    doc.save(DOCX_PATH)
    print(f"Generated: {DOCX_PATH}")
    print(f"Source: {MD_PATH}")


if __name__ == "__main__":
    os.chdir(PROJECT_ROOT)
    main()

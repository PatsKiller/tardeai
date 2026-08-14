#!/usr/bin/env python3
"""Render the Trade AI Institutional Report v2 to PDF (Chromium) and DOCX (python-docx).

Reads the live /api/v2/cio/report-v2 model (fetched to /tmp by default, or a path
passed as argv[1]) and emits:
  exports/cio_institutional_report_v2.pdf
  exports/cio_institutional_report_v2.docx

Read-only. Produces advisory artifacts only.
"""
import json
import pathlib
import sys

REPO = pathlib.Path(__file__).resolve().parent.parent
EXPORTS = REPO / "exports"


def _fmt(v, digits=2, suffix=""):
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{digits}f}{suffix}"
    except (TypeError, ValueError):
        return str(v)


def _pct(v):
    if v is None:
        return "—"
    try:
        return f"{float(v):.2f}%"
    except (TypeError, ValueError):
        return str(v)


def _rows_to_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def build_docx(model: dict, out: pathlib.Path) -> None:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    doc.core_properties.title = "Trade AI Institutional Report v2"
    doc.core_properties.author = "Trade AI Investment Office (Alex, CIO)"

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    # ── Title ──────────────────────────────────────────────────────────
    t = doc.add_heading("Trade AI — Institutional Report", level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run(
        f"{model.get('report_version', '')}  ·  {model.get('authority', '')}  ·  "
        f"as-of {model.get('as_of', '')}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    src = (model.get("manifest") or {}).get("source_sha")
    if src:
        p = doc.add_paragraph()
        pr = p.add_run(f"Source SHA: {src}")
        pr.font.size = Pt(8)
        pr.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    pa = model.get("part_a") or {}
    pb = model.get("part_b") or {}

    # ── Part A: CIO Letter ─────────────────────────────────────────────
    letter = pa.get("letter") or {}
    doc.add_heading("Part A — CIO Investment Committee", level=1)

    doc.add_heading("CIO Letter", level=2)
    if letter.get("stance"):
        doc.add_paragraph(f"Stance: {letter['stance']}")
    if letter.get("risk_posture"):
        doc.add_paragraph(f"Risk posture: {letter['risk_posture']}")

    if letter.get("priorities"):
        doc.add_paragraph("Priorities:").runs[0].bold = True
        for item in letter["priorities"]:
            doc.add_paragraph(item, style="List Bullet")
    if letter.get("what_changed"):
        doc.add_paragraph("What changed:").runs[0].bold = True
        for item in letter["what_changed"]:
            doc.add_paragraph(item, style="List Bullet")
    if letter.get("what_not_to_do"):
        doc.add_paragraph("Guardrails (what not to do):").runs[0].bold = True
        for item in letter["what_not_to_do"]:
            doc.add_paragraph(item, style="List Bullet")

    # ── Decisions Now ──────────────────────────────────────────────────
    decs = pa.get("decisions_now") or []
    doc.add_heading("Decisions Now", level=2)
    if decs:
        _rows_to_table(
            doc,
            ["Symbol", "Stance", "Value (USD)", "Weight", "Delta (USD)", "Why now", "Risk"],
            [
                [
                    d.get("symbol", "—"),
                    d.get("stance", "—"),
                    _fmt(d.get("current_value_usd")),
                    _pct(d.get("current_weight_pct")),
                    _fmt(d.get("recommended_delta_usd")),
                    d.get("why_now", "—"),
                    d.get("risk", "—"),
                ]
                for d in decs
            ],
        )
    else:
        doc.add_paragraph("No actionable decisions this cycle.")

    # ── Capital Plan ───────────────────────────────────────────────────
    cp = pa.get("capital_plan") or {}
    doc.add_heading("Capital Plan", level=2)
    band = cp.get("cash_policy_band") or {}
    cp_lines = [
        ("Cash total", _fmt(cp.get("cash_total_usd"))),
        ("Cash reserved", _fmt(cp.get("cash_reserved_usd"))),
        ("Cash investable", _fmt(cp.get("cash_investable_usd"))),
        ("Policy band", f"{_pct(band.get('min_pct'))} – {_pct(band.get('max_pct'))}"),
        ("Recommended deploy", _fmt(cp.get("recommended_deploy_usd"))),
        ("Recommended raise", _fmt(cp.get("recommended_raise_usd"))),
        ("Post-plan cash", _fmt(cp.get("post_plan_cash_usd"))),
    ]
    _rows_to_table(doc, ["Metric", "Value"], cp_lines)

    if cp.get("sources") or cp.get("uses"):
        doc.add_paragraph("Sources of funds:").runs[0].bold = True
        for k, v in (cp.get("sources") or {}).items():
            doc.add_paragraph(f"{k}: {_fmt(v)}", style="List Bullet")
        doc.add_paragraph("Uses of funds:").runs[0].bold = True
        for k, v in (cp.get("uses") or {}).items():
            doc.add_paragraph(f"{k}: {_fmt(v)}", style="List Bullet")

    # ── Portfolio Posture ──────────────────────────────────────────────
    pp = pa.get("portfolio_posture") or {}
    doc.add_heading("Portfolio Posture", level=2)
    top = pp.get("top_position") or {}
    rh = pp.get("risk_heat") or {}
    posture_lines = [
        ("Top position", f"{top.get('symbol', '—')} ({_pct(top.get('weight_pct'))})"),
        ("Concentration fire %", _pct(pp.get("concentration_fire_pct"))),
        ("Max drawdown", _pct(rh.get("max_drawdown_pct"))),
        ("Sharpe", _fmt(rh.get("sharpe"), digits=3)),
        ("Sortino", _fmt(rh.get("sortino"), digits=3)),
    ]
    _rows_to_table(doc, ["Metric", "Value"], posture_lines)

    if pp.get("sector_posture"):
        doc.add_paragraph("Sector posture:").runs[0].bold = True
        _rows_to_table(
            doc,
            ["Sector", "State", "Exposure", "Target", "Recommendation"],
            [
                [
                    s.get("sector", "—"),
                    s.get("state", "—"),
                    _pct(s.get("exposure_pct")),
                    _pct(s.get("target_pct")),
                    s.get("recommendation", "—"),
                ]
                for s in pp["sector_posture"]
            ],
        )

    # ── Opportunity Funnel ─────────────────────────────────────────────
    of = pa.get("opportunity_funnel") or {}
    doc.add_heading("Opportunity Funnel", level=2)
    if of.get("sector_opportunities"):
        doc.add_paragraph("Sector opportunities:").runs[0].bold = True
        for s in of["sector_opportunities"]:
            doc.add_paragraph(
                f"{s.get('sector', '—')} ({s.get('state', '—')}): {s.get('recommendation', '—')}",
                style="List Bullet",
            )
    if of.get("research_gaps"):
        doc.add_paragraph("Research gaps:").runs[0].bold = True
        for g in of["research_gaps"]:
            doc.add_paragraph(f"{g.get('symbol', '—')} ({g.get('sector', '—')})", style="List Bullet")

    # ── Counter-Thesis / Risks ─────────────────────────────────────────
    ct = pa.get("counter_thesis_risks") or {}
    doc.add_heading("Counter-Thesis & Risks", level=2)
    if ct.get("highest_impact_unknowns"):
        doc.add_paragraph("Highest-impact unknowns:").runs[0].bold = True
        for u in ct["highest_impact_unknowns"]:
            doc.add_paragraph(u, style="List Bullet")

    # ── Part B: Portfolio Book ─────────────────────────────────────────
    doc.add_heading("Part B — Portfolio Book", level=1)
    pf = pb.get("portfolio") or {}
    book_lines = [
        ("Total portfolio value", _fmt(pf.get("total_value"))),
        ("Cash", _fmt(pf.get("cash_value"))),
        ("Cash %", _pct(pf.get("cash_pct"))),
        ("Positions", pf.get("positions_count", "—")),
    ]
    _rows_to_table(doc, ["Metric", "Value"], book_lines)

    accounts = pb.get("accounts") or []
    if accounts:
        doc.add_heading("Accounts", level=2)
        _rows_to_table(
            doc,
            ["Account", "Broker", "Value", "Weight", "Gain/Loss", "Status"],
            [
                [
                    a.get("display_name") or a.get("account_id", "—"),
                    a.get("broker", "—"),
                    _fmt(a.get("total_value")),
                    _pct(a.get("weight_pct")),
                    _fmt(a.get("gain_loss")),
                    a.get("status", "—"),
                ]
                for a in accounts
            ],
        )

    alloc = pb.get("allocation") or {}
    if alloc:
        doc.add_heading("Asset Allocation", level=2)
        # Phase 3: model allocation is USD dollars — never format $ as weight %.
        weights = pb.get("allocation_weight_pct") or {}
        if not weights:
            try:
                from scripts.lib.cio_decision_semantics import (
                    allocation_weights_from_usd, looks_like_dollar_allocation,
                )
                if looks_like_dollar_allocation(alloc):
                    weights = allocation_weights_from_usd(alloc)
            except Exception:
                weights = {}
        if weights:
            rows = []
            for k, w in weights.items():
                usd = alloc.get(k)
                try:
                    usd_s = f"${float(usd):,.2f}" if usd is not None else "—"
                except (TypeError, ValueError):
                    usd_s = "—"
                rows.append((k, usd_s, _pct(w)))
            _rows_to_table(doc, ["Class", "Value (USD)", "Weight"], rows)
        else:
            _rows_to_table(doc, ["Class", "Weight"], [(k, _pct(v)) for k, v in alloc.items()])

    perf = pb.get("performance") or {}
    if perf:
        doc.add_heading("Performance", level=2)
        pr_lines = [
            ("YTD return", _pct(perf.get("ytd_return"))),
            ("Inception return", _pct(perf.get("inception_return"))),
            ("Portfolio CAGR", _pct(perf.get("port_cagr"))),
            ("Benchmark CAGR", _pct(perf.get("bench_cagr"))),
            ("Alpha (annualized)", _pct(perf.get("alpha_annualized"))),
            ("Sharpe", _fmt(perf.get("sharpe"), digits=3)),
            ("Sortino", _fmt(perf.get("sortino"), digits=3)),
            ("Max drawdown", _pct(perf.get("max_drawdown"))),
        ]
        _rows_to_table(doc, ["Metric", "Value"], pr_lines)

    # ── Coverage / Checkpoint ──────────────────────────────────────────
    cov = model.get("coverage") or {}
    ck = model.get("checkpoint") or {}
    doc.add_heading("Data Coverage & Provenance", level=1)
    cov_lines = [
        ("Fields tracked", cov.get("field_count", "—")),
        ("Source traceability", _pct(cov.get("source_traceability_pct"))),
        ("Fields present", len(cov.get("fields_present") or [])),
        ("Improved vs. reference", len(cov.get("fields_improved_vs_reference") or [])),
        ("Explicitly unavailable", len(cov.get("fields_unavailable") or [])),
        ("Quality flags", len(cov.get("quality_flags") or [])),
    ]
    _rows_to_table(doc, ["Metric", "Value"], cov_lines)

    if cov.get("fields_unavailable"):
        doc.add_paragraph("Unavailable fields:").runs[0].bold = True
        for f in cov["fields_unavailable"]:
            doc.add_paragraph(str(f), style="List Bullet")

    # ── Disclosure ─────────────────────────────────────────────────────
    doc.add_heading("Disclosure", level=1)
    doc.add_paragraph(
        "This report is generated by the Trade AI investment-office automation in an "
        "advisory-only capacity. No broker, order, or stop authority is exercised. "
        "Figures are composed from canonical portfolio state and are not investment advice."
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main() -> int:
    src = sys.argv[1] if len(sys.argv) > 1 else "/tmp/cio_report_v2_model.json"
    model = json.load(open(src))
    if "data" in model and "report_version" not in model:
        model = model["data"]
    out = build_docx(model, EXPORTS / "cio_institutional_report_v2.docx")
    print(f"DOCX bytes: {out.stat().st_size}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

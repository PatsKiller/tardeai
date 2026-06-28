#!/usr/bin/env python3
"""Append the Social → Momentum Scalp lifecycle hardening section to the Reference Architecture
DOCX. APPEND-ONLY, idempotent (marker-guarded), python-docx only — never raw XML. A timestamped
.bak is written before saving. As-built 2026-06-28 (PRs #7 + #8, merges fbe2048e-era branch)."""
import shutil
from datetime import datetime, timezone

from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Social → Momentum Scalp lifecycle hardening (2026-06-28)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)

    doc.add_paragraph(
        "A focused pass closed the lifecycle gaps in the intraday Momentum Scalp strategy and the "
        "Social Scalp discovery path that feeds it, landed on main through pull requests seven and "
        "eight with the release-readiness CI green. The work changed no trading behaviour and "
        "crossed no live broker write path; social-only signals remain advisory and can never become "
        "auto-tradeable without deterministic confirmation, large-language models remain advisory "
        "only, and the operator confirmation and two-factor path was treated as immutable throughout. "
        "The independently-scored lifecycle maturity rose from roughly four to 4.4 out of five — the "
        "engineering earns a raw weighted five, but the score is honestly capped at 4.4 because the "
        "empirical validation sample is still accumulating: Momentum Scalp stays in TESTING until it "
        "has at least thirty closed paper trades over six months, and the score does not pretend "
        "otherwise.")

    _head(doc, 3, "One source of truth, enforced before approval")
    doc.add_paragraph(
        "The strategy's intraday execution block is now the single source of truth for the trading "
        "window, the proposal time-to-live, the fast-path account, and the maximum price drift; a "
        "config validator fails the build whenever another block contradicts it, so the float ceiling, "
        "the entry window, and the thirty-minute proposal life can no longer drift apart. The "
        "automated approver consults that thirty-minute time-to-live before every approval decision: a "
        "scalp proposal older than its window is expired and recorded as such, and can never be "
        "approved, while the legacy four-hour rule survives only as a fallback for non-intraday "
        "proposals. A scalp whose freshness cannot even be established is blocked rather than approved. "
        "Liquidity is treated the same way — for this liquidity- and time-sensitive strategy a missing, "
        "stale, or errored quote now defers the proposal instead of failing open, and an operator force "
        "override is always logged, never silent.")

    _head(doc, 3, "Social discipline, explicit routing, and end-to-end lineage")
    doc.add_paragraph(
        "On the discovery side, alerting now keys off the final, capped decision rather than the raw "
        "score, so a social-only surge that has been downgraded for lack of a verified catalyst can no "
        "longer masquerade as an actionable signal or mirror itself into the proposals channel. A "
        "deterministic routing policy sends each social candidate to exactly one destination — "
        "watch-only, the micro-cap momentum scalp, a manual-review meme-squeeze path, the portfolio "
        "agents, or rejection — and an unverified catalyst can never reach an actionable route. Every "
        "social candidate now carries a stable discovery trace identifier that is threaded the whole "
        "length of the pipeline, from the social scan through the harmonised scan record, the strategy "
        "signal, the proposal, and finally the paper trade, so a single candidate's journey can be "
        "reconstructed end to end. A funnel report, a transparent capped maturity score, and a bounded "
        "advisory outcome-learning loop turn that lineage into evidence — and none of it can unlock "
        "execution or loosen a deterministic gate.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    bak = f"{DOCX_PATH}.bak_scalp_lifecycle_{stamp}"
    shutil.copy2(DOCX_PATH, bak)
    print(f"backup written: {bak}")
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

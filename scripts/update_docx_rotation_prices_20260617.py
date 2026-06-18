#!/usr/bin/env python3
"""Append the Rotation Intelligence live-prices / review-quantities addendum to the Reference Architecture
DOCX. APPEND-ONLY, idempotent (marker-guarded). Mirrors update_docx_* protocol.
As-built 2026-06-17; full detail in docs/project/ROTATION_LLM_ADVISOR.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation Intelligence - live prices + advisory review quantities (2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "The rotation summary now attaches a live price and an advisory review share quantity to every idea "
        "and candidate, so the operator can see what a trim or add would actually look like. This stays "
        "advisory: read-only data, no broker call, no order, and no live network request from the summary "
        "request itself. Quantities are review ranges derived from the advisory dollar range; nothing is "
        "sized or placed.")
    doc.add_paragraph(
        "Prices come from the latest market_quotes row per symbol (populated by the existing quote-ingest "
        "pipeline: Alpaca, then finviz, then yfinance) with a fallback to the price and share count already "
        "in the holdings snapshot. Each rebalance idea gains a from price, a to price, the held share count, "
        "and a sell share range and buy share range computed by dividing the dollar review range by each "
        "leg's price. The card reads, for example, thirty-one ninety-three times four hundred three shares "
        "held into six sixty-eight, with two chips: review trimming roughly two hundred to six hundred "
        "shares of the source, and approximately one to three thousand shares of the target. The review and "
        "research candidates gain a price and a colored day change, and held review candidates also show an "
        "estimated share count. Symbols with no quote, such as a retirement-plan fund code, simply omit the "
        "price rather than fabricate one. The language stays advisory throughout - review trimming and "
        "approximately - never sell now or buy now.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation prices addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation prices addendum appended + saved")


if __name__ == "__main__":
    main()

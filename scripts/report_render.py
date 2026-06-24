"""report_render.py — single HTML/CSS source of truth → paginated PDF (Playwright) + DOCX (python-docx).

One section model (the analyst_report_builder JSON) renders to:
  render_html(report) -> str        Jinja2 template + assets/analyst_report.css
  render_pdf(report, path)          headless Chromium (CSS Paged Media, running header/footer, inline charts)
  render_docx(report, path)         styled python-docx from the same section model

WeasyPrint/Pandoc are intentionally NOT required (blocked by sudo in this env); Playwright + python-docx
are the in-environment renderers. Charts are inlined as base64 data URIs so the HTML is self-contained.
"""
from __future__ import annotations

import base64
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES = PROJECT_ROOT / "templates"
ASSETS = PROJECT_ROOT / "assets"
REPORT_OUT = PROJECT_ROOT / "data" / "portfolios" / "reports" / "analyst"

# Sections that earn a compact KPI table beneath their prose (prose-first; no blanket dumps).
_KPI_TABLE_SECTIONS = {
    "header_context", "executive_summary", "fundamental_valuation", "analyst_predictions",
    "analyst_commentary", "options_income", "risk_assessment", "peer_comparison", "action_plan",
    "earnings_estimates", "fundamentals_deep", "valuation_context", "scenario_targets",
    "catalyst_risk", "tax_position", "portfolio_fit",
}
_COVER_KPIS = (
    ("recommendation", "REC"), ("price", "PRICE"), ("day_change_pct", "DAY"),
    ("confidence_label", "CONF"), ("thesis_status", "THESIS"),
    ("unrealized_pnl_pct", "UNREALIZED"), ("portfolio_pct", "WEIGHT"),
)


def _f(v: Any, d: float = 0.0) -> float:
    try:
        return float(v) if v is not None else d
    except (TypeError, ValueError):
        return d


def _md(text: Any) -> str:
    """Strip markdown emphasis (**bold**/*italic*) — the HTML template autoescapes, so raw ** would show."""
    import re
    s = str(text if text is not None else "")
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", s)
    return s


def _fmt(key: str, val: Any) -> str:
    if val is None:
        return "—"
    k = key.lower()
    if isinstance(val, bool):
        return "Yes" if val else "No"
    if isinstance(val, (int, float)):
        if k == "confidence" and val <= 1:
            return f"{val * 100:.0f}%"
        if k == "analysts" or k.endswith("_contracts") or k == "n":
            return f"{int(val)}"
        if k.endswith("_pct"):
            return f"{val:+.2f}%"
        if any(t in k for t in ("price", "target", "valid_", "value", "pnl", "mean", "low", "high",
                                "bear", "base", "bull", "gain", "tax", "cost", "loss", "benefit")) and k != "case":
            return f"${val:,.2f}" if abs(val) < 100000 else f"${val:,.0f}"
        if k in ("pe", "forward_pe", "p_fcf"):
            return f"{val:.1f}×"
        if k in ("peg", "debt_equity", "beta"):
            return f"{val:.2f}"
        if k in ("reward_risk",):
            return f"{val:.1f}:1"
        return f"{val:,.2f}".rstrip("0").rstrip(".") if abs(val) < 1000 else f"{val:,.0f}"
    return str(val)


def _rel_url_safe(path: Path) -> str:
    try:
        return "/" + str(Path(path).resolve().relative_to(PROJECT_ROOT)).replace("\\", "/")
    except ValueError:
        return str(path)


def _resolve_png(chart_path: str | None) -> Path | None:
    if not chart_path:
        return None
    p = Path(str(chart_path))
    if p.exists():
        return p
    s = str(chart_path).lstrip("/")
    cand = PROJECT_ROOT / s
    if cand.exists():
        return cand
    for base in (REPORT_OUT / "charts",):
        c = base / p.name
        if c.exists():
            return c
    return None


def _data_uri(chart_path: str | None) -> str | None:
    png = _resolve_png(chart_path)
    if not png:
        return None
    try:
        b = png.read_bytes()
        return "data:image/png;base64," + base64.b64encode(b).decode()
    except Exception:
        return None


def _prepare(report: dict) -> dict:
    """Build the template context: cover KPI tiles, sections w/ inlined figures + formatted KPI tables."""
    meta = report.get("meta") or {}
    kpis = meta.get("kpis") or {}
    overlay = (meta.get("claude_oversight") or {})

    cover_tiles = []
    for key, label in _COVER_KPIS:
        v = kpis.get(key)
        cls = ""
        if key.endswith("_pct") and isinstance(v, (int, float)):
            cls = "pos" if v >= 0 else "neg"
        cover_tiles.append({"label": label, "value": _fmt(key, v), "cls": cls})

    sections = []
    for s in (report.get("sections") or []):
        sid = s.get("id")
        figs = []
        for fig in (s.get("figures") or []):
            uri = _data_uri(fig.get("chart_path"))
            if uri:
                figs.append({"uri": uri, "caption": fig.get("caption") or ""})
        rows = []
        if sid in _KPI_TABLE_SECTIONS:
            for k, v in (s.get("metrics") or {}).items():
                if v is None or k in ("text", "what_to_do_now"):
                    continue
                rows.append({"label": str(k).replace("_", " ").title(), "value": _fmt(k, v)})
        # single callout overlay (Senior Analyst Overlay) — dedupe by label
        callouts, seen = [], set()
        for c in (s.get("callouts") or []):
            lab = str(c.get("label") or "")
            if lab in seen:
                continue
            seen.add(lab)
            callouts.append({"label": lab, "text": str(c.get("text") or "")})
        sections.append({
            "id": sid,
            "title": s.get("title") or sid,
            "content": _md(s.get("content") or ""),
            "callouts": [{"label": c["label"], "text": _md(c["text"])} for c in callouts],
            "kpi_rows": rows[:8],
            "bullets": [_md(b).lstrip("• ") for b in (s.get("bullets") or [])][:5],
            "agents": s.get("agents") or [],
            "peer_table": s.get("peer_table") or [],
            "figures": figs,
        })

    return {
        "title": meta.get("title") or f"{meta.get('symbol', '')} Analyst Report",
        "symbol": meta.get("symbol") or "",
        "company": meta.get("company") or "",
        "sector": meta.get("sector") or "",
        "generated_at": str(meta.get("generated_at") or datetime.now(timezone.utc).isoformat())[:19].replace("T", " "),
        "version": meta.get("version") or "4.0",
        "doc_class": meta.get("document_class") or "summary_prospectus",
        "cover_tiles": cover_tiles,
        "sections": sections,
        "oversight": {
            "verdict": overlay.get("verdict"),
            "fixes_applied": overlay.get("fixes_applied"),
            "model": overlay.get("model"),
        },
        "sources": report.get("sources") or [],
    }


def render_html(report: dict) -> str:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES)),
        autoescape=select_autoescape(["html"]),
    )
    css = (ASSETS / "analyst_report.css").read_text() if (ASSETS / "analyst_report.css").exists() else ""
    tmpl = env.get_template("analyst_report.html.j2")
    return tmpl.render(css=css, **_prepare(report))


def render_pdf(report: dict, output_path: Path) -> dict:
    """Paginated PDF via headless Chromium with running header/footer + page numbers."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    html = render_html(report)
    meta = report.get("meta") or {}
    sym = meta.get("symbol") or ""
    gen = str(meta.get("generated_at") or "")[:19].replace("T", " ")
    header = (
        '<div style="font-size:7px;width:100%;padding:0 14mm;color:#7a8699;'
        'font-family:Inter,Helvetica,Arial,sans-serif;display:flex;justify-content:space-between;">'
        f'<span>Trade AI v12 · Analyst Research</span><span>{sym}</span></div>'
    )
    footer = (
        '<div style="font-size:7px;width:100%;padding:0 14mm;color:#7a8699;'
        'font-family:Inter,Helvetica,Arial,sans-serif;display:flex;justify-content:space-between;">'
        f'<span>Produced by TradeAI v3.0 · Generated {gen} · Advisory — not investment advice</span>'
        '<span>Page <span class="pageNumber"></span> of <span class="totalPages"></span></span></div>'
    )
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox"])
            page = browser.new_page()
            page.set_content(html, wait_until="networkidle")
            page.pdf(
                path=str(output_path), format="Letter", print_background=True,
                display_header_footer=True, header_template=header, footer_template=footer,
                margin={"top": "16mm", "bottom": "14mm", "left": "0mm", "right": "0mm"},
            )
            browser.close()
    except Exception as e:
        return {"ok": False, "error": f"playwright pdf failed: {str(e)[:200]}"}
    return {"ok": True, "format": "pdf", "path": str(output_path), "engine": "playwright",
            "size_kb": round(output_path.stat().st_size / 1024, 1)}


def render_docx(report: dict, output_path: Path) -> dict:
    """Styled DOCX from the same section model (python-docx — pandoc-free path)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ctx = _prepare(report)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    h = doc.add_heading(ctx["title"], level=0)
    sub = doc.add_paragraph()
    sub.add_run("Trade AI v12 · Analyst Research\n").bold = True
    sub.add_run(f"{ctx['company']} · {ctx['sector']}\n")
    sub.add_run(f"Generated {ctx['generated_at']} · v{ctx['version']} · Advisory — not investment advice\n")

    # KPI band as a 2-row table
    tiles = [t for t in ctx["cover_tiles"] if t["value"] != "—"]
    if tiles:
        t = doc.add_table(rows=2, cols=len(tiles))
        t.style = "Light Grid Accent 1"
        for i, tile in enumerate(tiles):
            t.rows[0].cells[i].text = tile["label"]
            t.rows[1].cells[i].text = tile["value"]
    doc.add_paragraph("")
    doc.add_heading("Contents", level=1)
    for i, s in enumerate(ctx["sections"], 1):
        doc.add_paragraph(f"{i}. {s['title']}", style="List Number")
    doc.add_page_break()

    for s in ctx["sections"]:
        doc.add_heading(s["title"], level=1)
        for co in s["callouts"]:
            p = doc.add_paragraph()
            p.add_run(f"{co['label']}: ").bold = True
            p.add_run(co["text"]).font.color.rgb = RGBColor(0x0f, 0x6d, 0x3a)
        if s["content"]:
            doc.add_paragraph(s["content"])
        if s["kpi_rows"]:
            tb = doc.add_table(rows=1 + len(s["kpi_rows"]), cols=2)
            tb.style = "Light Grid Accent 1"
            tb.rows[0].cells[0].text = "Metric"; tb.rows[0].cells[1].text = "Value"
            for i, r in enumerate(s["kpi_rows"], 1):
                tb.rows[i].cells[0].text = r["label"]; tb.rows[i].cells[1].text = r["value"]
        for b in s["bullets"]:
            doc.add_paragraph(b.lstrip("• "), style="List Bullet")
        if s["agents"]:
            at = doc.add_table(rows=1 + len(s["agents"]), cols=3)
            at.style = "Light Grid Accent 1"
            for j, lbl in enumerate(("Agent", "Rec", "Weight")):
                at.rows[0].cells[j].text = lbl
            for i, ag in enumerate(s["agents"], 1):
                at.rows[i].cells[0].text = str(ag.get("agent") or "—")
                at.rows[i].cells[1].text = str(ag.get("recommendation") or "—")
                at.rows[i].cells[2].text = str(ag.get("weight") or ag.get("relevance") or "—")[:40]
        # embed figures via original section paths (data-URIs are for HTML; docx wants the file)
        orig = next((x for x in report.get("sections", []) if x.get("id") == s["id"]), {})
        for fig in (orig.get("figures") or []):
            p = _resolve_png(fig.get("chart_path"))
            if p:
                try:
                    doc.add_picture(str(p), width=Inches(6.3))
                    if fig.get("caption"):
                        cap = doc.add_paragraph(fig["caption"]); cap.runs[0].italic = True
                        cap.runs[0].font.size = Pt(8)
                except Exception:
                    pass

    foot = doc.add_paragraph("— End of Report · Produced by TradeAI v3.0 · Advisory, not investment advice —")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.save(str(output_path))
    return {"ok": True, "format": "docx", "path": str(output_path),
            "size_kb": round(output_path.stat().st_size / 1024, 1)}

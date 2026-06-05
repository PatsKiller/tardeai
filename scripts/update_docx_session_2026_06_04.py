#!/usr/bin/env python3
"""Append Session 2026-06-04 addenda to Reference Architecture DOCX. APPEND-ONLY,
idempotent (each section marker-guarded independently):
  1. Premarket proposal freshness alignment + proposal-alert dedup persistence fix.
  2. RTH-gate on intraday proposal generation (supersedes the 'still open' note in #1)."""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Premarket Proposal Freshness Alignment & Alert Dedup Fix (Session 2026-06-04)"
RTH_MARKER = "RTH-Gate on Intraday Proposal Generation (Session 2026-06-04)"
CATALYST_MARKER = "Catalyst Pipeline Repair — news to catalyst to fusion (Session 2026-06-04)"
BRIDGE_MARKER = "Hermes News Bridge + Tiered Catalyst Cadence (Session 2026-06-04)"
MONITOR_MARKER = "System Freshness Monitor + Staleness Scan (Session 2026-06-04)"
SENTIMENT_MARKER = "Sentiment Lane Repoint + Dormant-Subsystem Decisions (Session 2026-06-04)"
FAILTEST_MARKER = "Watchdog Fail-Test PASSED + Device Delivery Confirmed (Session 2026-06-04)"
HEARTBEAT_MARKER = "Watch-the-Watchman Dead-Man's Switch (Session 2026-06-04)"
RESEARCH_MARKER = "Research Lane Revived — fusion 4 of 5 (Session 2026-06-04)"
TOPICENG_MARKER = "topic_ingestion Root-Cause + Fix + Cadence + Monitor (Session 2026-06-04)"
LIQUIDITY_MARKER = "Liquidity Pre-Screen at Proposal Generation (Session 2026-06-04)"
BTFILTER_MARKER = "v3 Backtesting Tab Respects Shared Filters (Session 2026-06-04)"


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def append_main(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MARKER

    doc.add_paragraph(
        "Context: on 2026-06-03/04, with auto-trade active, premarket paper proposals (ARTL "
        "recovery_watch, FOFO momentum_scalp) showed ACTIONABLE_READY and re-alerted on Telegram "
        "every 2 minutes. No order executed — the execution guards held — but two defects in the "
        "proposal/alert layer were confirmed and fixed. The approver correctly rejected a manual "
        "FOFO approval with 'quote is 1030.9min old. Need fresh market data.'")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Readiness/executor freshness alignment (intraday-only tightening)"
    doc.add_paragraph(
        "proposal_execution_readiness.py relaxed the quote-age ceiling to 24h for ALL strategy "
        "classes outside regular hours. For intraday strategies (momentum_scalp, gap_and_go) this "
        "let a stale premarket quote surface as ACTIONABLE_READY, only for the approver to reject "
        "it — the executor (paper_trade_logger.py) enforces a flat 15-min ceiling "
        "(max_quote_age_minutes=15, no caller override) with no after-hours relaxation. Fix: the "
        "after-hours relaxation is now gated to swing/position only (is_intraday guard from "
        "get_timeframe_class); intraday keeps its 300s ceiling year-round so readiness can no "
        "longer show a green card the executor will reject. Swing/position after-hours pre-staging "
        "on a day-old quote is unchanged (per operator decision to keep pre-staging).")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Proposal-alert dedup persistence fix"
    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Aspect", "Detail"
    for c, d in [
        ("Symptom", "Identical proposal cards re-sent every 2 minutes (e.g. ARTL ~10x). The 30-min "
                    "per-proposal dedup keyed on paper_trade_proposals.last_alert_at never fired."),
        ("Root cause", "send_telegram_proposal_alert.py _db_query never committed (db_adapter "
                       "singleton conn, autocommit=False) and ran the dedup UPDATE through the "
                       "fetchall() path, which raises on a no-result statement; the bare except "
                       "swallowed it and left the shared transaction aborted (poisoning later "
                       "queries). last_alert_at was never persisted -> dedup always read NULL."),
        ("Fix", "_db_query gained a fetch='none' write mode that commits and returns rowcount, plus "
                "rollback-on-error; the last_alert_at UPDATE now uses fetch='none'."),
        ("Known limitation", "Dedup is purely time-based: a genuine ACTIONABLE -> BLOCKED transition "
                             "within the 30-min window is also suppressed until the window elapses or "
                             "the proposal auto-rejects after 30-min blocked. State-aware dedup is a "
                             "possible follow-up."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d

    doc.add_paragraph(
        "Still open (scoped out by operator on 2026-06-04): no regular-hours gate on proposal "
        "GENERATION/alert crons (they fire from 7:30 AM premarket). See MASTER lifecycle note and "
        "alert dispatcher section. [SUPERSEDED — see the RTH-Gate section below; the generation gate "
        "was implemented later the same day.]")


def append_rth(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = RTH_MARKER

    doc.add_paragraph(
        "Implements the 'still open' generation gate noted in the freshness-alignment section above. "
        "The freshness fix stopped the false ACTIONABLE_READY green light, but intraday proposals "
        "were still GENERATED premarket (04:00-09:00) on stale quotes, so they expired or "
        "auto-rejected (price-drift / blocked-too-long) before the 9:30 open — a real GO signal "
        "(XOS, FOFO on 2026-06-04) never became an approvable proposal. Investigation: today's 6 GO "
        "names produced 12 proposals (IDs 165-176), all EXPIRED (10) or REJECTED (XOS "
        "auto_stale_price_drift_22pct; FOFO auto_blocked_32min); 0 paper_trades created that day.")

    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Aspect", "Detail"
    for c, d in [
        ("Where", "auto_proposal_generator.py -> run_auto_proposals(), first check in the per-signal "
                  "loop. Single gate covers both the direct */30 9-16 cron AND catalyst_momentum_engine.py "
                  "premarket_scalp band, which shells out to auto_proposal_generator.py."),
        ("Rule", "If strategy is intraday (get_timeframe_class == 'intraday'; momentum_scalp, gap_and_go) "
                 "AND current_market_session() != 'regular', skip with decision SKIPPED_OUTSIDE_RTH "
                 "(recorded in auto_proposal_decisions). Regenerates on the next regular-hours run on "
                 "live quotes."),
        ("Scope", "Intraday only. Swing/position generation unchanged (after-hours pre-staging kept). "
                  "Incubator promoter unaffected — its _CLASSIFICATION_STRATEGIES excludes intraday, so "
                  "it never emits momentum_scalp/gap_and_go."),
        ("Override", "force=True (manual operator runs) bypasses the gate for on-demand off-hours "
                     "generation."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d


def append_catalyst(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = CATALYST_MARKER

    doc.add_paragraph(
        "Root cause (STEP 0): catalyst_events was never a scheduled pipeline — all 345 rows came "
        "from a single manual run on 2026-04-27, and news_to_catalyst.py was never in any crontab. "
        "news_ingestion.py (scheduled 3x/day) DID try to write catalyst_events inline but used "
        "columns title/relevance_score that do not exist (table has headline/impact_score), so the "
        "INSERT failed on every article inside a SAVEPOINT rollback + except:pass — a silent "
        "failure inside a green job. Downstream, signal_fusion.py read empty catalyst_events "
        "(catalyst_score=0) and itself went dormant ~2026-05-11 (fused_signals stopped); "
        "cio_decision_engine reads fused_signals via LEFT JOIN so it degraded to NULL with no "
        "error. No health agent or Hermes check covers data-table write-health, so it hid ~5 weeks. "
        "See docs/HERMES_NEWS_TO_SCALP_CATALYST_INTEGRATION_2026_06_04.md.")

    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Change", "Detail"
    for c, d in [
        ("Dedup index", "UNIQUE INDEX uniq_catalyst_symbol_headline ON catalyst_events(symbol, headline); 0 pre-existing dups."),
        ("Writer A fix", "news_ingestion.py inline INSERT: corrected columns, classifies via news_to_catalyst._classify, ON CONFLICT (symbol,headline) DO NOTHING."),
        ("Writer B fix", "news_to_catalyst.py INSERT: ON CONFLICT (symbol,headline) DO NOTHING + None-guard (both writers run in parallel, deduped)."),
        ("Fusion re-wire", "signal_fusion.py --full scheduled (dormant since 2026-05-11); no code change needed."),
        ("Staleness monitor", "NEW intel_table_staleness_monitor.py: SIEM (alert_events data_integrity, deduped) if an intel table gets 0 rows/24h while its input is fresh. Tracks catalyst_events<-news_articles, fused_signals<-catalyst_events."),
        ("Cron", "news_to_catalyst 45 6,12,18 *; signal_fusion --full 0 7,13 * 1-5 (timeout 20m); staleness monitor 15 9 * 1-5 --send."),
        ("Proof", "news_to_catalyst created 487 catalyst_events (345 existing unchanged); NOC catalyst_score 0.9 vs ZZZZ 0; fused_signals max date 2026-05-11 -> 2026-06-04."),
        ("Follow-ups", "keyword classifier under-matches (~94% 'other'); fusion confidence x impact_score over-weights weak catalysts; signal_fusion --full heavy at 2,714 symbols; catalyst_sentiment_analysis still empty."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d


def append_bridge(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = BRIDGE_MARKER
    doc.add_paragraph(
        "Prompt #2 (APPLIED): scripts/hermes_news_bridge.py bridges Hermes ticker-level news "
        "(hermes_research_intelligence momentum_catalyst) into news_articles (source='hermes'), "
        "after which the repaired chain classifies it into catalyst_events and signal_fusion "
        "consumes it. READ-ONLY on hermes_* (Hermes wall intact); dedup tracked TradeAI-side via "
        "raw_payload.hermes_research_id. Proven: 28 bridged, ABTS traced hermes#302 -> news_articles "
        "-> catalyst_events -> signal_fusion catalyst_score 0.9; re-run 0 (dedup). Cron 40 6,12,18.")
    doc.add_paragraph(
        "Prompt #3 (DESIGN ONLY — NO CRONS ENABLED, pending approval): the live scalp engine uses "
        "SearXNG (local, no API quota), so an urgent scalp tier driven by bridge + news_to_catalyst "
        "(DB-internal) is quota-free. Proposed: SCALP 04:00-12:00 every 10 min over the tight 38 "
        "(screener GO/WAIT) or 93 (incubator score>=40) set — NOT the 2,747 watchlist; PROPOSALS+ATM "
        "(3 symbols) every 15-30 min; WATCHLIST+HELD 2-3x/day (existing signal_fusion --full). "
        "External catalyst_enrichment (NewsAPI ~100/day, AlphaVantage ~25/day) cannot drive a 5-min "
        "tier — one pass over 38 tickers blows those quotas; keep it low-cadence on cheap APIs only. "
        "See docs/HERMES_NEWS_TO_SCALP_CATALYST_INTEGRATION_2026_06_04.md §13-14.")


def append_monitor(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MONITOR_MARKER
    doc.add_paragraph(
        "STEP 0 system-wide staleness scan (read-only) found the catalyst death was NOT unique: a "
        "family of identical silent failures — code inserting into tables whose columns were renamed "
        "by a migration, failing inside green jobs (savepoint + except:pass), plus unscheduled jobs. "
        "Death clusters ~2026-04-27, ~05-10, ~05-11. The scan also caught an incomplete earlier fix "
        "(catalyst Writer A untested + a twin sentiment_observations column-bug).")
    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Item", "Detail"
    for c, d in [
        ("Phase A fix", "sentiment_observations twin column-bug fixed in news_ingestion; Writer A verified end-to-end (both catalyst writers now parallel + deduped, verified)."),
        ("Monitor", "system_freshness_monitor.py — registry-driven (fresh / empty-vs-input / cron-logfile checks), weekday-aware, severity by blast radius."),
        ("Escalation", "SIEM (alert_events data_integrity, deduped 12h) for all; Telegram for P0/P1 out-of-band so the next silent death is loud within ~20 min."),
        ("Auto-fix (narrow)", "Allowlist {news_to_catalyst, hermes_news_bridge} only — DB-only, idempotent, dedup-guarded, capped 2/day, ALWAYS logged + escalated. NEVER schema/column/trading writes (root cause needs a human)."),
        ("Cron", "*/20 * * * * --send --auto-fix, flock-guarded; supersedes intel_table_staleness_monitor.py."),
        ("Left for operator", "learning batch (frozen 05-11), social_mentions/catalyst_sentiment_analysis (empty, feed signal_fusion), market_ohlcv_bars (05-07) — revive vs deprecate is a human decision, not auto-revived."),
        ("Trust framing", "Verified both ways: 'all fresh' on healthy system; forced-stale entry fires. Trust should come from watching it page on the next real failure, not this doc."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d


def append_sentiment(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = SENTIMENT_MARKER
    doc.add_paragraph(
        "Read-only revive-vs-deprecate review of the 3 dormant subsystems from the staleness scan. "
        "Facts corrected two initial steers; net result is one small build, the rest left off (high "
        "bar to wake a subsystem at low trust). BUILD: signal_fusion read catalyst_sentiment_analysis "
        "(no producer -> sentiment_score stuck at 0.5 default); repointed to sentiment_observations "
        "(same overall_sentiment/confidence columns, now flowing). One line, reversible. Verified: RTX "
        "sentiment_score 0.5 -> 0.567, sentiments 0 -> 3; fusion now on 3 of 5 lanes (was 2). "
        "DECISIONS: social_mentions = deprecate (no producer; fusion reweights for it). "
        "market_ohlcv_bars = deprecate/defer (producer + strategy consumers all dormant; only a "
        "rarely-hit volume fallback in proposal_execution_readiness is live). Learning batch = stays "
        "asleep (it is operator-approved propose-with-governance, not a blind auto-tuner, but the bar "
        "to wake auto-tuning at 1.25 trust is not met; revive insight/shadow-only if ever, never "
        "auto-apply). See docs/HERMES_NEWS_TO_SCALP_CATALYST_INTEGRATION_2026_06_04.md §17.")


def append_failtest(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = FAILTEST_MARKER
    doc.add_paragraph(
        "The key verification: a forced failure pushed through the monitor's real run() path proved "
        "the whole escalation chain rather than trusting it. A clearly-labeled [FAIL-TEST] synthetic "
        "P1 finding was injected; test artifacts (SIEM row + auto-fix counter) were cleaned afterward. "
        "Results — detection fired; SIEM written (P1/urgent, id 466); auto-fix attempted and escalated "
        "EVEN ON SUCCESS (auto_remediation in payload); cap halts at 2 -> 'cap-reached, escalate to "
        "human'; Telegram accepted by API for both operator chats; and the operator CONFIRMED the "
        "message arrived on-device (verbatim). The */20 cron is confirmed firing (ran 12:00, flock "
        "held, logged). This is the proof point, not an assurance: the out-of-band link whose absence "
        "hid the catalyst death ~5 weeks is now demonstrably working. STILL OPEN (deliberate next "
        "task): watch-the-watchman — the monitor cannot detect its own death; needs an independent "
        "dead-man's switch (separate cron on its log freshness; ideally an off-host external "
        "heartbeat). See docs/HERMES_NEWS_TO_SCALP_CATALYST_INTEGRATION_2026_06_04.md §18.")


def append_heartbeat(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = HEARTBEAT_MARKER
    doc.add_paragraph(
        "Answers the recursive question 'who watches the watchman' — the freshness monitor cannot "
        "detect its own death. Solved in two complementary layers. IN-HOST (no external dependency): "
        "system_freshness_monitor writes logs/.freshness_monitor.heartbeat only on a completed run; a "
        "SEPARATE, self-contained cron freshness_watchdog_heartbeat.py (*/30, deliberately does not "
        "import the monitor so a broken monitor can't break its watcher) pages P0 if the heartbeat is "
        ">70 min stale (~3 missed */20 cycles). Verified: fresh->OK; backdated 90m->STALE + P0 SIEM "
        "(id 469); restored->OK. Catches monitor-down-while-host-alive. OFF-HOST (terminus): if env "
        "FRESHNESS_HEARTBEAT_PING_URL is set, the monitor pings an external uptime service each run; "
        "if pings stop (total-host death) that service pages — the only layer surviving the whole box "
        "going down. Env-gated; setting the URL is the one remaining operator step for full coverage. "
        "Honest limit: the in-host checker is itself unwatched; the off-host ping is the terminating "
        "layer. See docs/HERMES_NEWS_TO_SCALP_CATALYST_INTEGRATION_2026_06_04.md §18.")


def append_research(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = RESEARCH_MARKER
    doc.add_paragraph(
        "Operator-approved revive of the last dormant fusion lane. research_insights was frozen "
        "2026-04-27 because its producer (research_insight_extractor.py) was never scheduled — NOT a "
        "missing producer (unlike social_mentions). The extractor is keyword-based (extraction_model "
        "keyword_v1, no LLM/API cost), DB-only, idempotent (NOT EXISTS dedup), and reads news_articles "
        "+ catalyst_events (both now flowing). Scheduled 50 6,12,18. Verified: 200 new insights (200 "
        "existing preserved, no mutation); SCHD research_score 0.24 (5 inputs); signal_fusion now runs "
        "on 4 of 5 lanes (catalyst, news, sentiment, research; only deprecated social dark). Added to "
        "system_freshness_monitor registry (11 entries) and the safe-auto-fix allowlist (DB-only "
        "idempotent re-run). Dormant-subsystem tally: research revived, sentiment built, social + "
        "market_ohlcv_bars deprecated, learning batch asleep. See "
        "docs/SYSTEM_HEALTH_BASELINE_2026_06_04.md.")


def append_topiceng(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = TOPICENG_MARKER
    doc.add_paragraph(
        "Discovered during Research Topic Registry STEP 0: topic_ingestion.py (topic_monitor content "
        "aggregation -> news/YouTube -> RAG) was dead ~25 days (topic_monitor.last_searched frozen "
        "2026-05-10, 0/17 topics searched in 7d, 85 data_staleness topic alerts unactioned). ROOT "
        "CAUSE: the cron passed '--all', which the script's argparse does not accept -> every run "
        "exited immediately ('unrecognized arguments: --all') doing nothing. Confirmed by the log, "
        "the archived crontabs (--all), and job_coverage_monitor.py citing 'topic_ingestion.py --all' "
        "as its canonical 'scheduled but produces nothing' example. The live cron had already been "
        "corrected to '--use-llm-queries' (valid) but never re-run. FIX/VERIFY: proven with a single "
        "topic (ssdi: 26 articles, last_searched unstuck 05-10 -> today; youtube_api + google_news_rss "
        "carried it, Brave depleted but keyless fallbacks covered); full 17-topic run launched (processing one-by-one; daily cron backstops the rest). "
        "CADENCE: bumped 2x/week (Wed/Sat) -> daily (0 9 * * *) — YouTube/Google-RSS quota fine, "
        "Brave-independent. MONITOR: topic freshness added to system_freshness_monitor (12 entries; "
        "last_searched check via new per-entry ts_col; detect+escalate only — NO auto-fix, since "
        "topic_ingestion hits external APIs and is not the DB-only idempotent class. So the next stall "
        "pages instead of joining unactioned SIEM rows.")


def append_liquidity(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = LIQUIDITY_MARKER
    doc.add_paragraph(
        "Closes the upstream half of the 2026-06-04 proposal-noise problem. The RTH-gate + "
        "freshness fixes stopped premarket false-greens, but the screener still fed structurally "
        "untradeable microcaps that reached Telegram and were only rejected downstream by the "
        "execution-readiness layer — pure noise (FOFO live spread 19.8% / 8.5K shares; BNKK 15% / "
        "1.6K shares). Root cause: strategy_signals carries no absolute volume or spread, only rvol "
        "(a ratio — 22x a tiny base is still tiny), float_m and price; and no stored source is "
        "reliable for thin names (trade_ai_scans.volume NULL, market_quotes has no row for FOFO/BWEN). "
        "A live quote is the only dependable liquidity source, so the gate reads one at generation.")
    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Aspect", "Detail"
    for c, d in [
        ("Where", "auto_proposal_generator.py run_auto_proposals() step 2e — after strategy-criteria "
                  "validation, before sizing (only quotes candidates that passed the cheap structural "
                  "filters). New _liquidity_prescreen(symbol, rules) helper."),
        ("Rule", "Regular-hours only. Block if live spread_pct > max_spread_pct (5.0), or if BOTH the "
                 "share floor (min_day_volume_shares 25000) AND the dollar floor (min_dollar_day_volume "
                 "$100K) are breached. Records SKIPPED_LIQUIDITY in auto_proposal_decisions + "
                 "liquidity_rejected stat."),
        ("Thresholds", "config/strategies/shared_risk_rules.yaml -> liquidity_prescreen (enabled, "
                       "max_spread_pct, min_day_volume_shares, min_dollar_day_volume). Looser than the "
                       "execution layer's ~1% momentum spread on purpose — generation gate catches "
                       "structurally-untradeable, approval stays the tighter backstop."),
        ("Fail-open", "After-hours / no quote / provider error / config-disabled => NEVER blocks. "
                      "Swing/position after-hours pre-staging unaffected; force=True bypasses. Never "
                      "block on bad data."),
        ("Verified", "Blocks FOFO/BNKK/XOS on live quotes, passes AAPL; all four fail-open paths "
                     "confirmed; dry-run integration clean (liquidity_rejected stat wired)."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d


def append_btfilter(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = BTFILTER_MARKER
    doc.add_paragraph(
        "Dashboard bug (v3 Journal hub): the Backtesting tab ignored the shared account chips and "
        "time-range selection that drive every other Journal tab. JournalHub.tsx rendered "
        "<BacktestPanel onDrill> passing NONE of its filter state (the code comment even claimed an "
        "'all-accounts filter built in'), so the panel ran on its own independent filters and never "
        "reacted to the operator's account/range choice. Fix: JournalHub now passes "
        "sharedAccount={acctFilter} + sharedDateFrom={cutoff}; BacktestPanel seeds its internal "
        "account/dateFrom state from those props and re-syncs on change, driving its backend reload. "
        "The account keys line up exactly — strategy_backtest_trades.account stores the same "
        "normalized values JournalHub uses (alpaca_paper, schwab_rollover_ira, schwab_taxable) — so "
        "no translation is needed; the time range maps to the backend start_date. Selecting an "
        "account narrows to that account's replays (account-agnostic champion sims show only under "
        "'All'); schwab_roth_ira has no backtest rows so it correctly shows empty. Verified: tsc+vite "
        "build clean; live API (:7777) /api/v2/backtesting/trades returns all=48, alpaca_paper=11, "
        "schwab_rollover_ira=28, alpaca_paper since 2026-05-01=10.")


def main():
    doc = Document(DOCX_PATH)
    changed = False
    if not _has(doc, MARKER):
        append_main(doc)
        changed = True
        print("main addendum appended")
    if not _has(doc, RTH_MARKER):
        append_rth(doc)
        changed = True
        print("RTH-gate section appended")
    if not _has(doc, CATALYST_MARKER):
        append_catalyst(doc)
        changed = True
        print("catalyst-repair section appended")
    if not _has(doc, BRIDGE_MARKER):
        append_bridge(doc)
        changed = True
        print("bridge+tiered section appended")
    if not _has(doc, MONITOR_MARKER):
        append_monitor(doc)
        changed = True
        print("monitor section appended")
    if not _has(doc, SENTIMENT_MARKER):
        append_sentiment(doc)
        changed = True
        print("sentiment-lane section appended")
    if not _has(doc, FAILTEST_MARKER):
        append_failtest(doc)
        changed = True
        print("fail-test section appended")
    if not _has(doc, HEARTBEAT_MARKER):
        append_heartbeat(doc)
        changed = True
        print("heartbeat section appended")
    if not _has(doc, RESEARCH_MARKER):
        append_research(doc)
        changed = True
        print("research-lane section appended")
    if not _has(doc, TOPICENG_MARKER):
        append_topiceng(doc)
        changed = True
        print("topic_ingestion section appended")
    if not _has(doc, LIQUIDITY_MARKER):
        append_liquidity(doc)
        changed = True
        print("liquidity pre-screen section appended")
    if not _has(doc, BTFILTER_MARKER):
        append_btfilter(doc)
        changed = True
        print("backtest-filter section appended")
    if changed:
        doc.save(DOCX_PATH)
        print("saved")
    else:
        print("all sections present, skip")


if __name__ == "__main__":
    main()

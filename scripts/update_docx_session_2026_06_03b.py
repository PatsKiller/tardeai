#!/usr/bin/env python3
"""Append Session 2026-06-03b addendum (server hang fix + dashboard accuracy + v3 additions)
to Reference Architecture DOCX. APPEND-ONLY, idempotent (marker-guarded)."""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Server Hang Fix, Dashboard Accuracy & v3 Additions (Session 2026-06-03b)"


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def main():
    doc = Document(DOCX_PATH)
    for p in doc.paragraphs:
        if MARKER in (p.text or ""):
            print("present, skip")
            return

    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MARKER

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Critical: Portfolio Server hang + prevention"
    doc.add_paragraph(
        "The portfolio server deadlocked (process alive, serving nothing; /v2/ and /v3/ unreachable). "
        "Root cause: portfolio_server.py called importlib.reload(api_v2) on EVERY /api/v2/ request; under "
        "the v3 dashboard's concurrent polling on the threaded server this caused a reload race / deadlock. "
        "systemd Restart=always does not catch a hang (the process never exits). Fixes: (1) mtime-gated, "
        "lock-guarded reload in portfolio_server.py:_get_api_v2() — reloads only when api_v2.py changes "
        "(also ~10x latency improvement); (2) scripts/portfolio_server_watchdog.sh (cron */2) probes "
        "/api/health and kills the johnclaw-owned pid on repeated failure so systemd respawns it. "
        "See MASTER §18 Failure Modes.")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Agent dashboard accuracy + new v3 surfaces"
    t = doc.add_table(rows=1, cols=2)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text = "Change", "Detail"
    for c, d in [
        ("Agent 'last run' fix", "agents/summary derived last_run only from watchlist_agent_results; iris writes to iris_run_log.ran_at. Added to _AGENT_HOME_TABLES. aegis was already correct (nightly brief). Staleness coloring added as prevention."),
        ("Worker 'no handoffs' clarified", "agent_handoffs is written by synthesis/Alex/auto_research/system, not individual worker agents — drawer copy now says so (not a fault)."),
        ("v3 Trade AI tab", "Ported v2 /v2/trade-ai market-opportunities scanner (GO/WAIT/NO-GO, KPIs, per-ticker scan) into v3 Trading hub."),
        ("v3 Portfolio filters", "Per-account filter chips on Holdings."),
        ("Human Review actions", "Escalation-queue drawer gained navigation links (Inbox, Agent Collaboration) + action explanation. DetailDrawer gained read-only links."),
        ("Drive sync", "Files >1 MB upload raw (no Google Doc conversion) — fixes recurring http2 timeout on a 3.4 MB report."),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text = c, d

    doc.save(DOCX_PATH)
    print("addendum appended")


if __name__ == "__main__":
    main()

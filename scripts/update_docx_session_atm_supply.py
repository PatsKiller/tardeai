#!/usr/bin/env python3
"""Append ATM v1 + Supply Pipeline content to Reference Architecture DOCX.
Uses python-docx APPEND-ONLY operations per DOCX_UPDATE_PROTOCOL.
Session date: 2026-05-22"""

from docx import Document
from docx.shared import Pt, Inches
from lxml import etree
import shutil, os
from datetime import datetime

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    tblPr.append(etree.fromstring(borders))

def add_row(tbl, cells):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = str(text)

def main():
    # Backup
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    backup = f"{DOCX_PATH}.bak_atm_supply_{ts}"
    shutil.copy2(DOCX_PATH, backup)
    print(f"Backup: {backup}")

    doc = Document(DOCX_PATH)
    h1 = get_heading_style(doc, 1)
    h2 = get_heading_style(doc, 2)
    h3 = get_heading_style(doc, 3) or h2

    # ═══════════════════════════════════════════════════════════════════
    # SECTION: ATM v1
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h1
    p.text = "Automated Trade Mode (ATM) v1 (2026-05-22)"

    doc.add_paragraph(
        "ATM v1 replaces the manual /ptapprove Telegram flow with an automated "
        "decision engine that evaluates pending proposals every 15 minutes during "
        "market hours. All existing safety gates (risk gate, execution readiness, "
        "stop-breach, spread, R:R) remain active. ATM adds its own gate layer: "
        "classifier health, position limits, kill switches, and B-1 observation tracking."
    )

    # ATM Architecture
    p = doc.add_paragraph()
    p.style = h2
    p.text = "ATM Architecture"

    doc.add_paragraph(
        "Components: atm_auto_approver.py (cron, */15 9-15 weekdays), "
        "atm_config_manager.py (YAML config with SHA256 hash versioning), "
        "atm_classifier_health.py (per-strategy health scoring from closed trade outcomes), "
        "atm_state (DB: mode, config_hash, last_evaluated_at), "
        "atm_decision_log (DB: per-proposal decision audit trail)."
    )

    doc.add_paragraph(
        "Modes: disabled (no evaluation), dry_run (evaluate + log, no execution), "
        "active (evaluate + execute), paused (temporarily halted with auto-resume). "
        "Mode transitions logged to atm_state_events table."
    )

    # ATM Gate Chain
    p = doc.add_paragraph()
    p.style = h2
    p.text = "ATM Gate Chain (Evaluation Order)"

    tbl = doc.add_table(rows=1, cols=3)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Gate"
    hdr[1].text = "Check"
    hdr[2].text = "Action on Fail"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    gates = [
        ("1. Per-proposal override", "atm_action field (force_approve/reject/skip)", "Execute override"),
        ("2a. Account check", "target_account in enabled_accounts", "Defer"),
        ("2b. B-1 exclusion", "strategy in bucket2 during observation window", "Defer"),
        ("2c. Same-day skip", "strategy in same_day_skip_strategies list", "Defer"),
        ("3. Classifier health", "get_health(strategy) >= min_classifier_health", "Reject"),
        ("4. Position limits", "max_concurrent, max_new_per_day, max_pct_per_trade", "Reject"),
        ("5. Kill switches", "daily_loss_pct per account and aggregate", "Pause ATM"),
    ]
    for gate in gates:
        add_row(tbl, gate)

    # ATM Config
    p = doc.add_paragraph()
    p.style = h2
    p.text = "ATM Configuration (config/atm_config.yaml)"

    doc.add_paragraph(
        "Version-controlled YAML with SHA256 hash tracking. Changes logged to "
        "atm_config_history with old/new config, diff, and backup path. "
        "Key settings: position_limits (max_concurrent: 10, max_new_per_day: 6, "
        "max_pct_per_trade: 1.0), strategy_filter (min_classifier_health, whitelist, blacklist), "
        "kill_switches (daily_loss_pct_hard_pause: 10.0), operating_hours (09:35-15:30 ET), "
        "same_day_skip_strategies (momentum_scalp, gap_and_go), "
        "b1_tracking (observation_end, bucket2_strategies)."
    )

    # Classifier Health
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Classifier Health Scoring"

    doc.add_paragraph(
        "Per-strategy health score based on closed paper trade outcomes in the last 30 days. "
        "Formula: 0.4 * win_rate + 0.3 * sample_factor + 0.3 * r_factor. "
        "Returns 0.0 if fewer than 3 closed trades (cold-start). "
        "get_health_detail() returns score, closed_trades, has_baseline, wins, avg_r. "
        "Dashboard shows '0.00 (no baseline)' for strategies without sufficient data."
    )

    # ATM Dashboard
    p = doc.add_paragraph()
    p.style = h2
    p.text = "ATM Dashboard (/v2/automated-trade-mode)"

    doc.add_paragraph(
        "React component: AutomatedTradeMode.tsx. "
        "Sections: status banner (mode, hash, staleness with market-hours awareness), "
        "per-account cards (positions, new today with ATM/manual breakdown, ghost cards for disabled), "
        "activity tiles (proposals seen, approved, rejected, queue depth), "
        "queue preview (predicted_decision per proposal with color-coded chips), "
        "strategy health table (health score, baseline status, closed trades, eligible, B-1, same-day), "
        "recent decisions table (with first blocker gate), "
        "settings modal (defaults, accounts, global, B-1, same-day tabs)."
    )

    doc.add_paragraph(
        "API endpoints: /api/v2/atm/status (with is_market_hours, next_expected_cycle), "
        "/api/v2/atm/strategy-health (with has_baseline, closed_trades, wins, avg_r), "
        "/api/v2/atm/queue-preview (with predicted_decision, predicted_reason), "
        "/api/v2/atm/decisions, /api/v2/atm/config (GET/POST), /api/v2/atm/mode (POST), "
        "/api/v2/atm/proposal-action (POST)."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION: Supply Pipeline
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h1
    p.text = "Proposal Supply Pipeline Architecture (2026-05-22)"

    doc.add_paragraph(
        "The proposal supply pipeline has two independent paths that feed proposals "
        "to ATM. Both paths terminate at paper_trade_proposals with status=PENDING."
    )

    # Path A
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Path A: Orchestrator Scoring Pipeline"

    doc.add_paragraph(
        "Flow: Finviz screeners → trade_ai_orchestrator.py (scoring) → "
        "scalp_critic_agent.py (LLM review) → strategy_signals → "
        "auto_proposal_generator.py → paper_trade_proposals. "
        "Scoring thresholds: GO >= 40, WAIT >= 30, AVOID < 30. "
        "Critic can downgrade GO→NO_GO or WAIT→NO_GO. "
        "Output: ~5 proposals/day, dominated by momentum_scalp."
    )

    doc.add_paragraph(
        "Orchestrator cron schedule: 0900, 1000, 1200, 1400, 1600, 1730 (weekdays). "
        "Pre-market runs (0400, 0700) via continuous_runner.py with --allow-underfilled. "
        "Screener windows: 18 active screeners across 8 time windows."
    )

    # Path B
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Path B: Incubator Promotion Pipeline"

    doc.add_paragraph(
        "Flow: Finviz screeners → finviz_screener_runner.py (discovery) → "
        "incubator_universe → incubator_llm_screener.py (qwen3:14b grading) → "
        "incubator_proposal_promoter.py → paper_trade_proposals. "
        "Two sub-paths with different thresholds:"
    )

    doc.add_paragraph(
        "Screener sub-path: score >= 38, catalyst_verified OR score >= 45, "
        "days_active >= 1, llm_screen_verdict != DROP. "
        "Classification sub-path (income/dividend/recovery/defense strategies): "
        "score >= 15 (DIVERSITY_SCORE_FLOOR), joined to ticker_strategy_classifications. "
        "Output: ~4 proposals/day across diverse strategies."
    )

    doc.add_paragraph(
        "Promoter runs hourly 7am-5pm (weekdays). "
        "Pre-promoter incubator quote refresh at :45 past hours 7,10,12,13,16. "
        "Risk gate (RiskGate.check()) runs at proposal creation time. "
        "Spread gate, RSI gate, and price floor checks applied before INSERT."
    )

    # Execution Readiness
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Execution Readiness (proposal_execution_readiness.py)"

    doc.add_paragraph(
        "Evaluates proposals for execution eligibility. Thresholds vary by timeframe class:"
    )

    tbl = doc.add_table(rows=1, cols=5)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Timeframe"
    hdr[1].text = "Max Quote Age"
    hdr[2].text = "Max Price Drift"
    hdr[3].text = "Max Spread"
    hdr[4].text = "Min Volume"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    thresholds = [
        ("Intraday", "300s", "2.0%", "1.0%", "100,000"),
        ("Short Swing", "24h", "5.0%", "3.0%", "50,000"),
        ("Medium Swing", "24h", "8.0%", "3.0%", "50,000"),
        ("Position", "24h", "12.0%", "5.0%", "25,000"),
    ]
    for t in thresholds:
        add_row(tbl, t)

    doc.add_paragraph(
        "Readiness states: READY_FOR_PAPER_SUBMIT, BLOCKED_SPREAD, BLOCKED_PRICE_MOVED, "
        "BLOCKED_RISK_GATE, BLOCKED_NO_VOLUME, BLOCKED_NO_QUOTE, BLOCKED_MISSING_TECHNICALS. "
        "Revalidation cron runs every 30 min during market hours."
    )

    # Supply Funnel Metrics
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Supply Funnel Metrics (Baseline 2026-05-22)"

    tbl = doc.add_table(rows=1, cols=3)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Stage"
    hdr[1].text = "Per Day"
    hdr[2].text = "Retention"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    metrics = [
        ("Screener universe", "2,034", "—"),
        ("Screener hits (scanned)", "~1,800", "—"),
        ("Scored (trade_ai_scans)", "~1,380", "77%"),
        ("Score >= 30 (WAIT+)", "~25", "1.8%"),
        ("Score >= 40 (GO)", "~4", "0.3%"),
        ("After scalp critic (net GO)", "~3", "0.2%"),
        ("Auto proposals (Path A)", "~5", "—"),
        ("Incubator promotions (Path B)", "~4", "—"),
        ("Total proposals", "~9", "—"),
        ("Execution ready (READY)", "~3", "~37%"),
    ]
    for m in metrics:
        add_row(tbl, m)

    # Updated Operating Numbers
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Updated Operating Numbers (2026-05-22)"

    tbl = doc.add_table(rows=1, cols=2)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    numbers = [
        ("Active screeners", "18 (across 8 time windows)"),
        ("Orchestrator cron windows", "0900, 1000, 1200, 1400, 1600, 1730"),
        ("Incubator universe", "1,533 active candidates"),
        ("Promotable pool (screener path)", "69 candidates (score >= 38)"),
        ("LLM screen model", "qwen3:14b (upgraded from gemma3:4b)"),
        ("ATM mode", "dry_run"),
        ("ATM classifier_health threshold", "0.0 (cold-start bypass, restore to 0.50)"),
        ("B-1 observation end", "2026-05-25"),
        ("Strategies with health baseline", "0 of 9 (need 3+ closed trades each)"),
        ("Portfolio value", "$1,192,610 / 47 positions"),
    ]
    for n in numbers:
        add_row(tbl, n)

    # Save
    doc.save(DOCX_PATH)
    print(f"Updated: {DOCX_PATH}")
    print(f"Paragraphs: {len(doc.paragraphs)}")

if __name__ == "__main__":
    main()

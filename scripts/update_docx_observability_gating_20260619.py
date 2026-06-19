#!/usr/bin/env python3
"""Append the LLM-health observability + strategy-performance-gate sections to the Reference Architecture
DOCX. APPEND-ONLY, idempotent (marker-guarded). As-built 2026-06-19 (audit Tasks 2 + 4)."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "LLM review health observability + strategy-performance gate (2026-06-19)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)

    _head(doc, 3, "LLM review health observability")
    doc.add_paragraph(
        "A read-only health endpoint surfaces the state of the three free LLM review lanes — local Ollama "
        "gemma, the Grok xAI-OAuth proxy, and the ChatGPT codex-OAuth proxy — by delegating to the existing "
        "lane-availability check rather than duplicating it, and pairs that with the realized quality of the "
        "review corpus, reported as a valid rate by model. The intent is honest observability: a degraded or "
        "expired lane shows up immediately as a lane marked unavailable and as falling corpus validity, so a "
        "silent provider outage can no longer quietly corrupt the review record. As built, all three lanes are "
        "available and the corpus is about ninety-seven percent valid.")

    _head(doc, 3, "Strategy-performance gate on proposal generation")
    doc.add_paragraph(
        "A read-only performance gate sits in front of both proposal sources — the automated generator and the "
        "incubator promoter. A strategy with fewer than five closed paper trades is always eligible (insufficient "
        "data); a strategy with at least ten closed trades whose realized win rate has fallen below twenty-five "
        "percent stops generating new proposals until its edge recovers, and each suppression is written to an "
        "append-only suppression log surfaced on the strategy leaderboard. The gate never modifies a trade or "
        "proposal and fails open on any error. It complements the per-strategy allocation tilt: the tilt shifts "
        "candidate flow and risk budget toward live winners while tightening their position cap, and the gate is "
        "the hard floor that culls a strategy only once it has a real losing record. As built the gate is dormant "
        "— no strategy yet meets the closed-trade and win-rate thresholds — and it self-activates as the paper "
        "sample accumulates, with no further code change required.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

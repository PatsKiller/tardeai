#!/usr/bin/env python3
"""Append the rotation stop-risk-in-engine + clarity + card-freshness addendum to the Reference Architecture
DOCX. APPEND-ONLY, idempotent (marker-guarded). As-built 2026-06-18; detail in ROTATION_LLM_ADVISOR.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation stop-risk in engine + plain-English clarity + card freshness (2026-06-18)"


def h(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "Rotation Intelligence reconciles its valuation engine with stop risk and is clearer to read. The "
        "engine was blind to protective-stop state because the holdings snapshot carries none, so a name "
        "whose stop had broken still scored a trim of zero. The engine now accepts a stops map, which the "
        "summary builds from the latest Aegis nightly thesis, and a triggered, danger, or broken stop adds "
        "forty trim points - applied after the income and tax reductions so a broken stop is not softened. "
        "Deteriorating names now become trim-review items with real scores, and the engine and the stop "
        "signal are reconciled in a single number rather than disagreeing across columns.")
    doc.add_paragraph(
        "Each rebalance card now leads with a plain-English sentence, and the suggested size is "
        "context-aware: a five-to-fifteen-percent trim for a concentration flag, but a reduce-or-close of "
        "half to all of the position for a broken stop, with the wording changed from trim to review "
        "reducing or closing instead of suggesting a nonsensical fraction of a share on a tiny position. The "
        "summary tiles are clickable and filter the review candidates, and the word shares is spelled out.")
    doc.add_paragraph(
        "The unified symbol-card file that the engine and card layer read had no refresh job and went stale, "
        "which is why newly surfaced research candidates showed sector or analyst pending. A new refresh "
        "script, run on a weekday morning cron, refreshes the company profiles for watch-grade symbols and "
        "then materializes the card file from the symbol-cards endpoint. Research candidates now carry a "
        "sector and a description uniformly, with analyst upside where coverage exists and an honest none "
        "where a microcap has no coverage.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation stop-clarity addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation stop-clarity addendum appended + saved")


if __name__ == "__main__":
    main()

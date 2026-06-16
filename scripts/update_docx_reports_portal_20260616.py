#!/usr/bin/env python3
"""Append the Reports Portal + LM Feedback Loop section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-16; full detail in docs/MASTER_SYSTEM_DOCUMENTATION.md §15."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Reports Portal & LM Feedback Loop (2026-06-16)"


def h(doc, lvl):
    # Resolve the heading STYLE object by name (key lookup fails on duplicate style names).
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "Every recurring report the operator receives on Telegram is now also persisted and browsable in the "
        "v3 Reports hub, and the Central Intelligence board can run a local-LLM (and optional Grok) review of "
        "the live signals whose result feeds the self-learning loop. Full detail: "
        "docs/MASTER_SYSTEM_DOCUMENTATION.md section 15 (Notification & Alerting).")

    _head(doc, 3, "Capture at the send source")
    doc.add_paragraph(
        "Recurring reports (Incubator LLM Screen, Auto-Research, EOD Open Trades, Trade AI Critique, Strategy "
        "and Alex weekly reviews, Portfolio and Monthly reports, CIO and Daily Intelligence, overnight and "
        "pre-market briefs, Learning Digest, Stop briefs) were sent to Telegram but never stored, so the Reports "
        "hub could not show them. report_capture.py classify_report() recognizes 20 report headers and capture() "
        "writes each recognized report to a new telegram_outbox store. It runs at the Telegram send chokepoint "
        "telegram_alert._raw_send_telegram(), is best-effort (never raises, never blocks a send), and skips "
        "unrecognized or transient messages that already self-log to notification_log or alert_events so nothing "
        "is duplicated. Nine senders that posted to the Telegram API directly were routed through send_telegram() "
        "so they are captured and FQDN/v3-normalized; their DOCX sendDocument paths stay direct because "
        "send_telegram is text-only.")

    _head(doc, 3, "The portal: four stores, one feed")
    doc.add_paragraph(
        "reports_portal.py unions four stores — notification_log (briefs/alerts), alert_events (SIEM "
        "advisories), telegram_outbox (captured reports), and ai_reports (LLM monthly/weekly history) — into one "
        "categorized, searchable, paginated news-reader feed with full-article rendering and a purge tool. "
        "Seventeen category tabs: Morning Briefs, Digests, Portfolio Briefs, Monthly Reports, Weekly Reviews, "
        "Incubator Screen, Research and Intel, Trade Reports, Trade Critique, Learning Digest, Alerts, "
        "Advisories, Recovery Watch, Dividends, Regime/Rebalance, Paper Trading, and System Health. Monthly and "
        "Weekly tabs pull real history from ai_reports. Served at GET /api/v2/reports/categories, /list, /item "
        "and POST /api/v2/reports/purge; the UI is /v3/reports.")

    _head(doc, 3, "Link integrity — semantically correct, never dead")
    doc.add_paragraph(
        "Report bodies link to dashboard pages. The send-time normalizer notification_url_builder._to_v3 and the "
        "portal _PAGE map resolve every legacy /v2/<slug> to the v3 page that ACTUALLY contains that content, "
        "not merely a valid route: recovery maps to /v3/risk (where the Recovery Watch section lives, not "
        "/v3/retirement which is tax/Roth only), actions maps to /v3/ (the Home Action Inbox), approvals and "
        "proposals map to /v3/trading. Slugs that would resolve to a dead /v3 route are dropped rather than "
        "emitted as broken buttons. All recurring brief slugs were validated to resolve to a real route.")

    _head(doc, 3, "Central Intelligence — live LM review & feedback loop")
    doc.add_paragraph(
        "The operator's 'Agent / Local LM Review & Feedback' widget previously POSTed to a non-existent endpoint "
        "(404, falling back to localStorage). POST /api/v2/agents/intelligence-feedback now runs the local gemma "
        "LLM (llm_lane.generate lane='local') over the operator's question, critique, and the visible signals "
        "and returns the review synchronously. The operator may also opt into the Grok lane (use_grok; the free "
        "OAuth proxy, only when authenticated), and the UI shows the local and Grok reviews side by side. The "
        "submission is persisted to intelligence_feedback and one learning observation per lane is written to "
        "llm_feedback_observations (workflow='central_intelligence_review'), so operator critique of the "
        "intelligence feeds the self-learning loop.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Reports Portal section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Reports Portal & LM Feedback Loop section appended + saved")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Append Session 30c — Automated Journal consolidation to Reference Architecture DOCX."""

from docx import Document
from docx.shared import Pt
import os, shutil

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    add_table_borders(tbl)
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

def main():
    shutil.copy2(DOCX_PATH, DOCX_PATH + ".bak_session30c")
    doc = Document(DOCX_PATH)
    h1 = get_heading_style(doc, 1)
    h2 = get_heading_style(doc, 2)
    h3 = get_heading_style(doc, 3)

    p = doc.add_paragraph()
    p.style = h1
    p.text = "Session 30c: Automated Journal — Unified Trade Intelligence (2026-05-12)"

    doc.add_paragraph(
        "The Paper Journal has been renamed and replaced by the Automated Journal — a single, "
        "comprehensive trade journal for all automated trading accounts. It provides the same "
        "depth, transparency, and auditability as the regular manual journal, while representing "
        "system-driven decisions and Alpaca broker integration in real time."
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Architecture Change"

    doc.add_paragraph(
        "Before: Trade Journal had 5 tabs — Entries, Analytics, Reports, Paper Journal, "
        "Automated Journal (two separate, incomplete views of automated trades).\n\n"
        "After: Trade Journal has 4 tabs — Entries, Analytics, Reports, Automated Journal. "
        "The Automated Journal consolidates all automated trading accounts into one view, "
        "filtered by account. Currently: Alpaca Paper. Extensible to future accounts."
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Automated Journal Content Per Trade"

    doc.add_paragraph(
        "Every trade — open or closed — expands to show the full execution lifecycle. "
        "This replicates all elements from the regular journal plus Alpaca-specific data:"
    )

    add_table(doc, ["Section", "Fields"], [
        ("Position & Execution",
         "Direction, shares, entry price, planned entry, entry slippage, stop loss, "
         "target, current price, dollar size, dollar risk, order type, broker order ID, "
         "broker status, bracket order flag, submitted/filled timestamps"),
        ("Alpaca Real-Time",
         "Live price, market value, cost basis, unrealized P&L, unrealized %, "
         "intraday P&L, intraday %, change today, side (long/short). "
         "Active orders: type, side, stop price, limit price, qty, status, order ID"),
        ("Strategy & Market Context",
         "Strategy ID, account, market regime, VIX at entry, score at entry, "
         "RVOL at entry, float (M), intel readiness, risk gate result, "
         "opened via, logged by, proposal ID"),
        ("Entry Rationale",
         "Strategy logic, catalyst (verified/unverified with text), "
         "risk gate reason codes, execution notes"),
        ("Stop & Limit Logic",
         "Initial stop, planned stop, stop slippage, target 1, target 2, "
         "R-multiple, MFE (max favorable excursion), MAE (max adverse excursion)"),
        ("Execution Log",
         "Full timeline of all lifecycle events: MONITOR_ADJUST_STOP, "
         "MONITOR_TIGHTEN_NEAR_TARGET, MONITOR_CLOSE_TARGET, MONITOR_ADD_STOP, "
         "MONITOR_PHANTOM_CLOSED, IRIS_OUTCOME_WRITEBACK, AEGIS_POST_TRADE_SYNTHESIS, "
         "OUTCOME_LESSON_CAPTURED, PATTERN_CHECK, alerts (NEAR_STOP, NEAR_TARGET, "
         "STALE_TRADE, NEGATIVE_NEWS, CRITICAL_NEWS_CLOSE, EXTENDED_PROFIT). "
         "Each with timestamp, agent name, summary, and payload"),
        ("Exit & Outcome",
         "Exit price, P&L, P&L %, R realized, verdict (WIN/LOSS), "
         "exit reason, closed via, hold time, closed at, "
         "post-trade analyzed flag, Iris curated flag, Aegis summarized flag"),
        ("Journal Review",
         "Setup family/name, timeframe, direction, market regime, catalyst type, "
         "execution quality score (1-5), sizing quality score (1-5), "
         "risk management score (1-5), followed plan flag, well executed flag, "
         "planned R, realized R, confidence before, stress level, "
         "mistake tags, strength tags, lesson learned, review notes, "
         "coach/system notes, entry signals, exit signals"),
    ])

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Alpaca Integration"

    doc.add_paragraph(
        "The Automated Journal pulls real-time data from Alpaca Paper on every page load "
        "(30-second auto-refresh):\n\n"
        "1. GET /v2/positions — live price, market value, cost basis, unrealized P&L, "
        "intraday P&L, change today for each open position\n"
        "2. GET /v2/orders?status=open — active stop and limit orders per symbol\n\n"
        "Alpaca data is merged inline with DB records. A connection status badge "
        "(ALPACA LIVE / ALPACA OFFLINE) shows integration health. If Alpaca is "
        "unreachable, the journal falls back to DB-cached prices."
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Analytics"

    doc.add_paragraph(
        "The Automated Journal header provides aggregate analytics:\n\n"
        "Stats bar: Open count, Closed count, Wins, Losses, Win Rate, Avg R, "
        "Realized P&L, Unrealized P&L.\n\n"
        "Strategy Performance: Per-strategy breakdown showing trade count, total P&L, "
        "and win rate. Allows quick identification of which strategies are performing.\n\n"
        "Closed trades table: Sortable with Iris/Aegis/RAG curation indicators "
        "showing whether each trade has been analyzed, learned from, and indexed."
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "API Endpoint"

    doc.add_paragraph(
        "GET /api/v2/automated-trade-journal?account=ALPACA_PAPER\n\n"
        "Returns:\n"
        "- trades[]: All paper_trades (40+ fields) enriched with execution_log[], "
        "alerts[], journal_reviews[], and alpaca{} real-time data\n"
        "- open[]: Open trades subset\n"
        "- closed[]: Closed trades subset\n"
        "- summary{}: open_count, closed_count, wins, losses, win_rate, avg_r, "
        "total_pnl, unrealized_pnl, by_strategy[]\n"
        "- alpaca_connected: boolean indicating Alpaca API health"
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Design Intent"

    doc.add_paragraph(
        "The Automated Journal is designed to provide the same trust and auditability "
        "as a human-maintained trade journal, but for system-driven decisions. Every "
        "automated action — from entry to stop adjustment to exit — is recorded with "
        "timestamp, agent, and rationale. This allows:\n\n"
        "1. Post-trade review: Did the system follow the plan? Did stops adjust correctly?\n"
        "2. Learning validation: Did Iris capture the right lesson? Did Aegis synthesize correctly?\n"
        "3. Execution quality: Was there slippage? Did the fill match the plan?\n"
        "4. Strategy assessment: Which strategies are profitable? Which need adjustment?\n"
        "5. System debugging: If a trade went wrong, the execution log shows exactly what "
        "the system did, when, and why — no black boxes."
    )

    doc.save(DOCX_PATH)
    print(f"Session 30c appended to {DOCX_PATH}")

if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")
    main()

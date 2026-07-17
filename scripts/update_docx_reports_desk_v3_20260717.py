#!/usr/bin/env python3
"""Append the Reports Desk v3 section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded), python-docx only, timestamped .bak first.
As-built 2026-07-17 (commits a405b3c3..0be5a136)."""
import shutil
from datetime import datetime, timezone
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Reports Desk v3 — one corpus, system rollup, zero garbage (2026-07-17)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "Reports Desk v3 closed the honesty gaps that the v1 ship's own instrumentation exposed. The "
        "governing rule is one corpus per visual group: every rendered number names its data store, and "
        "no two numbers in the same group may come from different stores. The archive quick views now "
        "filter on the server inside the same query that serves the list, and the chip counts are "
        "computed in that same pass, so the class of disagreement where a chip said four hundred "
        "eighty-six open actions beside a list of eighteen is structurally impossible. The severity "
        "headline lives inside the raw-events analytics fold it belongs to, and the indexing policy — "
        "which event classes get archive rows and why — is configuration rendered as a legend, not lore.")
    doc.add_paragraph(
        "The Reports hub gained a fifth tab, System, rendering the whole-system rollup the operator "
        "asked for: pipeline runs with a failure rail, agent analyses with confidence, proposals by "
        "status, closed validation trades with profit and loss, raw alert volumes, research production, "
        "directive activity, and a health strip that reuses the existing health snapshot, data-source "
        "health, and consumption aggregates rather than recomputing them. A nightly job snapshots the "
        "rollup into a daily table for fourteen-day trend sparklines (which stay honestly empty until "
        "three days accrue) and renders a deterministic Daily System Digest into the report catalog and "
        "archive, with a single Telegram line rather than the body.")
    doc.add_paragraph(
        "Advisory hygiene: conversational LLM preambles are stripped deterministically at write time, a "
        "preamble-leak flag was added to the shared research QA lint, the one writer that bypassed lint "
        "now cleans before storing, and sixteen of thirty-one stored advisories were cleaned in place. "
        "Preamble-only output degrades to an explicit research-pending stub instead of rendering filler "
        "as findings. The analyst tab finished its truth pass: the two need-refresh numbers are labeled "
        "with their genuinely different scopes, holdings under one thousand dollars fold into a residual "
        "positions group instead of standing as peers of core holdings, held names display current "
        "holdings vocabulary with the original registry verb preserved in a tooltip, and the dead "
        "acknowledged column was removed because no acknowledgement write path exists. Producer strings "
        "in the analytics band are normalized through a configuration registry with kind chips, and any "
        "unmapped producer renders a raw chip as visible debt.")


def main():
    bak = DOCX_PATH + ".bak." + datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("marker present — idempotent no-op")
        return 0
    shutil.copy2(DOCX_PATH, bak)
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended Reports Desk v3 section (backup: {bak})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

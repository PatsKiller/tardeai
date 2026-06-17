#!/usr/bin/env python3
"""Append the Rotation Intelligence section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-17; full detail in docs/project/ROTATION_LLM_ADVISOR.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation Intelligence (advisory-only, 2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "Rotation Intelligence is a Command Center v3 feature that helps the operator review whether to rotate "
        "exposure between holdings. It is strictly advisory: it never places, cancels, buys, sells, or routes "
        "anything, it calls no broker API, and it uses no paid LLM API and no API keys. Grok runs only on a "
        "free OAuth session. Full detail: docs/project/ROTATION_LLM_ADVISOR.md.")

    _head(doc, 3, "Grounded-first answer")
    doc.add_paragraph(
        "Every question is first answered by a deterministic grounding report built from the real holdings, "
        "the symbol-card intelligence, and the advisory rotation scorer. If grounding shows no model-supported "
        "trim, add, or rotation signal, the answer mode is grounded_no_supported_action and the review range is "
        "reported as unavailable - no numeric trim amount is invented. The local model and Grok are second "
        "opinions only and can never override grounding.")

    _head(doc, 3, "Ask paths")
    doc.add_paragraph(
        "Ask Local returns the grounded review instantly. Run Grok Review gets a free OAuth Grok second opinion "
        "inline through the local OAuth proxy (no API key, no paid API); if the proxy is offline it falls back "
        "to a prompt the operator pastes into Grok manually. Validate with local model and Run Dual Review run "
        "the local model for an extra opinion and can take a few minutes under GPU load. The hardened dual "
        "advisor script itself never calls Grok over an API - the inline call lives only in the API layer "
        "through the free OAuth proxy.")

    _head(doc, 3, "Endpoints and pages")
    doc.add_paragraph(
        "Endpoints (all advisory-only, none call a broker): GET /api/v2/rotation/summary runs the local "
        "rotation engine (cached) and returns the summary counts, data-quality, missing-sector and "
        "missing-analyst-upside counts, real rotation pairs, and per-symbol review candidates (worthless or "
        "delisted rows filtered, missing sectors backfilled from symbol_profiles). POST /api/v2/rotation/ask "
        "runs the grounded / local / oauth-prompt / dual paths via safe subprocess arguments with a hard "
        "timeout. POST /api/v2/rotation/grok-prompt builds the manual prompt. POST /api/v2/rotation/grok-review "
        "runs the inline free OAuth Grok second opinion. Pages: /v3/rotation (Rotation Intelligence) and "
        "/v3/advisor-changes (Advisor Changes), plus a Rotation tab on the Intelligence hub and a Rotation "
        "Advisor card on the Portfolio hub. All UI language is advisory - review, watch, second opinion, range "
        "unavailable - never execute, buy, sell, or place order.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation Intelligence section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation Intelligence section appended + saved")


if __name__ == "__main__":
    main()

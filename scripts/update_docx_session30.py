#!/usr/bin/env python3
"""Append Session 30 content to Reference Architecture DOCX.
Uses python-docx APPEND-ONLY operations per DOCX_UPDATE_PROTOCOL."""

from docx import Document
from docx.shared import Pt, RGBColor
import os
import shutil

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

def add_metric_table(doc, rows_data):
    tbl = doc.add_table(rows=1, cols=2)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for metric, value in rows_data:
        row = tbl.add_row().cells
        row[0].text = metric
        row[1].text = value

def add_feature_table(doc, headers, rows_data):
    tbl = doc.add_table(rows=1, cols=len(headers))
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

def main():
    # Backup first
    if os.path.exists(DOCX_PATH):
        shutil.copy2(DOCX_PATH, DOCX_PATH + ".bak_session30")

    doc = Document(DOCX_PATH)
    h1 = get_heading_style(doc, 1)
    h2 = get_heading_style(doc, 2)

    # ══════════════════════════════════════════════════════════════════
    # SESSION 30 HEADER
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h1
    p.text = "Session 30: Research Pipeline, Trading Autonomy, and Learning Loop (2026-05-12)"

    doc.add_paragraph(
        "Session 30 closes critical gaps in the research-to-action pipeline, paper trading "
        "autonomy, and system learning. Research findings are now indexed into RAG and injected "
        "into agent context. Paper trade outcomes auto-embed into the intelligence pipeline. "
        "The monitoring loop is fully self-checking with phantom position detection. "
        "A new Automated Trade Journal provides professional execution-log-level detail for "
        "every paper trade."
    )

    # ── Updated Operating Numbers ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Updated Operating Numbers"

    add_metric_table(doc, [
        ("Database Tables", "320+"),
        ("API Endpoints", "280+"),
        ("Command Center Pages", "43 (was 42)"),
        ("Cron Jobs", "152+"),
        ("RAG Source Types", "12 (was 10)"),
        ("Agent Context Injections", "13 (was 11)"),
        ("Morning Brief Sections", "7 (was 6)"),
        ("Portfolio Value", "$1,192,663"),
        ("Open Paper Trades", "2 (INFU, XMTR)"),
        ("Paper Trade Account", "ALPACA_PAPER (was TOS_PAPER)"),
    ])

    # ══════════════════════════════════════════════════════════════════
    # RESEARCH PIPELINE
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Research Advisory Pipeline — Now Fully Connected"

    doc.add_paragraph(
        "Research findings from user_research_topics were previously a write-only dead end — "
        "stored in PostgreSQL and sent to Telegram but invisible to agents, dashboards, and RAG. "
        "Session 30 wires the complete loop:"
    )

    add_feature_table(doc, ["Step", "Detail"], [
        ("Source", "user_research_topics table — created via Telegram 'research <topic>'"),
        ("Iteration", "iterate_research_topics.py runs daily, LLM with prior findings as context"),
        ("RAG Index", "Indexed as research_finding source type (1.25x boost) in content_embeddings"),
        ("Agent Injection", "Active Research Advisories block injected into all agent prompts"),
        ("Morning Brief", "RESEARCH ADVISORIES section (priority 6) — top 3 findings with preview"),
        ("Command Center", "Intelligence > Research Topics page (/v2/research-topics)"),
        ("Email Digest", "Included in daily GOG Gmail digest"),
    ])

    # ══════════════════════════════════════════════════════════════════
    # PAPER TRADING AUTONOMY
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Paper Trading — Fully Autonomous After Approval"

    doc.add_paragraph(
        "Paper trading is now fully autonomous beyond initial trade approval. "
        "The system detects price spikes, dynamically adjusts stops to lock in gains, "
        "monitors for negative news and volume fade, and auto-closes phantom positions. "
        "Alpaca is the source of truth — no DB-only positions can persist."
    )

    p = doc.add_paragraph()
    p.style = h2
    p.text = "Autonomous Monitoring Stack"

    add_feature_table(doc, ["Monitor", "Schedule", "Actions"], [
        ("paper_trade_monitor.py", "Every 5 min (market hours)",
         "R-multiple trailing stops, target hit close, near-target tighten, "
         "MFE/MAE tracking, stop placement, phantom detection"),
        ("open_trade_monitor.py", "Every 15 min (market hours)",
         "Near-stop alerts, near-target alerts, negative news scan, "
         "volume fade detection, stale trade warnings, extended profit alerts"),
        ("alpaca_paper_adapter.py", "On demand + cron sync",
         "Position sync, closed position detection, outcome curation trigger"),
    ])

    p = doc.add_paragraph()
    p.style = h2
    p.text = "R-Multiple Trailing Stop Logic"

    add_feature_table(doc, ["R-Multiple", "Stop Adjustment", "Effect"], [
        (">= 1.0R", "Move stop to breakeven", "Eliminate loss risk"),
        (">= 1.5R", "Lock 0.5R profit", "Guarantee partial gain"),
        (">= 2.0R", "Lock 1.0R profit", "Secure meaningful gain"),
        (">= 3.0R", "Lock 2.0R profit", "Capture extended move"),
        (">= 80% to target", "Tighten to 65% of target move", "Lock most of the gain"),
        ("Target hit", "Market sell", "Auto-close at target"),
    ])

    doc.add_paragraph(
        "Stops only move UP, never down. All adjustments are executed on Alpaca via API "
        "and logged as MONITOR_ADJUST_STOP curation events for full audit trail."
    )

    # ══════════════════════════════════════════════════════════════════
    # LEARNING LOOP
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Trade Outcome Learning Loop"

    doc.add_paragraph(
        "When a paper trade closes, the system now runs a complete learning pipeline:"
    )

    add_feature_table(doc, ["Hook", "Agent", "Output"], [
        ("iris_record_trade_outcome", "Iris", "Outcome lesson to agent_intelligence_rules"),
        ("aegis_write_post_trade_synthesis", "Aegis", "Synthesis paragraph for intelligence whiteboard"),
        ("trigger_outcome_lessons", "System", "Outcome scorer + feedback loop"),
        ("check_pattern_confirmation", "Iris", "Pattern library validation"),
        ("_index_trade_outcome_to_rag", "RAG", "Embed outcome into content_embeddings (1.35x boost, NEW)"),
    ])

    doc.add_paragraph(
        "Trade outcomes are the highest-boosted RAG source (1.35x), ensuring agents "
        "prioritize learning from real trades over news or social signals."
    )

    # ══════════════════════════════════════════════════════════════════
    # PHANTOM DETECTION
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Phantom Position Detection"

    doc.add_paragraph(
        "Session 30 discovered and fixed phantom positions — trades marked as 'open' in the "
        "database but with no corresponding position on Alpaca. Root causes: (1) Alpaca orders "
        "that were canceled but not checked, (2) Account mismatch (TOS_PAPER vs ALPACA_PAPER). "
        "The paper_trade_monitor now cross-checks every cycle: if DB says open but Alpaca has "
        "no position, auto-close with 'phantom_no_alpaca_position' reason and full audit trail."
    )

    # ══════════════════════════════════════════════════════════════════
    # AUTOMATED TRADE JOURNAL
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Automated Trade Journal"

    doc.add_paragraph(
        "New 'Automated Journal' tab under Trade Journal (5th tab). Shows professional "
        "execution-log-level detail for every paper trade, filtered by account:"
    )

    add_feature_table(doc, ["Section", "Content"], [
        ("Position Details", "Shares, long/short, entry, stop, target, current, risk $, MFE, MAE, VIX, regime"),
        ("Entry Rationale", "Strategy, opened_via, catalyst (verified/unverified), risk gate result"),
        ("Execution Log", "Timeline of all MONITOR_* events, alerts, stop adjustments, system observations"),
        ("Exit & Outcome", "Exit reason, verdict (WIN/LOSS), closed_via"),
        ("Journal Review", "Mistake tags, strength tags, lesson learned, system fixes applied"),
    ])

    doc.add_paragraph(
        "API endpoint: /api/v2/automated-trade-journal?account=ALPACA_PAPER. "
        "Returns enriched trade data with execution_log, alerts, and journal_reviews per trade."
    )

    # ══════════════════════════════════════════════════════════════════
    # EMAIL DIGEST FIX
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Email Digest Fix"

    doc.add_paragraph(
        "The daily email digest via GOG Gmail stopped firing on 2026-05-06. Root cause: "
        "portfolio_orchestrator.py called 'gog' without full path, which fails in cron's "
        "minimal PATH. Fixed to use ~/.local/bin/gog. Research updates now emailed as digest."
    )

    # ══════════════════════════════════════════════════════════════════
    # PHASE 1D
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "LLM Fleet Phase 1D — Gemma Pilot Limit 5"

    doc.add_paragraph(
        "Controlled expansion of gemma3-overnight pilot from max 3 to max 5 symbols. "
        "Manual test completed: 5 symbols classified in 7 min 41 sec (under 10 min timeout). "
        "Per-symbol average: ~91 sec. gemma3-overnight VRAM: 13.64 GB GPU + 4.79 GB CPU. "
        "qwen3:14b + nomic-embed-text fully restored after each run. "
        "Recommendation: Phase 1 stops at manual limit 5. Nightly cron requires "
        "queue drain check and model-swap lock before scheduling."
    )

    # ══════════════════════════════════════════════════════════════════
    # COMMITS
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Session 30 Commits"

    add_feature_table(doc, ["Commit", "Description"], [
        ("6231e21", "Fix email digest GOG path bug"),
        ("6300e42", "Add Research Topics page to command center"),
        ("986e4e2", "Wire research into RAG, Morning Brief, agent context"),
        ("867b7a0", "Phase 1D: raise gemma pilot max limit to 5"),
        ("336e645", "Phase 1D: limit-5 pilot report and deployment log"),
        ("7b05010", "Fix paper trade monitor: INFU orphan, sync mismatch"),
        ("30cef84", "Replace TOS_PAPER with ALPACA_PAPER everywhere"),
        ("70b3ddf", "Close trading loop: phantom detection, RAG outcomes, journal"),
        ("32e8331", "Pending changes: telegram smart split, youtube channels, docs"),
        ("769b519", "Add Automated Trade Journal with execution log"),
    ])

    doc.save(DOCX_PATH)
    print(f"Session 30 appended to {DOCX_PATH}")
    print(f"Backup at {DOCX_PATH}.bak_session30")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")
    main()

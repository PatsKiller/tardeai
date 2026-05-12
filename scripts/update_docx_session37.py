#!/usr/bin/env python3
"""Append Session 37 content to Reference Architecture DOCX.
Uses python-docx APPEND-ONLY operations per DOCX_UPDATE_PROTOCOL."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy
import os

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    """Get heading style from existing paragraphs (safe method)."""
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    """Add borders to a table via XML."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

def main():
    doc = Document(DOCX_PATH)
    h1 = get_heading_style(doc, 1)
    h2 = get_heading_style(doc, 2)

    # ── Section Header ──
    p = doc.add_paragraph()
    p.style = h1
    p.text = "Session 37: Closed-Loop Intelligence Pipeline (2026-05-10)"

    doc.add_paragraph(
        "Session 37 transforms the system from a data collector into a closed-loop intelligence engine. "
        "Every data source feeds into correlation, every agent analysis feeds back into new searches, "
        "and every failure triggers a Telegram notification with reply-to-retry."
    )

    # ── Operating Numbers ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Updated Operating Numbers"

    tbl = doc.add_table(rows=1, cols=2)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    stats = [
        ("News articles (total)", "2,880"),
        ("News approved for RAG", "2,328"),
        ("News sentiment scored", "2,880 (100%)"),
        ("YouTube transcripts", "771"),
        ("YouTube approved for RAG", "748"),
        ("Social posts (with sentiment)", "2,245"),
        ("Sentiment observations", "2,890"),
        ("Fused signals (strategy-weighted)", "1,264"),
        ("Content entity links", "21 (curator extracting)"),
        ("Intelligence entities (active)", "612"),
        ("Active topics (LLM-curated)", "17"),
        ("Incubator candidates (ACTIVE)", "632"),
        ("Paper trade proposals (PENDING)", "13"),
        ("Agent results (last 24h)", "115"),
    ]
    for metric, value in stats:
        row = tbl.add_row().cells
        row[0].text = metric
        row[1].text = value

    # ── Closed-Loop Intelligence Flow ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Closed-Loop Intelligence Flow"

    doc.add_paragraph(
        "The system operates as a feedback loop, not a data warehouse. "
        "Data flows: INGEST -> CORRELATE -> SENTIMENT -> CURATE -> AGENTS -> DEMAND SIGNAL -> back to INGEST."
    )

    tbl2 = doc.add_table(rows=1, cols=4)
    add_table_borders(tbl2)
    hdr2 = tbl2.rows[0].cells
    for cell, text in zip(hdr2, ["Layer", "Scripts", "Output", "Cadence"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    layers = [
        ("Ingest", "news_ingestion, social_ingest, youtube_transcript_ingest, sec_data_ingest",
         "Raw rows in news_articles, social_posts, youtube_transcripts", "2-3x daily + on-demand"),
        ("Correlate", "intelligence_entity_manager, topic_curator (extract_and_link_entities)",
         "content_entity_links, intelligence_entities (per-symbol score)", "After each ingest"),
        ("Sentiment", "sentiment_processor, signal_fusion",
         "sentiment_observations, fused_signals (per-symbol composite)", "2x daily (7 AM, 12 PM)"),
        ("Curate", "topic_curator (rate, extract, improve_queries)",
         "rag_status, llm_generated_queries, content_entity_links", "Daily 7 AM"),
        ("Agent Analysis", "process_watchlist_agent_jobs",
         "watchlist_agent_results (recommendation + narrative)", "Every 15 min"),
        ("Demand Signal", "agent_event_router (handle_content_gap, handle_research_more_demand)",
         "Auto-triggered: topic_ingestion -> sentiment -> RAG -> re-analysis", "On event"),
        ("Feedback", "agent_outcome_scorer, learning_governance",
         "agent_calibration (win rate, PnL), confidence adjustments", "Daily 5:30 AM"),
        ("RAG Index", "rag_indexer",
         "Vector embeddings for semantic search", "4x daily + on gap-fill"),
    ]
    for layer, scripts, output, cadence in layers:
        row = tbl2.add_row().cells
        row[0].text = layer
        row[1].text = scripts
        row[2].text = output
        row[3].text = cadence

    # ── Agent Context Stack ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Agent Context Injection (12-Layer Stack)"

    doc.add_paragraph(
        "Every agent analysis receives this full context per symbol:"
    )

    context_items = [
        "1. Scan Intelligence -- screener position, score, decision (GO/WAIT/AVOID)",
        "2. RAG Pre-Context -- top 5 prior intelligence items (news, transcripts, agent results)",
        "3. News Sentiment (7d) -- article count, avg score, headlines with sentiment labels",
        "4. Social Sentiment (7d) -- post count, bullish/bearish/neutral breakdown, top posts",
        "5. Fused Signal -- strategy-weighted composite (catalyst + news + social + sentiment)",
        "6. Peer Agent Notes -- what other agents concluded on this symbol recently",
        "7. Content Gap Warnings -- Iris librarian flags on missing coverage",
        "8. Technical Confluence -- RSI, SMA, ATR, confluence tier",
        "9. Prospects Context -- pipeline position (incubator, proposal, paper trade)",
        "10. Calibration Data -- agent's own win rate, avg confidence, past PnL",
        "11. Strategy Playbook -- role instructions, entry/exit rules, risk parameters",
        "12. Global Rules G1-G10 -- income protection, SSDI awareness, confidence gating",
    ]
    for item in context_items:
        doc.add_paragraph(item, style='List Paragraph')

    # ── Per-Agent Integration ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Per-Agent Full-Circle Integration"

    tbl3 = doc.add_table(rows=1, cols=5)
    add_table_borders(tbl3)
    hdr3 = tbl3.rows[0].cells
    for cell, text in zip(hdr3, ["Agent", "Reads", "Writes", "Triggers", "LLM"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    agents = [
        ("Maria", "RAG, sentiment, social, fused, peers, playbook, scans",
         "BUY/HOLD/AVOID + narrative", "Re-analysis on gap-fill; debate on SEC insider buy", "qwen3:14b (2-pass)"),
        ("Steph", "Portfolio state, allocations, income, sentiment",
         "ADD/TRIM/HOLD + allocation review", "Escalation for concentration risk; INCOME_CRITICAL", "qwen3:14b"),
        ("Alex", "Roth models, IRMAA, tax brackets, retirement RAG",
         "Research reports, Roth ladder plans", "Auto-queued on SEC insider consensus", "qwen3:14b + Claude"),
        ("Aegis", "All agent results, portfolio, overnight events",
         "Morning briefs, synthesis reports", "Morning brief delivery; post-trade synthesis", "qwen3:14b"),
        ("Iris", "Content freshness, RAG coverage, duplicates, staleness",
         "CONTENT_GAP events, hygiene proposals", "CONTENT_GAP -> auto-search; hygiene escalations", "qwen3:14b"),
        ("Scalp Critic", "Incubator candidates, catalyst, technicals, news/social",
         "Grade A-F, verdict PROMOTE/HOLD/DROP", "Gates incubator -> proposal promotion", "qwen3:14b"),
    ]
    for agent, reads, writes, triggers, llm in agents:
        row = tbl3.add_row().cells
        row[0].text = agent
        row[1].text = reads
        row[2].text = writes
        row[3].text = triggers
        row[4].text = llm

    # ── Demand-Driven Search ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Demand-Driven Search Loop"

    doc.add_paragraph(
        "When agents need more data, the system auto-responds with a full search-ingest-analyze cycle:"
    )

    tbl4 = doc.add_table(rows=1, cols=3)
    add_table_borders(tbl4)
    hdr4 = tbl4.rows[0].cells
    for cell, text in zip(hdr4, ["Trigger", "Source", "Action Chain"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    triggers = [
        ("CONTENT_GAP", "Iris librarian detects missing coverage",
         "topic_ingestion -> news search -> sentiment_processor -> RAG re-index -> Maria re-queued -> Telegram notify"),
        ("RESEARCH_MORE", "Agent outputs low-confidence RESEARCH_MORE",
         "Check watchdog_actions for recent fills -> fires synthetic CONTENT_GAP -> full search loop"),
        ("Improved Queries", "topic_curator generates better search terms",
         "Auto-runs topic_ingestion --use-llm-queries -> new content flows back to curation"),
    ]
    for trigger, source, chain in triggers:
        row = tbl4.add_row().cells
        row[0].text = trigger
        row[1].text = source
        row[2].text = chain

    # ── LLM Curation Schedule ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "LLM Curation Schedule"

    doc.add_paragraph(
        "The system gets smarter over time through scheduled LLM curation that improves search queries, "
        "rates content quality, and extracts entity relationships:"
    )

    tbl5 = doc.add_table(rows=1, cols=3)
    add_table_borders(tbl5)
    hdr5 = tbl5.rows[0].cells
    for cell, text in zip(hdr5, ["When", "What Happens", "LLM"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    schedule = [
        ("7 AM daily", "topic_curator rates pending content (approved/low_quality/blocked)", "qwen3:14b"),
        ("7 AM daily", "topic_curator extracts tickers + topics -> content_entity_links", "qwen3:14b"),
        ("7 AM daily", "topic_curator improves queries -> generates targeted queries per topic", "qwen3:14b"),
        ("7 AM daily", "Auto-ingests with improved queries (step 3b)", "N/A (search APIs)"),
        ("7 AM + 12 PM", "sentiment_processor scores all unscored news articles", "Lexicon v1"),
        ("7:15 AM + 12:15 PM", "signal_fusion fuses catalyst + news + social + sentiment per symbol", "N/A"),
        ("8:10 AM + 6 PM", "Incubator LLM screener grades candidates A-F", "qwen3:14b"),
        ("On CONTENT_GAP", "Auto-trigger: topic search -> news -> sentiment -> RAG -> re-analyze", "qwen3:14b"),
        ("On RESEARCH_MORE", "Agent demand -> synthetic CONTENT_GAP -> full search loop", "qwen3:14b"),
        ("5:30 AM daily", "Outcome scorer grades past recommendations (calibration update)", "Rule-based"),
        ("Sunday 6 AM", "Iris hygiene: demote stale content, detect superseded data", "Rule-based"),
    ]
    for when, what, llm in schedule:
        row = tbl5.add_row().cells
        row[0].text = when
        row[1].text = what
        row[2].text = llm

    # ── Failure Alerting ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Pipeline Failure Alerting (Session 37)"

    doc.add_paragraph(
        "Every critical cron job is wrapped with pipeline_alert.py. On non-zero exit, sends Telegram "
        "with error excerpt and reply-to-retry command. Wrapped pipelines: news_ingestion, youtube_ingest, "
        "overnight_batch, sec_data_ingest, event_detector, previously_traded, pipeline_watchdog."
    )

    doc.add_paragraph(
        "New Telegram commands for retry: 'run promoter' (retry incubator promoter), "
        "'run promoter dry' (dry-run first), 'status' (system health check)."
    )

    # ── 2Captcha Integration ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "2Captcha Integration"

    doc.add_paragraph(
        "API key in .env (TWOCAPTCHA_API_KEY). Enables automated data collection from CAPTCHA-protected sites. "
        "Supported: reCAPTCHA v2/v3, hCaptcha, Cloudflare Turnstile, image CAPTCHA, FunCaptcha, GeeTest."
    )

    tbl6 = doc.add_table(rows=1, cols=4)
    add_table_borders(tbl6)
    hdr6 = tbl6.rows[0].cells
    for cell, text in zip(hdr6, ["Site", "Data Value", "CAPTCHA Type", "Priority"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sites = [
        ("Seeking Alpha", "Premium analyst reports, earnings transcripts", "reCAPTCHA v2", "High"),
        ("TipRanks", "Analyst consensus, price targets, smart score", "reCAPTCHA v2", "High"),
        ("Finviz (rate-limited)", "Screener when cookie expires", "hCaptcha", "Medium"),
        ("MarketWatch", "Premium articles, options flow", "Cloudflare", "Medium"),
        ("Barron's", "Premium analysis, portfolio strategy", "Cloudflare", "Low"),
    ]
    for site, data, captcha, priority in sites:
        row = tbl6.add_row().cells
        row[0].text = site
        row[1].text = data
        row[2].text = captcha
        row[3].text = priority

    # ── Bugs Fixed ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Bugs Fixed (Session 37)"

    fixes = [
        "incubator_proposal_promoter.py: RealDictCursor fetchone()[0] KeyError -> ['id'] (blocked promotions 2+ days)",
        "youtube_transcript_ingest.py: Same dict cursor fix (blocked YouTube ingestion 2+ days)",
        "incubator_proposal_promoter.py: _queue_llm_review transaction poisoning -> savepoint isolation",
        "iris_taxonomy_agent.py: Hardcoded strategy_tag 'investment_general' -> resolved from source tables",
        "iris_taxonomy_agent.py: LIMIT 500 on whiteboard hygiene -> removed cap (882 items now processed)",
        "process_watchlist_agent_jobs.py: Missing sentiment/social/fused context in agent prompts -> added",
        "sentiment_processor.py and signal_fusion.py: Not scheduled -> added to crontab (7 AM + 12 PM)",
        "incubator_proposal_promoter and incubator_llm_screener: Not in pipeline_schedule -> registered for watchdog",
    ]
    for fix in fixes:
        doc.add_paragraph(fix, style='List Paragraph')

    # ── New Scripts ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "New Scripts (Session 37)"

    scripts = [
        "scripts/pipeline_alert.py -- Universal cron failure wrapper with Telegram notification",
        "scripts/cron_wrapper.sh -- Bash equivalent for shell-based cron entries",
    ]
    for s in scripts:
        doc.add_paragraph(s, style='List Paragraph')

    # ── Modified Scripts ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Modified Scripts (Session 37)"

    modified = [
        "agent_event_router.py -- Added handle_content_gap(), handle_research_more_demand() (+173 lines)",
        "incubator_proposal_promoter.py -- Dict cursor fix, savepoint, Telegram alerts (+89 lines)",
        "iris_taxonomy_agent.py -- Strategy tag resolution from source tables, LIMIT removed (+57 lines)",
        "process_watchlist_agent_jobs.py -- _get_sentiment_social_context() injected into prompts (+75 lines)",
        "telegram_command_handler.py -- Added 'run promoter' command (+248 lines)",
        "topic_curator.py -- Step 3b: auto-ingest with improved queries (+20 lines)",
        "topic_ingestion.py -- --use-llm-queries flag for curator-improved queries (+16 lines)",
        "youtube_transcript_ingest.py -- Dict cursor fix, Telegram failure alerting (+23 lines)",
    ]
    for m in modified:
        doc.add_paragraph(m, style='List Paragraph')

    # ── Version History ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Version History"

    doc.add_paragraph(
        "Session 37 (2026-05-10): Closed-loop intelligence pipeline. "
        "3 commits: c8fe976 (core fixes + loop), ea42b94 (docs), 61034b6 (workflow + 2captcha docs). "
        "10 files changed, 846 insertions. "
        "Total pipeline: 2,880 news, 771 YouTube, 2,245 social, 2,890 sentiment observations, "
        "1,264 fused signals, 612 intelligence entities, 17 LLM-curated topics."
    )

    # Save
    doc.save(DOCX_PATH)
    print(f"DOCX updated: {DOCX_PATH}")
    print(f"File size: {os.path.getsize(DOCX_PATH):,} bytes")


if __name__ == "__main__":
    os.chdir("/home/johnclaw/trade-ai-v12-rebuild/trade-ai-v12-rebuild")
    main()

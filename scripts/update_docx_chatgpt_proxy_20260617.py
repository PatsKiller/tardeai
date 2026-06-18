#!/usr/bin/env python3
"""Append the ChatGPT OAuth proxy + lane-monitor addendum to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). As-built 2026-06-17; detail in docs/project/ROTATION_LLM_ADVISOR.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "ChatGPT OAuth proxy + free-OAuth lane monitor (2026-06-17)"


def h(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "ChatGPT is now an inline free-OAuth lane, mirroring the Grok proxy. scripts/chatgpt_oauth_proxy.py "
        "runs a local OpenAI-compatible server on port 8646 that drives the operator's already-authenticated "
        "hermes openai-codex CLI inside a real pseudo-TTY. Hermes owns the OAuth; the proxy never reads or "
        "refreshes raw tokens. It is free under the ChatGPT subscription, not the metered OpenAI API, and "
        "exposes health, models, and chat-completions endpoints with a token-expired flag and a clean "
        "re-login hint when the session has ended. It runs as a user systemd service that restarts "
        "automatically, the same way the Grok xAI OAuth proxy runs on port 8645.")
    doc.add_paragraph(
        "The proxy is shared across all Hermes tasks: the external-researcher codex call now prefers the "
        "proxy and falls back to the pseudo-TTY CLI, so every Hermes task that uses the ChatGPT or codex lane "
        "- external research, curation, and rotation oversight - works headless through it. Grok was already "
        "proxy-backed.")
    doc.add_paragraph(
        "A Command-Center monitor endpoint, GET /api/v2/llm/oauth-lanes, probes all of the free OAuth lanes - "
        "Grok on 8645, ChatGPT on 8646, the Hermes Nous portal, and the local gemma lane via ollama - and "
        "returns each lane's reachable, authenticated, token-expired, status, and re-login hint. The Rotation "
        "Intelligence Independent Oversight panel shows a live lane-health strip: green for ready, amber for "
        "needs-login or session-expired, and red for offline, with a refresh control. To activate the ChatGPT "
        "lane the operator logs in once with hermes auth add openai-codex; the Hermes Nous lane is activated "
        "with hermes portal login. Everything stays advisory and free; no API key and no metered API.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("ChatGPT proxy addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("ChatGPT proxy addendum appended + saved")


if __name__ == "__main__":
    main()

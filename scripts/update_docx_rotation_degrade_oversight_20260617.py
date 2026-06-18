#!/usr/bin/env python3
"""Append the Rotation Intelligence degradation + oversight addendum to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). As-built 2026-06-17; detail in docs/project/ROTATION_LLM_ADVISOR.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation Intelligence - holdings degradation + independent oversight (2026-06-17)"


def h(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "The rotate-OUT side of Rotation Intelligence is now driven by real deterioration, and the engine's "
        "recommendations get independent oversight. Everything stays advisory: read-only data, no broker "
        "call, no order, no dollar amount auto-executed.")
    doc.add_paragraph(
        "Holdings degradation. The summary joins each held name to the latest Aegis nightly brief and returns "
        "a deteriorating-holdings list plus a degradation tag on each review candidate and rebalance idea, "
        "reordering the trim pool to put deteriorating names first. Accuracy is labeled per signal: a status "
        "of triggered, danger, or warning comes from deterministic stop-distance math in aegis_surveillance "
        "(high confidence, not an LLM); a status of weakening or broken comes from a local-LLM thesis read in "
        "aegis_synthesis using ollama gemma3 (softer, about 0.55 confidence); and near-52-week-low and "
        "analyst-recommendation flags come from the nightly finviz and yahoo snapshot. The UI badge shows "
        "whether each flag is stop math or an LLM read, so the operator knows hard versus soft. The weekly "
        "digest leads with the deteriorating holdings, split into deterministic-stop and soft counts.")
    doc.add_paragraph(
        "Independent oversight layers. A new endpoint, POST /api/v2/rotation/oversight, runs two independent "
        "oversight models over the rotate-out flags, rebalance ideas, and sector overweights - a second and "
        "third opinion on the engine itself. Both are free OAuth with no API key and no paid API: Grok through "
        "the local xAI OAuth proxy, and ChatGPT through the openai-codex OAuth that is free under the "
        "operator's ChatGPT subscription, not the metered OpenAI API. The endpoint hard-restricts the lanes to "
        "Grok and ChatGPT so the paid Claude and OpenAI key paths are never used. Because the Codex CLI needs "
        "a real terminal, in the headless server it may report itself unavailable; the endpoint then returns "
        "the prompt for a manual free-web paste. Each model is asked, as an oversight layer, whether the "
        "rotate-out flags are reasonable or an over-reaction to a stop wiggle on a tiny position, whether any "
        "rebalance idea is a poor fit such as trimming an income or core holding into a speculative microcap, "
        "what is missing, and an overall AGREE, CAUTION, or DISAGREE verdict - never a dollar amount, never an "
        "order. The page shows both verdicts side by side in an Independent Oversight panel.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation degrade/oversight addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation degrade/oversight addendum appended + saved")


if __name__ == "__main__":
    main()

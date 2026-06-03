#!/usr/bin/env python3
"""Append V3 Backtest Intelligence addendum to Reference Architecture DOCX.

APPEND-ONLY per docx protocol — uses python-docx, never raw XML edit.
Run once; idempotency guarded by a marker-heading check.
"""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "V3 Backtest Intelligence (Session 2026-06-02 Addendum)"


def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    tblPr.append(etree.fromstring(borders))


def main():
    doc = Document(DOCX_PATH)

    # Idempotency: skip if marker already present
    for p in doc.paragraphs:
        if MARKER in (p.text or ""):
            print("Addendum already present — skipping (idempotent).")
            return

    h2 = get_heading_style(doc, 2)
    h3 = get_heading_style(doc, 3)

    p = doc.add_paragraph()
    if h2:
        p.style = h2
    p.text = MARKER

    doc.add_paragraph(
        "The v2 Backtesting page was fully ported into the v3 Strategy -> Backtest tab "
        "(read-only), and a Backtest Intelligence layer was added: entry/exit grading, "
        "edge decay, capture/left-on-table, and potential-over-time. Default run_type is "
        "replay_trades (real data); champion rows are seeded simulations. Discovery: the "
        "closed-trade entry/exit grading engine (trade_backtest_engine.py) existed but was "
        "never scheduled, so trade_backtest_results was empty. It is now populated and on cron."
    )

    # Schema additions
    sp = doc.add_paragraph()
    if h3:
        sp.style = h3
    sp.text = "Schema additions"
    t1 = doc.add_table(rows=1, cols=3)
    add_table_borders(t1)
    for cell, text in zip(t1.rows[0].cells, ["Table", "Type", "Purpose"]):
        cell.text = text
        for r in cell.paragraphs[0].runs:
            r.bold = True
    schema = [
        ("backtest_result_history", "APPEND-ONLY (UNIQUE run_id, ON CONFLICT DO NOTHING)",
         "One permanent snapshot row per backtest run for potential-over-time; never overwrites. Backfilled 209 rows."),
        ("trade_backtest_results", "Populated (was empty)",
         "Closed-trade entry/exit A-D grades, RSI/SMA/ATR context, left-on-table 5/10/20d. Filled by trade_backtest_engine.py."),
    ]
    for name, typ, purpose in schema:
        row = t1.add_row().cells
        row[0].text = name
        row[1].text = typ
        row[2].text = purpose

    # Scheduling / endpoint
    cp = doc.add_paragraph()
    if h3:
        cp.style = h3
    cp.text = "Scheduling and endpoints"
    t2 = doc.add_table(rows=1, cols=3)
    add_table_borders(t2)
    for cell, text in zip(t2.rows[0].cells, ["Item", "Schedule / Route", "Notes"]):
        cell.text = text
        for r in cell.paragraphs[0].runs:
            r.bold = True
    rows = [
        ("trade_backtest_engine.py", "Weekdays 6:30 PM ET (cron + safe_flock)", "Grades newly closed trades."),
        ("backtest_history_snapshot.py", "Weekdays 6:10 AM + Sun 10:10 PM ET (cron)", "Append-only archiver; idempotent."),
        ("/api/v2/backtesting/result-history", "GET ?run_type=", "Read-only; defaults to replay_trades."),
        ("cron vs systemd", "services=systemd, batch=cron", "tradeai-portfolio-server.service is systemd; scheduled jobs use cron + safe_flock.sh."),
    ]
    for a, b, c in rows:
        row = t2.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    doc.add_paragraph(
        "Findings from first grading run (61 closed trades): entry grades B:1 / C:10 / D:48, "
        "avg entry RSI 67.8 (late entries), $210,141 left on table (20d), $189,001 of it from "
        "21 grade-D exits. Frontend: apps/command-center-v3/src/components/BacktestPanel.tsx (new), "
        "DetailDrawer.tsx (enriched with sparkline + grade badges + Journal<->Backtest link)."
    )

    doc.save(DOCX_PATH)
    print("Addendum appended to", DOCX_PATH)


if __name__ == "__main__":
    main()

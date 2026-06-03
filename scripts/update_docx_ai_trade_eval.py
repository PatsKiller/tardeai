#!/usr/bin/env python3
"""Append AI Trade Eval addendum to Reference Architecture DOCX (append-only, idempotent)."""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "AI Trade Eval — Structured LLM Trade Evaluation (Session 2026-06-02)"


def get_heading_style(doc, level):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {level}":
            return p.style
    return None


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    tblPr.append(etree.fromstring(borders))


def main():
    doc = Document(DOCX_PATH)
    for p in doc.paragraphs:
        if MARKER in (p.text or ""):
            print("Addendum already present — skipping (idempotent).")
            return
    h2 = get_heading_style(doc, 2)
    p = doc.add_paragraph()
    if h2:
        p.style = h2
    p.text = MARKER

    doc.add_paragraph(
        "Post-trade research / journaling and model evaluation — NOT live trading advice. "
        "trade_backtest_engine.py was extended to compute MACD, Bollinger, ADX, Fibonacci, daily "
        "candlestick and market-structure tags from the daily OHLCV it already fetches (13 new "
        "trade_backtest_results columns; VWAP/intraday remain out of scope). "
        "trade_close_llm_analyzer.py gained a self-contained --structured path that grades each "
        "enriched closed trade with local gemma3:12b, producing six 0-100 scores (confluence, "
        "entry_timing, exit_quality, risk_reward, management, overall) plus one of 12 verdict labels, "
        "entry/exit assessment, improvements and data_gaps. Outcome and quality are scored separately. "
        "Stored on trade_llm_reviews (review_stage='structured_backtest_eval', headline columns "
        "eval_overall_score + eval_verdict). Dedup is by setup-hash so identical setups across accounts "
        "collapse to one evaluation. Surfaced in the v3 Backtest 'AI Trade Eval' tab via "
        "/api/v2/backtesting/trade-evaluations. Cron: weekdays 9 PM ET, --limit 12, flock-guarded."
    )

    tbl = doc.add_table(rows=1, cols=2)
    add_table_borders(tbl)
    for cell, text in zip(tbl.rows[0].cells, ["Component", "Detail"]):
        cell.text = text
        for r in cell.paragraphs[0].runs:
            r.bold = True
    rows = [
        ("Indicators added", "MACD(line/signal/hist/state), Bollinger(%/state), ADX, Fibonacci(level+leg), daily candlestick, structure"),
        ("Model", "gemma3:12b local; ~1-4 min/trade on CPU; non-mutating, advisory preflight"),
        ("Output", "scores{6} + verdict{12 labels} + entry/exit assessment + improvements + data_gaps"),
        ("Storage", "trade_llm_reviews; eval_overall_score, eval_verdict, output_payload(jsonb)"),
        ("Endpoint", "/api/v2/backtesting/trade-evaluations (read-only; evaluations + verdict distribution + disclaimer)"),
        ("Dedup", "by input_snapshot_hash (setup-level); model_error rows retry"),
        ("Schedule", "weekday 9 PM ET cron, --limit 12, safe_flock guarded"),
    ]
    for a, b in rows:
        c = tbl.add_row().cells
        c[0].text = a
        c[1].text = b

    doc.save(DOCX_PATH)
    print("AI Trade Eval addendum appended.")


if __name__ == "__main__":
    main()

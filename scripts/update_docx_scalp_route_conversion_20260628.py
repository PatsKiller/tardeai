#!/usr/bin/env python3
"""Append the scalp route-persistence / large-float-scout / conversion-path section to the
Reference Architecture DOCX. APPEND-ONLY, idempotent (marker-guarded), python-docx only —
never raw XML. A timestamped .bak is written first. As-built 2026-06-28 (PR #10, merge bbf12d5a)."""
import shutil
from datetime import datetime, timezone

from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Scalp route persistence, large-float scout & paper conversion (2026-06-28)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)

    doc.add_paragraph(
        "A further pass completed the Social to Momentum Scalp lifecycle so the system generates both "
        "regular micro-float momentum scalps and social-derived scalps optimally, safely, and "
        "traceably, and so that names too large to be standard scalps are kept and labelled rather "
        "than discarded. The work landed through pull request ten with the release-readiness CI "
        "green. It preserved the earlier honest correction — there are still only two confirmed closed "
        "momentum scalp paper trades against a thirty-trade gate — so the combined lifecycle remains "
        "at 4.4 out of five: the engineering is mature, the empirical sample is not. No order behaviour "
        "changed, no live broker path was touched, quote-freshness and the time-to-live, liquidity, and "
        "trading-window gates were not weakened, and the operator confirmation and two-factor path "
        "remained immutable and out of scope.")

    _head(doc, 3, "Durable routing, the large-float scout, and route-aware injection")
    doc.add_paragraph(
        "Every social candidate is now routed by a single deterministic policy whose decision is "
        "persisted as first-class columns on the scan records, so a candidate's route, actionability, "
        "owning strategy, reason codes, and catalyst-verification state travel with it for audit and "
        "downstream optimisation. The standard momentum scalp is now strictly a micro-float setup — "
        "twenty million shares or fewer with a verified catalyst — and the legacy fallback that could "
        "infer a momentum scalp for floats up to a hundred million, with no catalyst requirement, was "
        "removed. A verified name whose float is above that micro-cap ceiling is no longer thrown away: "
        "it is retained as a large-float social scout, clearly labelled LARGE FLOAT and routed to manual "
        "review, so the operator sees it for what it is and it can never masquerade as, or be fast-pathed "
        "as, a standard low-float scalp. The live runner now injects social candidates by route rather "
        "than by score alone: only a verified micro-cap GO enters the tradeable path; large-float scouts "
        "enter as labelled manual-review candidates; and social-only names remain advisory, never "
        "actionable. Signal creation enforces the same durable route, so a watch-only or scout candidate "
        "can never become a standard momentum scalp signal further down the pipeline.")

    _head(doc, 3, "The paper conversion path and the freshness timing gap")
    doc.add_paragraph(
        "Diagnosis confirmed the conversion bottleneck is operational timing, not a defective gate: "
        "proposals reach the approver after their quote has gone stale, and the approval correctly "
        "refuses to act on a stale quote. A freshness service-level report makes this concrete — the "
        "failures are stale-quote refusals rather than time-to-live expiries, the median time from "
        "proposal creation to first approval attempt is about ten minutes with a long tail, and only a "
        "handful of proposals would have remained fresh had the approver run within a few minutes of "
        "creation. The response was a safe, paper-only, dry-run-first fast runner that finds a fresh, "
        "in-window, micro-cap momentum scalp proposal and routes it through the existing approval path "
        "before the thirty-minute time-to-live, reusing the existing risk and approval logic rather than "
        "reimplementing it, and rejecting anything stale, expired, out-of-window, social-only, or "
        "large-float. It never reaches the live broker path. The remaining gap to a higher maturity is "
        "therefore empirical and operational: accumulate enough confirmed closed paper trades by running "
        "that fresh, in-window approval cadence — not by loosening any safety gate.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    bak = f"{DOCX_PATH}.bak_scalp_route_conversion_{stamp}"
    shutil.copy2(DOCX_PATH, bak)
    print(f"backup written: {bak}")
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

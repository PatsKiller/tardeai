#!/usr/bin/env python3
"""Append the validation-taxonomy normalization section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded), python-docx only, timestamped .bak first. As-built
2026-06-28 (PR #13, merge 48a349c1)."""
import shutil
from datetime import datetime, timezone
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Momentum Scalp validation taxonomy normalization (2026-06-28)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "An operator-facing terminology pass renamed the momentum-scalp sample-collection lifecycle "
        "from \"paper\" to \"validation\". The word paper was misleading: this path has nothing to do "
        "with operator approval or live trading — it is sandbox/simulated strategy-validation execution "
        "used to gather an empirical sample before a strategy can be promoted. The canonical lifecycle "
        "terms are now validation execution, validation fast path, simulated validation trade, "
        "validation sample, sandbox validation account, and validation tracker. No database was renamed: "
        "the legacy storage and adapters (the paper_trades table, the paper-named submitter and logger "
        "modules, and the alpaca_paper sandbox account identifier) remain in place and are documented as "
        "backward-compatibility aliases, with clean validation-named wrapper modules layered over them. "
        "The change was committed through pull request thirteen with the release-readiness CI green, and "
        "a taxonomy audit now fails the build if forbidden operator-facing paper phrasing reappears "
        "outside those allowed legacy contexts.")
    _head(doc, 3, "What did and did not change")
    doc.add_paragraph(
        "The canonical validation fast path is a thin wrapper over the existing deterministic gate logic "
        "— it reuses exactly the same gates (verified micro-float route, intraday window, time-to-live, "
        "quote freshness, price drift, liquidity, trade-plan and reward-to-risk checks) so nothing was "
        "weakened, and a compatibility test proves the old and new entry points produce identical "
        "decisions. Reports and the config now speak in validation terms (validation_fast_path metrics, "
        "confirmed closed validation trades, a validation gate) while keeping legacy field names as "
        "documented aliases. A read-only validation operations report and a supplementary deterministic "
        "quality gate were added. Crucially, none of this touches live execution: validation is "
        "sandbox/simulated only, validation sample collection needs no human approval because "
        "deterministic gates replace it, but promotion off TESTING still requires human review and any "
        "live trading still requires the unchanged operator-confirmation and two-factor path. Large-float "
        "social scouts remain manual-review only and social-only candidates remain watch/wait only. "
        "Maturity is unchanged at 4.4 of five — the remaining gap is empirical sample collection, now "
        "framed honestly as validation, not approval.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    bak = f"{DOCX_PATH}.bak_validation_taxonomy_{stamp}"
    shutil.copy2(DOCX_PATH, bak)
    print(f"backup written: {bak}")
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

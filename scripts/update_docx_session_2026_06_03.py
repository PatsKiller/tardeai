#!/usr/bin/env python3
"""Append Session 2026-06-03 addendum to Reference Architecture DOCX.

APPEND-ONLY per docx protocol — uses python-docx, never raw XML edit.
Idempotent: guarded by a marker-heading check. Run once.

Covers: (1) entry/exit grade surfacing on Journal + open positions,
(2) agent & Hermes workflow documentation + Hermes live-coordinator state correction.
"""
from docx import Document
from lxml import etree

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Entry/Exit Grade Surfacing & Workflow Docs (Session 2026-06-03)"


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def borders(t):
    tbl = t._tbl
    pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b += f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b + '</w:tblBorders>'))


def main():
    doc = Document(DOCX_PATH)
    for p in doc.paragraphs:
        if MARKER in (p.text or ""):
            print("present, skip")
            return

    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MARKER

    doc.add_paragraph(
        "Surfacing fix: the entry/exit grade data (trade_backtest_results, 59 graded; ~74/76 Schwab "
        "closed trades) and AI Trade Eval already existed but were invisible outside the Backtest "
        "sub-tabs. This session wired them into the surfaces where the operator looks for them.")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Journal entry/exit grade column"
    doc.add_paragraph(
        "v3 JournalHub trade log gains an inline Grade column (entry/exit A–D, colored) for every "
        "closed trade. Backend: _attach_backtest_grades() enriches /api/v2/journal and "
        "/api/v2/automated-trade-journal rows by joining trade_backtest_results on trade_key "
        "({symbol}:{account}:{close_date}), adding entry_grade/exit_grade/entry_rsi/left_on_table_20d. "
        "Paper trades are not backtest-graded and show '—'. Retrospective/advisory only.")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Open-position entry-setup rating"
    doc.add_paragraph(
        "v3 Trading hub Open Trades tab shows an 'entry setup ~N' badge per open position, matched by "
        "symbol against the setup-quality prior (/api/v2/atm/setup-advisory). Open trades have no exit "
        "grade yet, so this reflects the RSI-band entry prior at proposal time. Advisory-only — never "
        "gates execution.")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Agent & Hermes workflow documentation"
    doc.add_paragraph(
        "New canonical reference docs/AGENT_AND_HERMES_WORKFLOWS.md captures both fleets end-to-end "
        "(roster, schedules, allocation chain Maria→Steph→Risk→Tax→Alex, Hermes coordinator workflow, "
        "per-tick caps, safety controls, v3 surfaces), linked from MASTER §11 and §18b. Corrected "
        "drift in §18b: the Hermes Chief Coordinator runs the fleet live on a */15 flock-guarded cron "
        "(Operator Directive B, 2026-06-02; verified live 2026-06-03) with bounded, reversible research "
        "auto-promotion — this concerns research intelligence only and does NOT relax any trade/proposal "
        "gate. Caps per tick: librarian 10, autonomous loop 3/sub-loop, promote 10, embed 2.")

    t = doc.add_table(rows=1, cols=3)
    borders(t)
    hd = t.rows[0].cells
    hd[0].text, hd[1].text, hd[2].text = "Surface", "What changed", "Source endpoint(s)"
    for s, w, e in [
        ("v3 Journal trade log", "Inline entry/exit Grade column (A–D)", "/api/v2/journal, /api/v2/automated-trade-journal"),
        ("v3 Trading → Open Trades", "Per-position 'entry setup ~N' badge", "/api/v2/open-trades + /api/v2/atm/setup-advisory"),
        ("v3 Strategy → Backtest", "Entry Quality + AI Trade Eval tabs (verified, coverage diagnosed)", "/api/v2/journal/backtest-summary, /api/v2/backtesting/trade-evaluations"),
        ("Docs", "AGENT_AND_HERMES_WORKFLOWS.md + MASTER §11/§18b", "n/a"),
    ]:
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text = s, w, e

    doc.save(DOCX_PATH)
    print("addendum appended")


if __name__ == "__main__":
    main()

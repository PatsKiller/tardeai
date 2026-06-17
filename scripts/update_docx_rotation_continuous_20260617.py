#!/usr/bin/env python3
"""Append the Rotation Intelligence sleeve-balance / amount-range / continuous-loop addendum to the
Reference Architecture DOCX. APPEND-ONLY, idempotent (marker-guarded). Mirrors update_docx_* protocol.
As-built 2026-06-17; full detail in docs/project/ROTATION_LLM_ADVISOR.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation Intelligence - sleeve balance, advisory amount ranges, continuous loop (2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "Rotation Intelligence is now sector-aware, range-aware, and continuous. Everything here stays "
        "advisory: nothing is placed, no broker API is called, no dollar amount is ever auto-executed.")
    doc.add_paragraph(
        "Sleeve balance. The rotation summary reads the portfolio look-through theme weights against operator "
        "comfort targets in config/rotation_sector_targets.json and returns sector_overweights (theme, "
        "percent, target, excess percent, excess dollars, and the top holdings driving the sleeve) and "
        "sector_underweights (theme, percent, floor, gap). A new Sleeve Balance section on the page shows the "
        "overweight sleeves with the dollars over target and the underweight sleeves with the redeploy gap. "
        "The excess dollars are a review range only - the operator confirms sizing, tax and account fit.")
    doc.add_paragraph(
        "Advisory amount ranges. Each rebalance-from-research idea now carries a review_amount_range of low, "
        "high and a basis string, computed as roughly five to fifteen percent of the trim holding's value. "
        "The card shows the range with the basis line 'advisory, operator-confirmed, not auto-placed'. No "
        "idea is ever sized or placed automatically.")
    doc.add_paragraph(
        "Continuous loop and research wiring. POST /api/v2/rotation/research-gaps seeds watch_directives for "
        "the underweight sleeves and the named rotate-in candidates so TradeAI and Hermes research the gaps; "
        "the directives feed the discovery and watchlist sweep, and the names return as rotate-in candidates. "
        "POST /api/v2/rotation/feedback records the operator's review of an idea into the "
        "llm_feedback_observations learning table. A weekly cron, scripts/rotation_rebalance_digest.py, "
        "computes the summary, seeds the research gaps and sends a Telegram digest of overweight sleeves, "
        "underweight sleeves being researched, and advisory rebalance ideas with their review ranges. It is "
        "localhost-only, uses no external network and no Grok or xAI API, and places nothing.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation continuous addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation continuous addendum appended + saved")


if __name__ == "__main__":
    main()

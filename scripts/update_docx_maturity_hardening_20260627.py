#!/usr/bin/env python3
"""Append the PO/P0/P1 execution-safety hardening section (maturity 4.95) to the Reference
Architecture DOCX. APPEND-ONLY, idempotent (marker-guarded), python-docx only — never raw XML.
A timestamped .bak is written before saving. As-built 2026-06-27 (PR #6, merge fbe2048e)."""
import shutil
from datetime import datetime, timezone

from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Execution-safety hardening to maturity 4.95 (PO/P0/P1, 2026-06-27)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)

    doc.add_paragraph(
        "A focused hardening pass raised the platform's independently-scored execution maturity from a "
        "release-blocked 3.75 to 4.95 out of 5, landed on main through pull request six (merge commit "
        "fbe2048e) with the release-readiness CI workflow green. The exercise changed no trading behaviour "
        "and crossed no live broker write path; it made the safety architecture that already existed legible, "
        "machine-provable, and resistant to the failure modes that let a plausible-but-wrong artifact pass as "
        "evidence. The operator confirmation and two-factor path was treated as immutable throughout — it was "
        "neither widened, weakened, nor bypassed — and large-language models remain advisory only, unable to "
        "unlock execution.")

    _head(doc, 3, "Readiness, evidence binding, and broker-truth")
    doc.add_paragraph(
        "Central execution readiness now evaluates in four explicit modes — preflight, submit, dry-run, and "
        "audit — so that an order awaiting the operator's confirmation is reported as operator-required rather "
        "than conflated with a deterministic safety block. The transport uses submit-mode, which fails closed "
        "and hard-requires the two-factor state before an order can become submit-ready; an unrecognised mode "
        "is treated as the strictest submit path rather than defaulting open. Operator approval is bound to its "
        "evidence through separate, like-to-like hashes — approval bundle, readiness, quote, risk, chain, and "
        "model are each hashed and revalidated against their own kind — so a quote that drifts past tolerance, "
        "a changed risk posture, a materially different option chain, an expired or reused approval, or a kill "
        "switch flipped after approval each blocks the submit, and an evidence bundle that cannot be regenerated "
        "fails closed instead of being trusted. Broker truth is authoritative after submit: order state is "
        "normalised across the full Schwab status taxonomy, partial fills are preserved explicitly, nothing is "
        "marked filled or working ahead of the broker's own acknowledgement, and an idempotency fence plus a "
        "repeatable reconciliation report prevent internal state from outrunning the broker.")

    _head(doc, 3, "Release readiness, the maturity score, and the self-referential fixpoint")
    doc.add_paragraph(
        "Release readiness is now tri-state — PASS, WARN_NON_LIVE_ADJACENT, or FAIL — and distinguishes a clean "
        "tree from one whose only dirty files are the regenerated diligence and evidence artifacts. Any dirty "
        "live-broker, secrets, or execution-adjacent source is a hard FAIL; documented runtime-generated "
        "artifacts are a justified non-live-adjacent warning that does not gate a release. The maturity score is "
        "a transparent weighted computation bounded by hard caps that can only lower it: a release FAIL caps it "
        "at 3.75, an uncommitted live-adjacent file caps it at 4.0, a missing readiness resolver or a failed "
        "no-bypass or write-policy check caps it lower still. The single insight that unlocked the score was "
        "that the platform was being held at 3.75 not by any defect but by five verified-yet-uncommitted broker "
        "scripts; committing them cleared the caps and the evidence earned 4.95. The diligence pack and release "
        "manifest are themselves generated artifacts, which produces a deliberate fixpoint: a clean checkout "
        "validates PASS, while any committed manifest necessarily reads WARN_NON_LIVE_ADJACENT because the act "
        "of generating it dirties itself — an honest, explicitly-justified state, not a defect.")

    _head(doc, 3, "Continuous proof and the source-only CI sandbox")
    doc.add_paragraph(
        "A single read-only proof now exists both locally and in GitHub Actions: it runs execution-state, "
        "release readiness, the Schwab write-policy validator, the no-broker-write-bypass scan, the evidence "
        "and readiness test suites, the order-lifecycle and reconciliation fixtures, the options hard-risk "
        "matrix, the audit-ledger chain, and a frontend smoke check, emitting one Markdown-and-JSON report with "
        "per-step command, exit code, duration, and result. Because a clean CI sandbox has no application "
        "database, the two write-policy guards that assert live deployment posture cannot run there; rather than "
        "fake a database — which would let those guards pass trivially and prove nothing — a source-only mode "
        "runs every code and source-level fence and skips the two database-state guards loudly, never silently "
        "passing them, deferring their proof to the deployed run where the validator scores a full twenty-seven "
        "of twenty-seven. The exit code propagates honestly, so the proof can no longer report green while its "
        "own report says fail. As built the workflow is green and the deployed proof is a full pass, with no "
        "broker writes performed at any point.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    bak = f"{DOCX_PATH}.bak_maturity_hardening_{stamp}"
    shutil.copy2(DOCX_PATH, bak)
    print(f"backup written: {bak}")
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

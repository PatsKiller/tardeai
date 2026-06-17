#!/usr/bin/env python3
"""Append the Unified Card Enrichment & Sweep Priority section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-17; full detail in docs/MASTER_SYSTEM_DOCUMENTATION.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Unified Card Enrichment & Sweep Priority (2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "Improvements to the unified card layer that powers the Watchlist, Portfolio, and Open Trades cards, "
        "plus the watchlist AI-enrichment cadence. Full detail: docs/MASTER_SYSTEM_DOCUMENTATION.md and "
        "CHANGELOG.md (2026-06-17).")

    _head(doc, 3, "Two-line company blurb and fund sectors")
    doc.add_paragraph(
        "The symbol_profiles table, served as /api/v2/symbol-cards, supplies every card a short 'what the "
        "company does' description, its sector, one-week performance versus its sector, analyst consensus, "
        "and recent news. build_symbol_profiles.py now writes the first TWO sentences of the business summary "
        "(a roughly two-line blurb) instead of one. ETFs have no sector in the data provider, so a curated ETF "
        "map fills them: the sector SPDRs map to their real GICS sector, while broad, bond, and income funds "
        "get an honest asset-class label (Fixed Income, Dividend Equity, and so on). Open-end mutual funds get "
        "a Morningstar-style category label (for example Fidelity Contrafund as Large-Cap Growth). Screener-"
        "discovered names that previously had no profile (JRSH, SPAI) are now profiled. Opaque retirement-plan "
        "fund codes and delisted CUSIPs stay blank because there is no name source to enrich them. A weekly "
        "cron keeps profiles fresh.")

    _head(doc, 3, "AI-enrichment stale flag tightened to one hour")
    doc.add_paragraph(
        "The watchlist AI enrichment runs about every 30 minutes during market hours, so the card's 'technicals "
        "stale' flag was tightened from two hours to one hour, and the 'AI Enriched' metric is now green only "
        "within one hour to match. The daily 'Validated' metric keeps its original color bands.")

    _head(doc, 3, "Hermes-rank priority tier in the enrichment sweep")
    doc.add_paragraph(
        "With more than 3,300 researched watchlist items and a Finviz-bounded per-run cap, a plain stalest-"
        "first rotation left high-value but no-directive researched cards (for example ELVN at Hermes rank 3, "
        "or SNOW) sitting 24 to 48 hours stale even though they are visible on the page. The sweep is now two-"
        "tier. A PRIORITY pool - directive-linked OR active OR Hermes rank within the top 150, about 162 items "
        "- is refreshed stalest-first at roughly 135 per run, a cycle of about 36 minutes, so the front-page "
        "cards stay within the one-hour stale flag. A reserved TAIL slice, one quarter of the cap, rotates the "
        "rest of the universe stalest-first so nothing is ever permanently starved. The per-run cap was raised "
        "from 150 to 180. Verified live: ELVN moved from 21.6 hours stale to fresh, SNOW from 42 hours to "
        "fresh, and the sweep enriched 174 of 174 selected items.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Card Enrichment section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Unified Card Enrichment & Sweep Priority section appended + saved")


if __name__ == "__main__":
    main()

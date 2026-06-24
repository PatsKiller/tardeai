#!/usr/bin/env python3
"""report_export.py — professional DOCX/PDF export for analyst reports."""
from __future__ import annotations

import json
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parent.parent
REPORT_OUT = PROJECT_ROOT / "data" / "portfolios" / "reports" / "analyst"


def _resolve_chart_path(chart_path: str | Path | None) -> Path | None:
    """Map web URL (/data/portfolios/...) or absolute path to a local PNG for embed."""
    if not chart_path:
        return None
    p = Path(str(chart_path))
    if p.exists():
        return p
    s = str(chart_path).lstrip("/")
    if s.startswith("data/"):
        candidate = PROJECT_ROOT / s
        if candidate.exists():
            return candidate
    name = p.name
    if name:
        for base in (
            REPORT_OUT / "charts",
            PROJECT_ROOT / "data" / "portfolios" / "reports" / "analyst" / "charts",
        ):
            candidate = base / name
            if candidate.exists():
                return candidate
    return None


def _rl(text: Any, *, br: bool = False) -> str:
    """Escape free text for reportlab Paragraph mini-markup (fixes the 'P&L;' bug where a
    bare ampersand was parsed as a malformed XML entity). Preserves intended <br/> breaks."""
    from xml.sax.saxutils import escape
    s = escape(str(text if text is not None else ""))
    if br:
        s = s.replace("\n", "<br/>")
    return s


def _rel_url(path: Path) -> str:
    try:
        rel = path.relative_to(PROJECT_ROOT)
        return "/" + str(rel).replace("\\", "/")
    except ValueError:
        return str(path)


_COVER_KPI_KEYS = (
    "recommendation", "price", "day_change_pct", "confidence_label",
    "thesis_status", "unrealized_pnl_pct", "portfolio_pct",
)

# Sections that render a curated KPI table beneath their prose (P2-1).
_KPI_TABLE_SECTIONS = (
    "header_context", "executive_summary", "fundamental_valuation",
    "analyst_predictions", "risk_assessment", "peer_comparison", "action_plan",
)


_CURRENCY_KPIS = (
    "price", "entry_price", "target_low", "target_mean", "target_high",
    "valid_low", "valid_high", "market_value", "unrealized_pnl",
)
_RATIO_KPIS = ("reward_risk", "pe", "forward_pe")


def _fmt_kpi(key: str, val: Any) -> str:
    if val is None:
        return "—"
    if key == "confidence" and isinstance(val, (int, float)) and val <= 1:
        return f"{float(val) * 100:.0f}%"
    if key == "analysts" and isinstance(val, (int, float)):
        return f"{int(val)}"
    if key.endswith("_pct") and isinstance(val, (int, float)):
        return f"{float(val):+.2f}%"
    if key in _CURRENCY_KPIS and isinstance(val, (int, float)):
        return f"${float(val):,.2f}"
    if key == "portfolio_value" and isinstance(val, (int, float)):
        return f"${float(val):,.0f}"
    if key in _RATIO_KPIS and isinstance(val, (int, float)):
        return f"{float(val):.1f}×" if key in ("pe", "forward_pe") else f"{float(val):.1f}:1"
    if isinstance(val, float):
        # never leak a raw float like 171.4635761769087
        return f"{val:,.2f}".rstrip("0").rstrip(".") if abs(val) < 1000 else f"{val:,.0f}"
    return str(val)


def _visual_caption(vis: dict) -> str:
    vtype = vis.get("type", "chart")
    if vis.get("caption"):
        return str(vis["caption"])
    if vtype == "thesis_validity_bar":
        return (
            f"Zone {vis.get('zone_status', '—')} · ${vis.get('price', '—')} "
            f"(entry ${vis.get('entry', '—')} · stop ${vis.get('stop', '—')} · target ${vis.get('target1', '—')})"
        )
    if vtype == "price_levels":
        return f"RSI {vis.get('rsi', '—')} · RVOL {vis.get('rvol', '—')}x"
    if vtype == "risk_profile":
        return f"Beta {vis.get('beta', '—')} · vol {vis.get('volatility_w_pct', '—')}%"
    return vtype.replace("_", " ").title()


def export_docx(report: dict, output_path: Path | None = None) -> dict:
    """Export structured analyst report to Word (v3 premium layout)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    meta = report.get("meta") or {}
    stem = meta.get("symbol") or meta.get("report_type") or "report"
    ts = datetime.now().strftime("%Y%m%d")
    if output_path is None:
        REPORT_OUT.mkdir(parents=True, exist_ok=True)
        output_path = REPORT_OUT / f"{stem}_{ts}.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    sections = report.get("sections") or []
    exec_sec = next((s for s in sections if s.get("id") == "executive_summary"), {})
    kpis = dict(meta.get("kpis") or {})
    for k, v in (exec_sec.get("metrics") or {}).items():
        if v is not None and k not in kpis:
            kpis[k] = v

    # Cover
    title = doc.add_heading(meta.get("title") or "Analyst Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    sub.add_run("Trade AI v12 · Analyst Research\n").bold = True
    sub.add_run(f"Generated: {meta.get('generated_at', datetime.now().isoformat())}\n")
    if meta.get("company"):
        sub.add_run(f"{meta.get('company')}")
        if meta.get("sector"):
            sub.add_run(f" · {meta['sector']}")
        sub.add_run("\n")
    if meta.get("symbol"):
        sub.add_run(f"Symbol: {meta['symbol']}\n")
    if meta.get("version"):
        sub.add_run(f"Report schema v{meta['version']}\n")
    doc.add_paragraph("")

    strip_keys = [k for k in _COVER_KPI_KEYS if kpis.get(k) is not None]
    if strip_keys:
        ktable = doc.add_table(rows=2, cols=len(strip_keys))
        ktable.style = "Light Grid Accent 1"
        for i, k in enumerate(strip_keys):
            ktable.rows[0].cells[i].text = str(k).replace("_", " ").title()
            ktable.rows[1].cells[i].text = _fmt_kpi(k, kpis[k])
        doc.add_paragraph("")

    for co in exec_sec.get("callouts") or []:
        p = doc.add_paragraph()
        p.add_run(f"{co.get('label', 'Note')}: ").bold = True
        run = p.add_run(str(co.get("text") or ""))
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    doc.add_paragraph("")

    doc.add_heading("Contents", level=1)
    for i, sec in enumerate(sections, 1):
        doc.add_paragraph(f"{i}. {sec.get('title', sec.get('id', 'Section'))}", style="List Number")
    doc.add_page_break()

    skip_ids = {"agent_synthesis", "agent_performance_note", "ensemble_validation"}
    for sec in sections:
        sid = sec.get("id")
        if sid in skip_ids and any(s.get("id") == "intelligence_view" for s in sections):
            continue
        doc.add_heading(sec.get("title") or sid or "Section", level=1)
        if sec.get("content"):
            p = doc.add_paragraph(str(sec["content"]))
            p.paragraph_format.space_after = Pt(10)

        for co in sec.get("callouts") or []:
            cp = doc.add_paragraph()
            cp.add_run(f"{co.get('label', 'Action')}: ").bold = True
            cr = cp.add_run(str(co.get("text") or ""))
            cr.font.color.rgb = RGBColor(0x0f, 0x6d, 0x3a)

        metrics = sec.get("metrics") or {}
        show_metrics = sid in _KPI_TABLE_SECTIONS
        if metrics and show_metrics:
            rows = [[str(k).replace("_", " ").title(), _fmt_kpi(k, v)]
                    for k, v in metrics.items() if v is not None and k != "text"]
            if rows:
                table = doc.add_table(rows=1 + len(rows), cols=2)
                table.style = "Light Grid Accent 1"
                table.rows[0].cells[0].text = "Metric"
                table.rows[0].cells[1].text = "Value"
                for i, (label, val) in enumerate(rows, 1):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = val
                doc.add_paragraph("")

        bullets = sec.get("bullets") or []
        if sid == "intelligence_view" and bullets:
            bullets = bullets[:1]
        for bullet in bullets[:4]:
            doc.add_paragraph(str(bullet), style="List Bullet")

        agents = sec.get("agents") or []
        if agents:
            atable = doc.add_table(rows=1 + len(agents), cols=3)
            atable.style = "Light Grid Accent 1"
            hdr = atable.rows[0].cells
            hdr[0].text = "Agent"
            hdr[1].text = "Recommendation"
            hdr[2].text = "Weight"
            for i, ag in enumerate(agents, 1):
                row = atable.rows[i].cells
                row[0].text = str(ag.get("agent") or "—")
                row[1].text = str(ag.get("recommendation") or "—")
                row[2].text = str(ag.get("weight") or ag.get("relevance") or "—")[:48]
            doc.add_paragraph("")

    chart_visuals = [v for v in (report.get("visuals") or []) if _resolve_chart_path(v.get("chart_path"))]
    if chart_visuals:
        doc.add_page_break()
        doc.add_heading("Visual Summary", level=1)
        for vis in chart_visuals:
            cap = _visual_caption(vis)
            cp = doc.add_paragraph()
            cp.add_run(cap).italic = True
            chart_path = _resolve_chart_path(vis.get("chart_path"))
            if chart_path:
                try:
                    doc.add_picture(str(chart_path), width=Inches(6.2))
                except Exception:
                    pass
            doc.add_paragraph("")

    sources = report.get("sources") or []
    if sources:
        doc.add_heading("Data Sources", level=1)
        for src in sources:
            doc.add_paragraph(
                f"{src.get('label', src.get('id'))}"
                + (f" ({src.get('count')} rows)" if src.get("count") else "")
                + (" ✓" if src.get("present") else ""),
                style="List Bullet",
            )

    doc.add_paragraph("")
    footer = doc.add_paragraph("— End of Report —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    size_kb = round(output_path.stat().st_size / 1024, 1)
    return {
        "ok": True,
        "format": "docx",
        "path": str(output_path),
        "url": _rel_url(output_path),
        "filename": output_path.name,
        "size_kb": size_kb,
    }


def export_pdf_reportlab(report: dict, output_path: Path) -> dict:
    """Native PDF export via reportlab (v3 premium layout)."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    meta = report.get("meta") or {}
    sections = report.get("sections") or []
    exec_sec = next((s for s in sections if s.get("id") == "executive_summary"), {})
    kpis = dict(meta.get("kpis") or {})
    for k, v in (exec_sec.get("metrics") or {}).items():
        if v is not None and k not in kpis:
            kpis[k] = v

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Heading1"], fontSize=18, spaceAfter=10)
    h1 = ParagraphStyle("H1", parent=styles["Heading2"], fontSize=13, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1e3a5f"))
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=0, spaceAfter=4)
    callout = ParagraphStyle("Callout", parent=body, fontSize=10, leading=13, textColor=colors.HexColor("#0f6d3a"), spaceAfter=6)

    story = []
    story.append(Paragraph(meta.get("title") or "Analyst Report", title_style))
    sub = f"Trade AI v12 · Analyst Research<br/>Generated: {meta.get('generated_at', '')}"
    if meta.get("symbol"):
        sub += f"<br/>Symbol: <b>{meta['symbol']}</b>"
    if meta.get("version"):
        sub += f" · v{meta['version']}"
    story.append(Paragraph(sub, body))
    story.append(Spacer(1, 0.15 * inch))

    strip_keys = [k for k in _COVER_KPI_KEYS if kpis.get(k) is not None]
    if strip_keys:
        rows = [strip_keys, [_fmt_kpi(k, kpis[k]) for k in strip_keys]]
        tbl = Table(rows, colWidths=[6.0 / max(len(strip_keys), 1) * inch] * len(strip_keys))
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.1 * inch))

    for co in exec_sec.get("callouts") or []:
        story.append(Paragraph(
            f"<b>{_rl(co.get('label', 'Note'))}:</b> {_rl(co.get('text') or '', br=True)}",
            callout,
        ))

    story.append(Paragraph("Contents", h1))
    for i, sec in enumerate(sections, 1):
        story.append(Paragraph(f"{i}. {_rl(sec.get('title', sec.get('id', 'Section')))}", body))

    skip_ids = {"agent_synthesis", "agent_performance_note", "ensemble_validation"}
    for sec in sections:
        sid = sec.get("id")
        if sid in skip_ids and any(s.get("id") == "intelligence_view" for s in sections):
            continue
        story.append(Paragraph(_rl(sec.get("title") or sid or "Section"), h1))
        if sec.get("content"):
            story.append(Paragraph(_rl(sec["content"], br=True), body))
        for co in sec.get("callouts") or []:
            story.append(Paragraph(
                f"<b>{_rl(co.get('label', 'Action'))}:</b> {_rl(co.get('text') or '')}",
                callout,
            ))
        show_metrics = sid in _KPI_TABLE_SECTIONS
        metrics = sec.get("metrics") or {}
        if metrics and show_metrics:
            rows = [["Metric", "Value"]] + [
                [str(k).replace("_", " ").title(), _fmt_kpi(k, v)]
                for k, v in metrics.items() if v is not None and k != "text"
            ]
            if len(rows) > 1:
                tbl = Table(rows, colWidths=[2.2 * inch, 3.8 * inch])
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.08 * inch))
        bullets = (sec.get("bullets") or [])[:4]
        if sid == "intelligence_view" and bullets:
            bullets = bullets[:1]
        for b in bullets:
            story.append(Paragraph(f"• {_rl(b)}", bullet))
        agents = sec.get("agents") or []
        if agents:
            rows = [["Agent", "Rec", "Weight"]] + [
                [
                    str(ag.get("agent") or "—"),
                    str(ag.get("recommendation") or "—"),
                    str(ag.get("weight") or ag.get("relevance") or "—")[:40],
                ]
                for ag in agents
            ]
            tbl = Table(rows, colWidths=[1.6 * inch, 0.9 * inch, 3.5 * inch])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.08 * inch))

    chart_visuals = [v for v in (report.get("visuals") or []) if _resolve_chart_path(v.get("chart_path"))]
    if chart_visuals:
        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Visual Summary", h1))
        for vis in chart_visuals:
            story.append(Paragraph(_visual_caption(vis), ParagraphStyle("Cap", parent=body, fontSize=9, textColor=colors.grey)))
            chart_path = _resolve_chart_path(vis.get("chart_path"))
            if chart_path:
                try:
                    story.append(Image(str(chart_path), width=6.2 * inch, height=2.4 * inch))
                except Exception:
                    pass
            story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("— End of Report —", ParagraphStyle("Footer", parent=body, alignment=1, textColor=colors.grey)))
    doc.build(story)
    return {
        "ok": True,
        "format": "pdf",
        "path": str(output_path),
        "url": _rel_url(output_path),
        "filename": output_path.name,
        "size_kb": round(output_path.stat().st_size / 1024, 1),
        "engine": "reportlab",
    }


def export_pdf(docx_path: Path, pdf_path: Path | None = None, report: dict | None = None) -> dict:
    """Convert DOCX to PDF via LibreOffice, or reportlab fallback when report JSON provided."""
    docx_path = Path(docx_path)
    pdf_path = pdf_path or docx_path.with_suffix(".pdf")
    if report is not None:
        try:
            return export_pdf_reportlab(report, pdf_path)
        except Exception as e:
            pass  # fall through to LibreOffice attempt
    if not docx_path.exists():
        if report is not None:
            return export_pdf_reportlab(report, pdf_path)
        return {"ok": False, "error": "docx not found"}
    soffice = PROJECT_ROOT / "scripts" / "office" / "soffice.py"
    cmd = None
    if soffice.exists():
        cmd = ["python3", str(soffice), "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(pdf_path.parent)]
    else:
        for bin_name in ("libreoffice", "soffice"):
            try:
                subprocess.run([bin_name, "--version"], capture_output=True, check=True)
                cmd = [bin_name, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)]
                break
            except Exception:
                continue
    if not cmd:
        if report is not None:
            return export_pdf_reportlab(report, pdf_path)
        return {"ok": False, "error": "PDF export requires LibreOffice (soffice) or report JSON for reportlab"}
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=120)
    except Exception as e:
        if report is not None:
            return export_pdf_reportlab(report, pdf_path)
        return {"ok": False, "error": str(e)[:200]}
    if not pdf_path.exists():
        alt = pdf_path.parent / f"{docx_path.stem}.pdf"
        if alt.exists():
            pdf_path = alt
        elif report is not None:
            return export_pdf_reportlab(report, pdf_path)
        else:
            return {"ok": False, "error": "PDF conversion produced no file"}
    return {
        "ok": True,
        "format": "pdf",
        "path": str(pdf_path),
        "url": _rel_url(pdf_path),
        "filename": pdf_path.name,
        "size_kb": round(pdf_path.stat().st_size / 1024, 1),
        "engine": "libreoffice",
    }


def export_report(
    report: dict,
    fmt: str = "docx",
    *,
    output_stem: str | None = None,
    in_place: bool = False,
) -> dict:
    """Export report to docx and optionally pdf.

    in_place=True writes to {stem}.docx / {stem}.pdf (no timestamp) — used for living prospectus files.
    """
    fmt = (fmt or "docx").lower()
    REPORT_OUT.mkdir(parents=True, exist_ok=True)
    stem = output_stem or report.get("meta", {}).get("symbol") or report.get("meta", {}).get("report_type") or "report"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = REPORT_OUT / (f"{stem}.docx" if in_place else f"{stem}_{ts}.docx")
    docx_result = export_docx(report, docx_path)
    if not docx_result.get("ok"):
        return docx_result
    if fmt == "docx":
        return docx_result
    if fmt == "pdf":
        pdf_path = REPORT_OUT / (f"{stem}.pdf" if in_place else f"{stem}_{ts}.pdf")
        pdf_result = export_pdf(docx_path, pdf_path, report=report)
        pdf_result["docx"] = docx_result
        return pdf_result
    if fmt == "json":
        json_path = REPORT_OUT / f"{stem}_{ts}.json"
        json_path.write_text(json.dumps(report, indent=2, default=str))
        return {"ok": True, "format": "json", "path": str(json_path), "url": _rel_url(json_path), "filename": json_path.name}
    return {"ok": False, "error": f"unsupported format: {fmt}"}
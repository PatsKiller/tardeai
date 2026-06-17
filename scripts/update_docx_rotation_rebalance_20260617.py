#!/usr/bin/env python3
"""Append the Rotation Intelligence rebalance-from-research addendum to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-17; full detail in docs/project/ROTATION_LLM_ADVISOR.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Rotation Intelligence - rebalance from research + Grok idea review (2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 3, MARKER)
    doc.add_paragraph(
        "Rotation Intelligence also recommends rotations into names the operator does not yet hold, sourced "
        "from research, and lets Grok review those ideas. This stays advisory: nothing is placed, no broker "
        "API is called, and no dollar amounts are invented.")
    doc.add_paragraph(
        "The rotation summary now returns two extra fields. research_candidates is the set of top non-held "
        "watchlist names with conviction - their Hermes rank, sector, analyst rating and upside - as rotate-in "
        "targets. research_rotation_ideas is a small set of advisory ROTATE_REVIEW pairs that match a "
        "trim-worthy real-ticker holding with a top research candidate; they carry no dollar amount and are "
        "explicitly flagged as not a model-supported signal, with sizing, tax and account fit left to the "
        "operator. Retirement-plan fund codes are excluded as rotate-from sources because a plan fund cannot "
        "rotate into an individual stock, and duplicate research names are removed.")
    doc.add_paragraph(
        "A new endpoint, POST /api/v2/rotation/grok-rebalance-review, sends those ideas plus the trim-holding "
        "facts and the research-candidate facts to Grok through the free OAuth proxy and asks for a per-idea "
        "verdict - reasonable to review or a poor fit, and why, considering conviction, concentration, and "
        "account or tax fit - followed by an overall WATCH or RESEARCH_MORE. It uses no API key and no paid "
        "API, invents no amounts, and calls no broker. The matching Run Grok Review and Grok Review These "
        "Ideas buttons surface the second opinions inline on the Rotation Intelligence page.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rotation rebalance addendum already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rotation rebalance addendum appended + saved")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Append the Hermes Global-Install Migration section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_session_* protocol.
Facts verified live on ms01 (2026-06-06): hermes v0.16.0, 5 profiles, tools disabled on tradeai*."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Hermes Global-Install Migration & Profile Model (2026-06-06)"


def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}":
            return p.style
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def append_hermes(doc):
    p = doc.add_paragraph()
    if h(doc, 2):
        p.style = h(doc, 2)
    p.text = MARKER

    doc.add_paragraph(
        "Hermes was promoted from a Trade-AI-gated sidecar install to a global/default install with named "
        "profiles. Verified live on ms01: CLI at ~/.local/bin/hermes -> ~/.local/share/hermes-agent-venv, "
        "Hermes Agent v0.16.0 (2026.6.5); default profile at ~/.hermes; per-profile configs under "
        "~/.hermes/profiles/. The old sidecar (hermes_sidecar/, v0.15.2) is RETAINED as rollback/migration "
        "source and audit evidence only — it is no longer the canonical runtime.")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Profiles & model policy"
    doc.add_paragraph(
        "Five profiles (all provider=custom -> local Ollama http://127.0.0.1:11434/v1): default (gemma3:4b, "
        "general assistant); tradeai (gemma3:4b, stable restricted Trade AI advisory); tradeai12b "
        "(gemma3:12b-ctx4k, experimental context-gated advisory); dev (unconfigured — future ChatGPT/Codex "
        "development route); serverops (unconfigured — future controlled server ops). Canonical entrypoints: "
        "hermes chat, tradeai chat, tradeai12b chat, dev chat, serverops chat.")
    doc.add_paragraph(
        "Model policy: gemma3:4b is the approved stable default (direct-Ollama exact-string + math canaries "
        "pass live). gemma3:12b is NOT approved unconstrained (operator-observed garbage under large/default "
        "context); it is stable only via native /api/generate at num_ctx=4096, exposed as the gemma3:12b-ctx4k "
        "alias used solely by tradeai12b (canaries pass live, advisory only). qwen3:14b must NOT be "
        "reintroduced as a Hermes default and is absent from the live model inventory. Codex is reserved for "
        "the future dev profile only (human-invoked development mode, never autonomous or Trade AI runtime).")

    p = doc.add_paragraph()
    if h(doc, 3):
        p.style = h(doc, 3)
    p.text = "Safety boundaries & retirement status"
    doc.add_paragraph(
        "tradeai and tradeai12b run with ALL toolsets DISABLED (web/browser/terminal/file/code_execution/"
        "vision/video) — advisory/research only, enforced by both config and SOUL. Their SOULs explicitly "
        "forbid placing orders, modifying stops, approving/promoting proposals, mutating broker/holdings/"
        "strategy data, and reading raw secrets/.env; the generic 'execute actions via your tools' instruction "
        "is absent from both. Sidecar retirement is staged and reversible: Stage A (preserve) is complete "
        "(snapshot backups/hermes_sidecar_snapshot_20260606_2007.tgz + SOUL archive under "
        "~/.hermes/migration_from_tradeai_sidecar_20260606/souls/); Stages B (curated migration), C (validate), "
        "and D (rename-not-delete + wrapper stubs) are operator-gated. No deletion, no gateway/Telegram/cron/"
        "systemd/Codex enablement, and no broker/trading changes were made. Detail: docs/hermes/"
        "HERMES_{GLOBAL_INSTALL_MIGRATION,PROFILE_MATRIX,MODEL_CANARY_STATUS,SIDECAR_RETIREMENT_PLAN}_20260606.md.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Hermes section already present, skip")
        return
    append_hermes(doc)
    doc.save(DOCX_PATH)
    print("Hermes global-install section appended + saved")


if __name__ == "__main__":
    main()

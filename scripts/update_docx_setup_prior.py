#!/usr/bin/env python3
"""Append Setup-Quality Prior (ATM advisory) addendum to Reference Architecture DOCX (append-only, idempotent)."""
from docx import Document
from lxml import etree
DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Setup-Quality Prior — ATM Proposal Advisory (Session 2026-06-02)"
def h(doc, lvl):
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {lvl}": return p.style
    return None
def borders(t):
    tbl=t._tbl; pr=tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    b='<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for e in ('top','left','bottom','right','insideH','insideV'): b+=f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    pr.append(etree.fromstring(b+'</w:tblBorders>'))
doc=Document(DOCX_PATH)
for p in doc.paragraphs:
    if MARKER in (p.text or ""): print("present, skip"); raise SystemExit
p=doc.add_paragraph();
if h(doc,2): p.style=h(doc,2)
p.text=MARKER
doc.add_paragraph(
 "Feedback loop closing the post-trade findings into forward advisement. ADVISORY-ONLY: never "
 "writes to paper_trade_proposals, never affects any gate or execution. setup_quality_prior.py "
 "(nightly 10 PM) distills entry grades (trade_backtest_results) + structured evals (trade_llm_reviews) "
 "into an RSI-band prior, then attaches a per-proposal advisory flag (caution/neutral/favorable) to "
 "recent proposals whose entry RSI falls in a historically weak band. Tables: setup_quality_prior, "
 "proposal_setup_advisory. Endpoint /api/v2/atm/setup-advisory. v3: prior panel in AI Trade Eval tab + "
 "caution badge on Trading-hub proposal rows. LLM prior is monotonic: RSI 40-55=60, 55-70=32, >70=10 "
 "('weak setup'). Every output labelled by sample size + confidence.")
t=doc.add_table(rows=1, cols=2); borders(t)
for c,x in zip(t.rows[0].cells,["RSI band (prior)","Learned (n / win / LLM score / verdict)"]):
    c.text=x;
    for r in c.paragraphs[0].runs: r.bold=True
for a,b in [("40-55","n=9 / 67% / 60 / good entry, poor exit"),("55-70","n=21 / 43% / 32 / weak setup"),(">70","n=27 / 59% / 10 / weak setup"),("<40","n=2 / 0% / — / low confidence")]:
    rc=t.add_row().cells; rc[0].text=a; rc[1].text=b
doc.save(DOCX_PATH); print("appended")

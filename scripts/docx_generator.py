"""docx_generator.py — Trade AI Word document generator.

Produces a rich .docx with:
  - Clickable Table of Contents with bookmarks
  - Executive summary
  - Delta / what changed section
  - Full scorecard table
  - Per-ticker dashboard cards (criteria met, pillar bars, LLM narrative)
  - Catalyst tape
  - TOS export block

Requires: python-docx >= 1.1
"""
from __future__ import annotations
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colors ────────────────────────────────────────────────────────────────────
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLUE       = RGBColor(0x1A, 0x73, 0xE8)
GREEN      = RGBColor(0x0F, 0x9D, 0x58)
AMBER      = RGBColor(0xF4, 0xB4, 0x00)
RED        = RGBColor(0xDB, 0x44, 0x37)
DARK_BG    = RGBColor(0x1E, 0x1E, 0x2E)
MUTED      = RGBColor(0x9A, 0x9A, 0xB0)

DECISION_COLOR = {"GO": GREEN, "WAIT": AMBER, "AVOID": RED}


# ── XML helpers ───────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str) -> None:
    """Apply background shade to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)

def _add_bookmark(paragraph, bookmark_id: int, bookmark_name: str) -> None:
    """Add a Word bookmark to a paragraph for ToC linking."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"),   str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, bm_start)
    paragraph._p.append(bm_end)

def _add_hyperlink_to_bookmark(paragraph, text: str, bookmark_name: str) -> None:
    """Add a clickable hyperlink pointing to an internal bookmark."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    r.append(rPr)
    r.append(t_elem)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


# ── Formatting helpers ────────────────────────────────────────────────────────

def _h(doc: Document, text: str, level: int = 1, bookmark_name: str = "", bookmark_id: int = 0) -> None:
    p = doc.add_heading(text, level=level)
    if bookmark_name:
        _add_bookmark(p, bookmark_id, bookmark_name)

def _p(doc: Document, text: str, bold: bool = False, color: RGBColor = None,
       size_pt: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    para.alignment = align

def _rule(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A73E8")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _criteria_bar(criteria_met: int, criteria_total: int) -> str:
    filled = round(criteria_met / criteria_total * 15) if criteria_total else 0
    return "●" * filled + "○" * (15 - filled) + f"  {criteria_met}/{criteria_total} criteria"

def _pillar_bar(score: int, max_score: int) -> str:
    filled = round(score / max_score * 12) if max_score else 0
    return "█" * filled + "░" * (12 - filled) + f" {score}/{max_score}"


# ── Table builders ────────────────────────────────────────────────────────────

def _scorecard_table(doc: Document, scored: List[Dict[str, Any]]) -> None:
    headers = ["Symbol", "Score", "Grade", "Decision", "RVOL", "Price", "Chg%", "Gap%", "Float M", "Catalyst"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = BLUE
        _shade_cell(hdr_cells[i], "#2C2C44")

    for t in scored:
        top = t.get("top_catalyst") or {}
        cat_snippet = (top.get("title") or "—")[:50]
        row = tbl.add_row().cells
        vals = [
            t["symbol"],
            str(t["score"]),
            t["grade"],
            t["decision"],
            f"{t.get('relative_volume', 0):.1f}x",
            f"${t.get('price', 0):.2f}",
            f"{t.get('change_percent', 0):+.1f}%",
            f"{t.get('gap_percent', 0):+.1f}%",
            f"{t.get('float_m', 0):.1f}M",
            cat_snippet,
        ]
        decision = t["decision"]
        dec_color = DECISION_COLOR.get(decision, MUTED)
        for i, val in enumerate(vals):
            row[i].text = val
            run = row[i].paragraphs[0].runs[0] if row[i].paragraphs[0].runs else row[i].paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if i == 3:  # decision column
                run.bold = True
                run.font.color.rgb = dec_color


def _ticker_card(doc: Document, ticker: Dict[str, Any], bookmark_id: int) -> None:
    sym = ticker["symbol"]
    company = ticker.get("company", "")
    score = ticker["score"]
    grade = ticker["grade"]
    decision = ticker["decision"]
    dec_color = DECISION_COLOR.get(decision, MUTED)

    # Section heading with bookmark
    heading_text = f"{sym}  ·  {company}" if company else sym
    _h(doc, heading_text, level=3, bookmark_name=f"ticker_{sym}", bookmark_id=bookmark_id)

    # Score + decision line
    score_line = doc.add_paragraph()
    r1 = score_line.add_run(f"Score: {score}/50  ·  Grade: {grade}  ·  ")
    r1.font.size = Pt(10)
    r2 = score_line.add_run(f"[ {decision} ]")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = dec_color

    # Criteria bar
    cbar = _criteria_bar(ticker.get("criteria_met", 0), ticker.get("criteria_total", 9))
    cp = doc.add_paragraph(cbar)
    cp.runs[0].font.name = "Courier New"
    cp.runs[0].font.size = Pt(9)

    # Market data mini-table
    md_tbl = doc.add_table(rows=2, cols=6)
    md_tbl.style = "Table Grid"
    headers = ["Price", "Change%", "Gap%", "RVOL", "Float M", "Volume"]
    vals = [
        f"${ticker.get('price', 0):.2f}",
        f"{ticker.get('change_percent', 0):+.1f}%",
        f"{ticker.get('gap_percent', 0):+.1f}%",
        f"{ticker.get('relative_volume', 0):.1f}x",
        f"{ticker.get('float_m', 0):.1f}M",
        f"{int(ticker.get('volume', 0)):,}",
    ]
    for i, h in enumerate(headers):
        c = md_tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        c.paragraphs[0].runs[0].font.color.rgb = MUTED
        _shade_cell(c, "#2C2C44")
    for i, v in enumerate(vals):
        c = md_tbl.rows[1].cells[i]
        c.text = v
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Pillar bars
    pb = ticker.get("pillar_breakdown", {})
    pillar_names = {"catalyst": "Catalyst", "relative_volume": "RVOL",
                    "price_action": "Price Action", "float": "Float", "price_range": "Price Range"}
    max_pts = {"catalyst": 15, "relative_volume": 12, "price_action": 10, "float": 8, "price_range": 5}
    for key, label in pillar_names.items():
        pts = pb.get(key, 0)
        maxp = max_pts[key]
        bar_text = f"  {label:<16}  {_pillar_bar(pts, maxp)}"
        bp = doc.add_paragraph(bar_text)
        bp.runs[0].font.name = "Courier New"
        bp.runs[0].font.size = Pt(8.5)

    # Criteria checklist
    criteria: Dict[str, bool] = ticker.get("criteria", {})
    check_labels = {
        "has_catalyst": "Has catalyst",
        "has_high_impact": "High-impact catalyst",
        "has_fresh_catalyst": "Fresh catalyst (<12h)",
        "rvol_above_3": "RVOL ≥ 3x",
        "rvol_above_5": "RVOL ≥ 5x",
        "float_under_50m": "Float ≤ 50M",
        "price_in_range": "Price in range ($0.50–$30)",
        "gap_above_5pct": "Gap ≥ 5%",
        "change_above_10pct": "Change ≥ 10%",
    }
    for key, label in check_labels.items():
        met = criteria.get(key, False)
        cp = doc.add_paragraph(f"  {'✅' if met else '⬜'}  {label}")
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        if cp.runs:
            cp.runs[0].font.size = Pt(9)

    # LLM Narrative
    narrative = ticker.get("llm_narrative", "")
    if narrative:
        doc.add_paragraph()
        np = doc.add_paragraph(narrative)
        if np.runs:
            np.runs[0].font.size = Pt(9.5)
            np.runs[0].italic = True

    # Top catalyst
    top = ticker.get("top_catalyst") or {}
    if top.get("title"):
        doc.add_paragraph()
        cat_p = doc.add_paragraph()
        r = cat_p.add_run(f"🔥 Top Catalyst  ·  {top.get('hours_old', '?'):.1f}h ago  ·  {top.get('source', '?')}  ·  {top.get('impact_tier', '').replace('_', ' ').title()}")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = BLUE
        tit_p = doc.add_paragraph(top.get("title", ""))
        if tit_p.runs:
            tit_p.runs[0].font.size = Pt(9.5)
        if top.get("url"):
            url_p = doc.add_paragraph(top["url"])
            if url_p.runs:
                url_p.runs[0].font.color.rgb = BLUE
                url_p.runs[0].font.size = Pt(8)

    _rule(doc)


# ── Main builder ──────────────────────────────────────────────────────────────

def generate_docx(
    scored_tickers: List[Dict[str, Any]],
    delta: Dict[str, Any],
    output_dir: Path,
    run_label: str,
    date_str: str,
) -> str:
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"trade_ai_{date_str}_{run_label}.docx"

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Cover ──────────────────────────────────────────────────────────────
    title_p = doc.add_heading("⚡ TRADE AI v10", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p = doc.add_paragraph(
        f"Run Window: {run_label}  ·  Date: {date_str}  ·  Generated: {datetime.now().strftime('%H:%M:%S ET')}"
    )
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub_p.runs:
        sub_p.runs[0].font.size = Pt(10)
    _rule(doc)

    # ── ToC (manual clickable links) ───────────────────────────────────────
    _h(doc, "Table of Contents", level=2, bookmark_name="toc", bookmark_id=1)
    toc_items = [
        ("Executive Summary",    "exec_summary"),
        ("What Changed",         "delta_section"),
        ("Full Scorecard",       "scorecard"),
        ("GO-Tier Picks",        "go_picks"),
        ("Catalyst Tape",        "cat_tape"),
        ("ThinkorSwim Export",   "tos_export"),
    ]
    for label, anchor in toc_items:
        tp = doc.add_paragraph("  → ")
        tp.paragraph_format.space_after = Pt(2)
        _add_hyperlink_to_bookmark(tp, label, anchor)

    # ── Executive Summary ──────────────────────────────────────────────────
    _h(doc, "Executive Summary", level=2, bookmark_name="exec_summary", bookmark_id=2)
    go_list   = [t for t in scored_tickers if t.get("decision") == "GO"]
    wait_list = [t for t in scored_tickers if t.get("decision") == "WAIT"]
    new_t     = delta.get("new_tickers", [])
    events    = delta.get("events", [])

    summary_rows = [
        ("Total Tickers Scanned", str(len(scored_tickers))),
        ("GO-Tier (≥40)",         str(len(go_list))),
        ("WAIT-Tier (30–39)",     str(len(wait_list))),
        ("New Tickers This Run",  str(len(new_t))),
        ("Delta Events",          str(len(events))),
        ("Top Score",             f"{scored_tickers[0]['score']} ({scored_tickers[0]['symbol']})" if scored_tickers else "—"),
    ]
    s_tbl = doc.add_table(rows=len(summary_rows)+1, cols=2)
    s_tbl.style = "Table Grid"
    for i, cell in enumerate(s_tbl.rows[0].cells):
        cell.text = ["Metric", "Value"][i]
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
        _shade_cell(cell, "#2C2C44")
    for i, (label, val) in enumerate(summary_rows, 1):
        s_tbl.rows[i].cells[0].text = label
        s_tbl.rows[i].cells[1].text = val

    # ── Delta section ──────────────────────────────────────────────────────
    _h(doc, "What Changed Since Last Run", level=2, bookmark_name="delta_section", bookmark_id=3)
    if events:
        for evt in events[:30]:
            etype = evt.get("event", "")
            sym   = evt.get("symbol", "")
            if etype == "NEW_TICKER":
                line = f"🆕 NEW  {sym}  —  Score: {evt.get('score')}  ·  {evt.get('decision')}"
            elif etype == "GRADE_UP":
                line = f"📈 UP   {sym}  —  {evt.get('prev_score')} → {evt.get('score')} (+{evt.get('score_delta')})"
            elif etype == "GRADE_DOWN":
                line = f"📉 DOWN {sym}  —  {evt.get('prev_score')} → {evt.get('score')} ({evt.get('score_delta')})"
            elif etype == "NEW_CATALYST":
                titles = "; ".join((evt.get("titles") or [])[:2])
                line = f"🔥 CAT  {sym}  —  {titles[:100]}"
            elif etype == "RVOL_THRESHOLD_CROSS":
                line = f"🚀 RVOL {sym}  —  Crossed {evt.get('threshold')}x  (now {evt.get('rvol'):.1f}x)"
            elif etype == "NEW_CRITERIA_MET":
                line = f"✅ CRIT {sym}  —  {', '.join(evt.get('new_criteria', []))}"
            elif etype == "TICKER_FADED":
                line = f"👻 FADE {sym}  —  Dropped out  (was {evt.get('last_score')})"
            else:
                line = f"• {etype} {sym}"
            ep = doc.add_paragraph(line)
            ep.paragraph_format.space_after = Pt(1)
            if ep.runs:
                ep.runs[0].font.size = Pt(9.5)
    else:
        doc.add_paragraph("No delta events — this is the first run or state is empty.")

    # ── Full Scorecard ─────────────────────────────────────────────────────
    _h(doc, "Full Scorecard", level=2, bookmark_name="scorecard", bookmark_id=4)
    _scorecard_table(doc, scored_tickers)
    doc.add_paragraph()

    # ── GO-tier cards ──────────────────────────────────────────────────────
    _h(doc, f"GO-Tier Picks — Top {min(5, len(go_list))}", level=2,
       bookmark_name="go_picks", bookmark_id=5)

    # Add ToC sub-links for each GO ticker
    for t in go_list[:5]:
        tp = doc.add_paragraph(f"  → ")
        tp.paragraph_format.space_after = Pt(1)
        _add_hyperlink_to_bookmark(tp, t["symbol"], f"ticker_{t['symbol']}")

    for i, ticker in enumerate(go_list[:5]):
        _ticker_card(doc, ticker, bookmark_id=100 + i)

    # ── Catalyst tape ──────────────────────────────────────────────────────
    _h(doc, "Catalyst Tape", level=2, bookmark_name="cat_tape", bookmark_id=6)
    for t in scored_tickers:
        cats = t.get("catalysts", [])
        if not cats:
            continue
        doc.add_paragraph(f"{t['symbol']}  ·  {t.get('company', '')}").runs[0].bold = True
        for cat in cats[:6]:
            age = cat.get("hours_old", 0)
            src = cat.get("source", "?")
            impact = cat.get("impact_tier", "").replace("_", " ").title()
            line = f"  [{impact}  {age:.1f}h  {src}]  {cat.get('title', '')}"
            cp = doc.add_paragraph(line)
            cp.paragraph_format.space_after = Pt(1)
            if cp.runs:
                cp.runs[0].font.size = Pt(8.5)

    # ── TOS Export ────────────────────────────────────────────────────────
    _h(doc, "ThinkorSwim Export", level=2, bookmark_name="tos_export", bookmark_id=7)
    tos_symbols = "  ".join(t["symbol"] for t in go_list[:20])
    doc.add_paragraph(f"GO-tier symbols:  {tos_symbols or '—'}")
    doc.add_paragraph("Import: Charts → Watchlists → Import Watchlist → select .tst file")

    doc.save(str(docx_path))
    return str(docx_path)

#!/usr/bin/env python3
"""Append the Stop Management Architecture section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-15; full detail in docs/brokers/stop-management-architecture.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Stop Management Architecture (Stage 2c, LIVE 2026-06-15)"


def h(doc, lvl):
    # Resolve the heading STYLE object by name. Iterate doc.styles (key lookup fails on the doc's
    # duplicate style names) and don't rely on an existing paragraph already using it — Heading 3 was
    # defined-but-unused, which made the old paragraph-search return None and drop subsections to body.
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "Protective stops are live across the whole book. Protective SELL stops are the safe direction — they "
        "sell-to-close an existing long and can never open a position or overspend — which is what makes a "
        "standing capability acceptable. Live accounts (Schwab taxable + both IRAs) are MANUAL with per-order "
        "two-factor approval; the Alpaca paper account is managed AUTOMATICALLY for best risk:reward. Routing is "
        "decided by the account, never the client: an account with a trading API + write enabled submits live; "
        "anything else (Fidelity 401k — no API) returns an exact thinkorswim ticket to place by hand. The broker "
        "is the source of truth for all monitoring. Full detail: docs/brokers/stop-management-architecture.md.")

    _head(doc, 3, "Account capability matrix")
    doc.add_paragraph(
        "schwab_taxable / schwab_roth_ira / schwab_rollover_ira: live protective stops, any symbol up to "
        "$250,000 per order, MANUAL + 2FA. fidelity_401k: no trading API — ticket only. alpaca_paper: managed "
        "AUTOMATICALLY (ratchet for R:R), no 2FA (it is paper).")

    _head(doc, 3, "The gating stack (Schwab live path)")
    doc.add_paragraph(
        "Every Schwab submit passes the full stack before any HTTP reaches the broker; any failure fails closed. "
        "(1) protective_stop_policy.evaluate() — the committed, tamper-evidenced envelope: ENABLED master gate; "
        "account in effective_account_allowlist() (taxable plus the IRAs when IRA_PROTECTIVE_ENABLED); "
        "instruction SELL (sell-to-close); order type STOP/STOP_LIMIT/TRAILING_STOP; stop strictly below current "
        "price; placed stop within +/-8% of the engine-advised stop; quantity <= shares held; notional "
        "<= $250,000. (2) execution_guard.authorize() — the meta marker PROTECTIVE_STOP_2C routes the intent "
        "through protective_stop_policy INSTEAD of the BUY canary gate and skips the canary order cap; the "
        "STANDING unlock _protective_unlocked() = policy ENABLED + system_controls['protective_stops_enabled']="
        "'true' (no expiring ARM session — that remains canary-only); per-order 2FA required. (3) "
        "schwab_transport._pilot_preconditions(kind='protective_stop') — account in the protective allowlist + "
        "broker_accounts.api_write_enabled true. (4) schwab_transport.place_order(kind='protective_stop') — "
        "persist a reconcile-anchor row, POST, consume the 2FA, read back; tagged so the canary order budget is "
        "untouched. The Stage-2b canary BUY pilot is unchanged and still requires the ARM session.")

    _head(doc, 3, "Per-order two-factor approval")
    doc.add_paragraph(
        "approval_service with REQUIRED_CHANNELS=1 — EITHER channel confirms: type the ticker exactly in the web "
        "modal (anti-fat-finger), OR enter a 6-digit one-time code sent to BOTH Telegram and email (or tap "
        "Approve in the Telegram message). One order at a time; codes are single-use, TTL 10 minutes, burned on "
        "submit. Every Schwab place and Modify requires this; Cancel is the safe direction and does not.")

    _head(doc, 3, "Lifecycle: place, modify, cancel")
    doc.add_paragraph(
        "Place: POST /api/v2/holdings/protective-stop (request) then /protective-stop/confirm (2FA + submit). "
        "Modify (one click, after a price rise): a replace_order_id threads through the intent; on confirm the "
        "old stop is cancelled FIRST and only then is the new one placed — if the cancel fails the new stop is "
        "NOT placed, so there is never a double or oversized stop. Cancel: /protective-stop/cancel via "
        "schwab_transport.cancel_order (no 2FA); it refuses non-pilot orders, so a stop placed manually in "
        "thinkorswim must be cancelled there.")

    _head(doc, 3, "Monitoring, health agent, Hermes, and Grok")
    doc.add_paragraph(
        "stop_lifecycle_monitor.py reads the broker's live working stops across Schwab and Alpaca, cross-"
        "references shares held now, and classifies each by lifecycle (working / near_trigger <=2% / triggered / "
        "filled / cancelled / orphaned), coverage (full / oversized / partial / orphaned — a standing GTC stop "
        "does not auto-resize when you trim), proximity (% from price down to the stop), and health (ok / warn / "
        "alert). It persists a snapshot to the stop_lifecycle table and is exposed at GET /api/v2/stops/"
        "lifecycle (cached 45s). The Open Trades card shows a green PROTECTED banner with Modify + Cancel, plus "
        "coverage and percent-from-stop warnings. The health agent stop_health_check.py (cron every 10 minutes, "
        "market hours) escalates orphaned / oversized / triggered / near-trigger conditions via the SIEM "
        "(save_alert_event), Telegram (central router, deduped), and the system_health_events watchdog log, AND "
        "writes each condition into the Hermes research stream (hermes_research_intelligence, "
        "research_type='stop_health') so Hermes monitors the stops. grok_stop_review.py adds an external-LLM "
        "risk:reward curation of every live stop beyond the technical rules (grade, should-trail, "
        "recommendation, confidence), persisted to stop_grok_reviews and surfaced on the card as 'reviewed by "
        "GROK' (advisory only). All monitoring layers are read-only on the broker.")

    _head(doc, 3, "Alpaca automatic stop management (paper only)")
    doc.add_paragraph(
        "alpaca_stop_manager.py (cron every 20 minutes, market hours) manages ONLY the Alpaca paper account. For "
        "each open paper position it asks strategy_trailing_policy.recommend_stop for the risk:reward-optimal "
        "stop and ratchets the live Alpaca stop UP to it (cancel + re-place via the paper API). It is "
        "ratchet-only (never lowers a stop, never places at or above the current price, ignores sub-0.25% "
        "nudges), automatic (no 2FA — paper), audited and SIEM-logged, and persists the new stop and order id to "
        "paper_trades. It respects Hard Rule 7 — the paper pipeline owns paper execution; it uses the Alpaca "
        "paper API directly and never the Schwab guard. Schwab accounts are never auto-managed.")

    _head(doc, 3, "Safety invariants")
    doc.add_paragraph(
        "SELL-to-close only; quantity <= held; stop strictly below price; notional <= $250k. Every Schwab place "
        "and modify requires per-order 2FA, without exception. Fidelity 401k can never API-submit. Alpaca "
        "auto-management is ratchet-only and paper-only and never touches a Schwab order. The risk envelope is "
        "commit-only and tamper-evidenced against git HEAD by the write-policy validator; UI, config, and the DB "
        "cannot widen it. Modify cancels the old stop before placing the new, and on a cancel failure it does "
        "not place. Revoke all Schwab protective submits instantly by setting "
        "system_controls['protective_stops_enabled'] to anything other than 'true'.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Stop Management section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Stop Management Architecture section appended + saved")


if __name__ == "__main__":
    main()

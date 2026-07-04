#!/usr/bin/env python3
"""Append the July-4 incident-set and infrastructure-hardening section to the Reference
Architecture DOCX. APPEND-ONLY, idempotent (marker-guarded), python-docx only, timestamped
.bak first. As-built 2026-07-04 (PR #83 merge f99133da, PR #84 merge 2416fa3a, plus
off-repo OpenClaw 2026.6.11 upgrade and systemd unit hardening)."""
import shutil
from datetime import datetime, timezone
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "July-4 incident set: queue SLA, governor transaction death, infra monitoring (2026-07-04)"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "A holiday-morning audit answering the question \"did everything run properly\" surfaced four "
        "distinct failure classes that had been running blind, all fixed and merged the same day through "
        "pull requests eighty-three and eighty-four with release-readiness green. First, decision-feeding "
        "agent jobs (full analysis, proposal review, research gaps, events) had been starved for over two "
        "days because the watchlist job worker ordered its queue by symbol tier alone — a continuous "
        "stream of scheduled research on high-ranked symbols always outranked them. A canonical "
        "time-sensitive request-type list now leads the worker's pick order and is imported by the health "
        "agent's SLA check so the drain order and the alarm can never drift apart; the backlog of "
        "seventy-four stuck jobs drained within hours. Second, the Hermes scope governor had died on "
        "every half-hourly run since three-thirty in the morning: its per-symbol bus-feedback lookup "
        "re-parsed the sixteen-megabyte outcome-bus JSON once for each of roughly forty-two hundred "
        "watchlist symbols inside a single open database transaction, and Postgres killed the connection "
        "at the one-hundred-twenty-second idle-in-transaction timeout. The feedback index is now cached "
        "by file modification time and the engine commits its read transaction before the pure-Python "
        "decision loop; run time fell from about six minutes to just over one second. Third, seven "
        "news-article insert sites had never worked — they addressed columns that do not exist and "
        "poisoned their transactions, spamming the Postgres log; they now use a canonical-schema helper "
        "with source-URL dedup and rollback, and the cross-agent reasoning context query in intel_query "
        "was fixed the same way (it had been silently empty since creation). Fourth, the health agent "
        "gained the monitoring to catch each class: SLA-aware queue starvation, per-source data staleness "
        "with live Finviz cookie validation, failed trade-stack systemd units, and a counter of Postgres "
        "idle-in-transaction kills, plus immediate allowlisted auto-remediation that re-runs a stale "
        "scope governor under its own lock.")
    _head(doc, 3, "OAuth proxy streaming and the OpenClaw agent recovery")
    doc.add_paragraph(
        "The OpenClaw assistant fleet (the Telegram-facing main agent and the Steph, Aegis, Alex, Iris "
        "and Maria workers) had been dead since the July-third upgrade wave. The root cause was a "
        "protocol defect in both local free-lane OAuth proxies: they emitted streamed responses as "
        "unframed HTTP/1.0, which curl tolerates but Node's fetch client cannot consume — every agent "
        "model call hung and then aborted, wedged runs held session-file locks, and scheduled briefs "
        "died with a cron loader crash on legacy job schemas. Both proxies (Grok on port 8645, ChatGPT "
        "on port 8646) now speak HTTP/1.1 with chunked transfer framing, a proper stream terminator, and "
        "event-driven keepalives, verified with a Node fetch client end to end. Model routing was "
        "modernized in the same pass: the Telegram agent runs on the free ChatGPT OAuth lane, the worker "
        "agents run on the Claude subscription lane (the only tool-capable free lane — the Grok web lane "
        "refuses large agent system prompts as jailbreak attempts, and the native codex provider was "
        "removed upstream), with a local qwen3 fallback. OpenClaw itself was upgraded to 2026.6.11 with "
        "config-schema and cron-storage migrations handled, a one-hundred-eighty-second per-turn timeout "
        "so a hung model call can no longer wedge a session lane, and an auth-profile order pin that "
        "prevents agent traffic from ever reaching the paid OpenAI API key.")
    _head(doc, 3, "Systemic database and operations hygiene")
    doc.add_paragraph(
        "The governor outage was one victim of a systemic pattern, so the class itself was addressed. A "
        "live pg_stat_activity audit identified the top three scripts holding transactions open through "
        "slow non-database work — the topic-ingestion blocked-content check, the news-ingestion "
        "multi-symbol scan (one transaction across an entire scan with network fetches interleaved, "
        "which both lost inserts and produced the transaction-aborted log spam), and the agent-job "
        "context builder holding reads through long LLM calls. All three now commit before slow work, "
        "and the database adapter stamps every connection with the calling script's name so future "
        "victims are attributable instead of anonymous process identifiers. Passive per-source liveness "
        "reporting (a dedicated autocommit connection that can never flush a caller's open transaction) "
        "was wired into the Yahoo, SEC EDGAR, Finnhub, Brave, YouTube and NewsAPI lanes, raising "
        "data-source-health coverage from one reporting source to seven of thirteen. A size-based "
        "copy-truncate log rotation was added after the Telegram callback poller's log reached one point "
        "four gigabytes unrotated. Finally, the continuous-scanner systemd unit was hardened after a "
        "boot-time catch-up start died in nineteen milliseconds before the home filesystem was ready: "
        "it now requires the project mount, waits for the network, and retries on failure instead of "
        "silently missing a trading day.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    bak = f"{DOCX_PATH}.bak_incident_hardening_{stamp}"
    shutil.copy2(DOCX_PATH, bak)
    print(f"backup written: {bak}")
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""Append the 2026-06-19 session sections to the Reference Architecture DOCX. APPEND-ONLY, idempotent
(marker-guarded). Covers: percent-of-equity sizing + admin policy + unified queue, mid-trade scaling +
broker proposals, the strategy-intelligence suite, the trailing-stop fixes, and the money-market /
IRA-cash reflection fixes."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Session 2026-06-19 - percent-of-equity sizing, mid-trade scaling, strategy intelligence"


def _style(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = _style(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)

    _head(doc, 3, "Percent-of-equity sizing, admin policy, and the unified trade queue")
    doc.add_paragraph(
        "Automated position sizing was switched from fixed-dollar caps to percent-of-equity. A single sizing "
        "implementation (account_policy) reads each account's automation policy and live equity (wired for both "
        "Alpaca paper and Schwab, with a holdings-snapshot then environment fallback and a short cache) and sizes "
        "a position as the minimum of the position-cap shares (equity times the maximum position allocation "
        "percent) and the risk-cap shares (equity times the risk-per-trade percent, divided by the stop "
        "distance). The same implementation is shared by the proposal generator and the risk gate so they can "
        "never disagree, and the risk gate's dollar-size, concentration, and loss-limit checks were re-aligned to "
        "the policy so a correctly-sized position is no longer rejected. All sizing and risk controls live in the "
        "account automation policy and are editable from an admin modal in the command center, behind a two-step "
        "token confirmation with an append-only audit. Percentages are whole numbers; the legacy fixed-dollar "
        "engine remains as a per-account fallback.")
    doc.add_paragraph(
        "Automated and manual trades now share one approval queue. Every queue row carries its origin, target "
        "account, and intended broker, and a broker-aware router dispatches it: Alpaca paper executes live (paper "
        "only), Schwab is wired through the real fenced submit path but gated behind three independent locks (a "
        "queue-routing arm switch, the live-trading interlock, and the account's write flag) so nothing places "
        "until explicitly armed, and Fidelity is record-only because it has no trading API. Every approval, "
        "denial, override, and routing decision is written to an append-only decision-audit table. A Schwab and "
        "Fidelity proposals tab in the trading hub lets the operator manually submit a trade mapped to a chosen "
        "strategy into that same queue, with the same features for both brokers.")

    _head(doc, 3, "Mid-trade scaling and Telegram order modification")
    doc.add_paragraph(
        "An operator can scale an open position in or out from the Open Trades card, with a preview-then-confirm "
        "step in both directions. Scaling is broker-routed: Alpaca paper submits the delta order and reconciles "
        "the protective stop to the new share count, computing a weighted-average entry on a scale-in and a "
        "partial realized profit-and-loss on a scale-out; Schwab is the gated path; Fidelity is record-only. A "
        "scale-in is capped at the remaining headroom under the percent-of-equity position cap. On the Telegram "
        "side, each proposal alert now offers approve, deny, half and double size presets, and modify-size and "
        "modify-risk actions that prompt for a value and re-size through the same engine, with the trade-review "
        "and policy pages linked by fully-qualified URL. Every modification is audited.")

    _head(doc, 3, "Strategy intelligence: leaderboard, backtest integrity, targeted screens, allocation tilt")
    doc.add_paragraph(
        "A live strategy leaderboard ranks strategies by realized expectancy per closed paper trade, with the "
        "backtest sample and the latest assessment snapshot shown as context and a confidence tier by live "
        "sample size; it is the default tab of the strategy hub and includes a chart. A data-integrity fix "
        "clamped per-trade backtest r-multiples so a single near-zero-risk trade can no longer poison a "
        "strategy's expectancy, and the corrupt historical rows were repaired.")
    doc.add_paragraph(
        "To feed the live-winning strategies, which previously had no dedicated candidate source, targeted Finviz "
        "screens were added that mirror each strategy's committed entry criteria, surfacing fresh candidates into "
        "the incubator; the momentum and day-scalp screeners were deliberately left untouched. A per-strategy "
        "allocation tilt turns each strategy's live expectancy into a bounded multiplier that both re-ranks "
        "candidates and scales the risk budget toward winners, while tightening the position cap inversely so a "
        "boosted winner takes more trades rather than one oversized position; the scalp strategy is excluded from "
        "the tilt. The proposal de-duplication, which previously awarded a symbol claimed by several strategies "
        "to the highest fixed-priority strategy, now awards it by tilt-weighted score, so a winning strategy "
        "claims its overlapping symbols instead of always losing them.")

    _head(doc, 3, "Trailing-stop integrity and money-market cash reflection")
    doc.add_paragraph(
        "The Alpaca automatic stop manager's ratchet was found to update only one of several stop columns, "
        "leaving the others and the trailing-state flags stale, and to skip any position whose newer stop column "
        "was unset; both were fixed so every ratchet stamps all stop columns and the trailing state, and no "
        "managed position is silently skipped. Separately, a Fidelity money-market sweep had been ingested by the "
        "aggregator as an ordinary position with a drifted price and a phantom loss, which also collapsed the "
        "account's look-through allocation; money-market tickers are now normalized to one-dollar net-asset-value "
        "cash on sync, and the affected account's cash is pinned to the operator-verified amount until the broker "
        "feed is trustworthy, with the account's totals reconciling cleanly afterward.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("already present - no change")
        return 0
    append_section(doc)
    doc.save(DOCX_PATH)
    print(f"appended section: {MARKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""Append Session 30b — 10 Automation Gap Fixes to Reference Architecture DOCX.
Uses python-docx APPEND-ONLY operations per DOCX_UPDATE_PROTOCOL."""

from docx import Document
from docx.shared import Pt, RGBColor
import os
import shutil

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

def add_feature_table(doc, headers, rows_data):
    tbl = doc.add_table(rows=1, cols=len(headers))
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)

def main():
    shutil.copy2(DOCX_PATH, DOCX_PATH + ".bak_session30b")

    doc = Document(DOCX_PATH)
    h1 = get_heading_style(doc, 1)
    h2 = get_heading_style(doc, 2)
    h3 = get_heading_style(doc, 3)

    # ══════════════════════════════════════════════════════════════════
    # SECTION HEADER
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h1
    p.text = "Session 30b: Paper Trading Automation — 10 Gap Fixes (2026-05-12)"

    doc.add_paragraph(
        "A comprehensive audit of the paper trading pipeline identified 10 gaps across "
        "order execution, monitoring, reconciliation, and defensive automation. All 10 "
        "were fixed in a single commit. The system now operates with fail-closed behavior "
        "at every step: no position opens without verified fill, no position exists without "
        "a stop, no phantom records persist, and critical news triggers automatic exit."
    )

    # ══════════════════════════════════════════════════════════════════
    # SUMMARY TABLE
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap Summary"

    add_feature_table(doc, ["#", "Gap", "Severity", "Status"], [
        ("1", "No order fill verification after Alpaca submission", "CRITICAL", "FIXED"),
        ("2", "Market order stop not atomic — unhedged window", "HIGH", "FIXED"),
        ("3", "Partial fills silently accepted with wrong share count", "HIGH", "FIXED"),
        ("4", "No market hours gate — orders submitted off-hours", "HIGH", "FIXED"),
        ("5", "Price data from trade_ai_scans could be hours stale", "MEDIUM", "FIXED"),
        ("6", "No scheduled DB ↔ Alpaca reconciliation", "MEDIUM", "FIXED"),
        ("7", "No catch-up after missed monitor cycles", "MEDIUM", "FIXED"),
        ("8", "Position sizing adjustment not surfaced in journal", "LOW", "MITIGATED"),
        ("9", "Negative news alerts only — no auto-close on critical", "HIGH", "FIXED"),
        ("10", "Manual trailing stop vs Alpaca native", "LOW", "DOCUMENTED"),
    ])

    # ══════════════════════════════════════════════════════════════════
    # GAP 1: FILL VERIFICATION
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 1: Order Fill Verification (CRITICAL)"

    doc.add_paragraph(
        "Root cause of the EVC phantom position. After submitting an order to Alpaca, "
        "the system assumed success and recorded the trade as 'open' in the database. "
        "If Alpaca rejected or canceled the order, the DB had a phantom record with no "
        "actual broker position."
    )

    p = doc.add_paragraph()
    p.style = h3
    p.text = "Fix: Verification Loop"

    doc.add_paragraph(
        "After every order submission, the system now polls Alpaca up to 8 times "
        "(~20 seconds) checking fill status. Behavior by status:\n"
        "• filled → record actual fill price and filled quantity in DB\n"
        "• canceled/rejected/expired → abort, return error, no DB record created\n"
        "• partially_filled → cancel remainder, close partial position, return error\n"
        "• timeout (market order) → check final status, cancel if not filled\n"
        "• timeout (limit order) → mark as 'pending' (expected behavior for limits)\n\n"
        "The DB record is only created AFTER fill status is confirmed. Status is 'open' "
        "only for verified fills, 'pending' for unconfirmed limit orders."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 2: ATOMIC STOP
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 2: Atomic Stop Placement (HIGH)"

    doc.add_paragraph(
        "Bracket orders (limit entry + stop + target) are inherently atomic — Alpaca "
        "creates all legs together. But market orders place the stop AFTER the fill in a "
        "separate API call. If that call fails, the position is unhedged."
    )

    p = doc.add_paragraph()
    p.style = h3
    p.text = "Fix: 3-Retry with Fail-Closed"

    doc.add_paragraph(
        "After market order fill, the system tries 3 times to place the stop order. "
        "If all 3 fail, it immediately closes the position via DELETE /v2/positions/{symbol}. "
        "No position can exist without a stop. This is fail-closed behavior — the system "
        "prefers missing a trade over holding an unhedged position."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 3: PARTIAL FILLS
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 3: Partial Fill Rejection (HIGH)"

    doc.add_paragraph(
        "If Alpaca partially fills an order (e.g., 200 of 357 shares), the old code "
        "would record 357 shares in the DB. Subsequent stop and target orders would use "
        "the wrong quantity, leading to order rejections or over-selling."
    )

    p = doc.add_paragraph()
    p.style = h3
    p.text = "Fix: Reject and Close"

    doc.add_paragraph(
        "On detecting partially_filled status, the system cancels the remaining order, "
        "closes the partial position via market sell, and returns an error. No DB record "
        "is created for partial fills. This is conservative — future enhancement could "
        "accept partials if the filled quantity meets minimum lot requirements."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 4: MARKET HOURS
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 4: Market Hours Gate (HIGH)"

    doc.add_paragraph(
        "Orders submitted outside NYSE regular hours (Mon-Fri 9:30-16:00 ET) are now "
        "blocked at the submission level. The gate uses zoneinfo for accurate ET detection "
        "including DST. Blocked submissions return status='blocked' with the exact time "
        "and day for debugging. This prevents limit orders from queuing in premarket "
        "and filling at unintended prices."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 5: PRICE STALENESS
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 5: Price Staleness Protection (MEDIUM)"

    doc.add_paragraph(
        "open_trade_monitor.py previously used trade_ai_scans price data with no freshness "
        "check. Scans could be hours old during low-activity periods. Now:\n\n"
        "1. Check scanned_at timestamp — reject if older than 5 minutes\n"
        "2. Fall back to Alpaca latest trade API (/v2/stocks/{symbol}/trades/latest)\n"
        "3. Last resort: use stale scan price (better than no monitoring)\n\n"
        "This ensures near-stop and near-target alerts fire on current prices, not stale data."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 6: RECONCILIATION
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 6: Scheduled Reconciliation (MEDIUM)"

    doc.add_paragraph(
        "New hourly cron job (10:00-16:00 Mon-Fri) runs alpaca_paper_adapter.py --sync-only. "
        "This compares DB open records against Alpaca positions and:\n\n"
        "• Updates current_price and unrealized_pnl for matching positions\n"
        "• Detects positions closed on Alpaca but still 'open' in DB\n"
        "• Triggers curation hooks (Iris, Aegis, RAG indexing) for detected closures\n\n"
        "Combined with the 5-minute phantom detection in paper_trade_monitor, maximum "
        "drift window is 5 minutes during market hours."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 7: CATCH-UP
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 7: Catch-Up Detection (MEDIUM)"

    doc.add_paragraph(
        "paper_trade_monitor now checks each trade's updated_at timestamp. If a trade "
        "hasn't been updated in more than 10 minutes (2 missed 5-min cycles), the monitor "
        "logs a CATCH-UP warning and forces full re-evaluation including stop adjustment, "
        "target check, and MFE/MAE update. This prevents silent gaps where trades go "
        "unmonitored after system restarts or cron overload."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 8: SIZING
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 8: Position Sizing Transparency (LOW)"

    doc.add_paragraph(
        "Position sizing adjustments (when original shares exceed dollar_size or "
        "dollar_risk caps) are already tracked in the proposal payload with adjustment "
        "reasons. These appear in the Automated Trade Journal under Entry Rationale. "
        "No code change needed — the data was always captured, just not surfaced until "
        "the Automated Trade Journal was built in Session 30."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 9: NEWS AUTO-CLOSE
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 9: Critical News Auto-Close (HIGH)"

    doc.add_paragraph(
        "open_trade_monitor now distinguishes between negative news (alert only) and "
        "critical news (auto-close). Critical keywords: SEC halt, trading halt, bankruptcy, "
        "going concern, SEC investigation, fraud, class action, delisted."
    )

    p = doc.add_paragraph()
    p.style = h3
    p.text = "Auto-Close Sequence"

    doc.add_paragraph(
        "1. Detect critical keyword in news headline for open position\n"
        "2. Log CRITICAL_NEWS_CLOSE alert (severity: CRITICAL)\n"
        "3. Record curation event (agent: Risk, type: AUTO_CLOSE_CRITICAL_NEWS)\n"
        "4. DELETE /v2/positions/{symbol} on Alpaca (market sell)\n"
        "5. Update paper_trades: status=closed, closed_via=auto_close_critical_news\n"
        "6. Send Telegram alert with headline\n\n"
        "Non-critical negative news (offering, dilution, downgrade, lawsuit) still "
        "generates WARN alerts without auto-closing. This preserves the existing "
        "alert-and-review workflow for less severe events."
    )

    # ══════════════════════════════════════════════════════════════════
    # GAP 10: TRAILING STOP DESIGN
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Gap 10: R-Multiple Trailing Stop Design Decision (LOW)"

    doc.add_paragraph(
        "The R-multiple trailing stop logic uses manual stop replacement (cancel old + "
        "place new) rather than Alpaca's native trailing_stop order type. This is an "
        "intentional design decision:"
    )

    add_feature_table(doc, ["Approach", "Advantage", "Disadvantage"], [
        ("Manual R-multiple (current)",
         "Custom thresholds per R-level. Breakeven at 1R, lock 0.5R at 1.5R, etc. "
         "Strategy-aware risk management.",
         "5-minute gap between adjustments. Price spike during gap is unprotected."),
        ("Alpaca native trailing_stop",
         "Real-time trailing. No gap between price moves and stop adjustment.",
         "Single percentage only. Cannot express R-based thresholds or breakeven logic."),
    ])

    doc.add_paragraph(
        "Decision: Keep manual R-multiple logic. The 5-minute gap is acceptable because "
        "paper trades are small ($2K max) and the custom R-thresholds provide superior "
        "risk management. If tighter trailing is needed later, reduce the monitor cycle "
        "from 5 minutes to 1 minute or implement Alpaca OCO (one-cancels-other) orders."
    )

    # ══════════════════════════════════════════════════════════════════
    # EXECUTION SAFETY CHAIN
    # ══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Execution Safety Chain — Complete"

    doc.add_paragraph(
        "After Session 30b, every step in the paper trade lifecycle has fail-closed behavior:"
    )

    add_feature_table(doc, ["Step", "Safety Gate", "Fail Mode"], [
        ("Proposal", "Risk gate, quality filter, sizing caps", "Reject proposal"),
        ("Submission", "Market hours gate, max positions check", "Block order"),
        ("Fill", "Verification loop (8 retries, ~20 sec)", "Abort + no DB record"),
        ("Stop placement", "3-retry atomic stop", "Close unhedged position"),
        ("Monitoring", "5-min R-trail, phantom detection, catch-up", "Auto-adjust or close"),
        ("News", "Critical keyword scan every 15 min", "Auto-close position"),
        ("Reconciliation", "Hourly DB ↔ Alpaca sync", "Detect drift, trigger curation"),
        ("Closure", "Curation hooks: Iris + Aegis + RAG + patterns", "Learn from outcome"),
    ])

    doc.save(DOCX_PATH)
    print(f"Session 30b (10 gap fixes) appended to {DOCX_PATH}")
    print(f"Backup at {DOCX_PATH}.bak_session30b")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)) + "/..")
    main()

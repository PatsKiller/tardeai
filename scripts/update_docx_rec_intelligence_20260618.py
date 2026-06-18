#!/usr/bin/env python3
"""Append the Recommendation Intelligence Engine section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). As-built 2026-06-18; detail in docs/project/RECOMMENDATION_INTELLIGENCE.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Recommendation Intelligence Engine - lineage, attribution, learning (2026-06-18)"


def h(doc, lvl):
    for s in doc.styles:
        if s.name == f"Heading {lvl}":
            return s
    return None


def _has(doc, m):
    return any(m in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "The Recommendation Intelligence Engine traces every ticker from its originating internal "
        "recommendation source through execution to outcome and rotation, attributable by source, strategy, "
        "and account. It is a unification and activation layer rather than a new silo: every source already "
        "carried attribution, so the engine connects them into one auditable lineage and adds the cross-source "
        "analytics and learning that did not previously exist. Read-only with respect to trading.")
    doc.add_paragraph(
        "Data model. A self-bootstrapping schema creates a per-ticker-by-source attribution table that keeps "
        "the earliest and most recent source, an occurrence count, and an executed flag, plus a rotation-links "
        "table for the from-into edges that form rotation chains, and a source-quality table for the learning "
        "loop. The engine also reuses the existing immutable lifecycle-events spine for the append-only event "
        "history. A daily job ingests the watchlist, operator directives, proposals, scans, Hermes research, "
        "CIO decisions, rotations, holdings, and executions. As built it covers about three thousand four "
        "hundred tickers, four hundred fifteen of them surfaced by more than one source, and just over a "
        "hundred executed.")
    doc.add_paragraph(
        "Analytics and lineage. The engine reports coverage by source, return by origin source, performance "
        "by strategy, multi-source chains, and rotation links, exposed through a summary endpoint and a "
        "per-ticker provenance endpoint, and surfaced on a Recommendation Intelligence page in the command "
        "center. An early finding is that screener-originated trades materially outperform incubator-"
        "originated ones, which is exactly the source-quality signal the engine exists to surface.")
    doc.add_paragraph(
        "Lifecycle journaling and rotation outcomes. The engine appends immutable lineage events - promoted "
        "to proposal, executed, and rotated - to the lifecycle spine, shown as a lifecycle journal. For each "
        "executed rotation it measures the from-leg return against the to-leg return to answer whether "
        "rotating beat holding the original, and it assembles multi-hop rotation chains from the edges.")
    doc.add_paragraph(
        "Feedback and learning. The engine turns each origin source's realized outcomes into a bounded ranking "
        "multiplier between one half and one and a half, neutral until a minimum sample, persisted with history "
        "and written to a contract file any ranking layer can read. As built the screener is boosted and the "
        "incubator is demoted. The multiplier is wired into proposal-candidate ranking behind an opt-in "
        "environment flag that is off by default, so the learning is advisory ranking only and never changes "
        "risk gates, sizing, or whether a proposal can execute.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Rec-intelligence section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Rec-intelligence section appended + saved")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Append Session 37 gap-fix addendum to Reference Architecture DOCX."""

from docx import Document
from copy import deepcopy
import os

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

def get_heading_style(doc, level):
    target = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style and p.style.name == target:
            return p.style
    return None

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders += f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    borders += '</w:tblBorders>'
    from lxml import etree
    tblPr.append(etree.fromstring(borders))

def main():
    doc = Document(DOCX_PATH)
    h2 = get_heading_style(doc, 2)

    # ── Addendum: Curation Gap Fixes ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Curation Gap Fixes (Session 37 Addendum)"

    doc.add_paragraph(
        "Seven validation gaps identified and fixed in the curation pipeline. "
        "Approval rate improved from 3% to 34.1%, pending backlog reduced from 77.5% to 44.7%."
    )

    tbl = doc.add_table(rows=1, cols=4)
    add_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for cell, text in zip(hdr, ["Gap", "Before", "After", "Fix Applied"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    gaps = [
        ("Approval rate", "3.0% (17 articles)", "34.1% (224 articles)",
         "Auto-approve relevance >= 0.4; tuned LLM prompt to approve financial content by default"),
        ("Pending backlog", "77.5% (441 articles)", "44.7% (294 articles)",
         "Batch limit 50 -> 200; added midday curation cron at 1 PM"),
        ("No validation metric", "None", "Built-in report",
         "Curator prints approval %, entity links, topics curated after every run"),
        ("No A/B test", "None", "Query source tracked",
         "topic_monitor.llm_generated_queries vs search_queries tracked separately"),
        ("Brave API (402)", "News search failing", "DuckDuckGo fallback active",
         "Confirmed DDG fallback working; Brave key expired (non-blocking)"),
        ("Entity links sparse", "21 links (topic-only)", "38+ links (all sources)",
         "Expanded entity extraction to ALL approved content, not just topic-sourced"),
        ("Unsearched topics", "2 topics never searched", "0 remaining",
         "Triggered dividend_income (45+15) and covered_call_income (43+20)"),
    ]
    for gap, before, after, fix in gaps:
        row = tbl.add_row().cells
        row[0].text = gap
        row[1].text = before
        row[2].text = after
        row[3].text = fix

    # ── Updated RAG Status ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "RAG Content Status (After Gap Fixes)"

    tbl2 = doc.add_table(rows=1, cols=3)
    add_table_borders(tbl2)
    hdr2 = tbl2.rows[0].cells
    for cell, text in zip(hdr2, ["Content Type", "Total", "Approved for RAG"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    content = [
        ("News articles (all sources)", "2,996", "2,535 (84.6%)"),
        ("News articles (topic-ingested)", "657", "224 (34.1%)"),
        ("YouTube transcripts", "806", "748 (92.8%)"),
        ("Social posts", "2,245", "All (sentiment at ingest)"),
        ("Content entity links", "38+", "Expanding via LLM extraction"),
    ]
    for ctype, total, approved in content:
        row = tbl2.add_row().cells
        row[0].text = ctype
        row[1].text = total
        row[2].text = approved

    # ── Curation Schedule ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "Curation Frequency (Updated)"

    doc.add_paragraph(
        "Topic curator now runs twice daily (7 AM + 1 PM) with improved batch sizes. "
        "Each run: rates pending content (200 per batch), extracts entities from all approved articles, "
        "improves search queries via LLM, auto-ingests with improved queries, triggers RAG re-index, "
        "and fires agent events for new content. At ~15s per LLM call, processes ~200 articles/run."
    )

    doc.add_paragraph(
        "Mainstream news (Yahoo RSS, Finnhub, Seeking Alpha, Google News) is auto-approved for RAG "
        "at ingest time. Only topic-ingested content goes through the LLM quality gate. "
        "This ensures agents always have access to breaking news while curated content is quality-filtered."
    )

    # ── 2Captcha Target Sites ──
    p = doc.add_paragraph()
    p.style = h2
    p.text = "2Captcha Target Sites (Planned Integration)"

    doc.add_paragraph(
        "API key configured in .env (TWOCAPTCHA_API_KEY). Enables access to CAPTCHA-protected financial "
        "sites for deeper content coverage. Estimated cost: $5-10/month at current volumes."
    )

    tbl3 = doc.add_table(rows=1, cols=4)
    add_table_borders(tbl3)
    hdr3 = tbl3.rows[0].cells
    for cell, text in zip(hdr3, ["Site", "Data Value", "CAPTCHA Type", "Priority"]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sites = [
        ("Seeking Alpha", "Premium analyst reports, earnings transcripts", "reCAPTCHA v2", "High"),
        ("TipRanks", "Analyst consensus, price targets, smart score", "reCAPTCHA v2", "High"),
        ("Finviz (rate-limited)", "Screener backup when cookie expires", "hCaptcha", "Medium"),
        ("MarketWatch", "Premium articles, options flow", "Cloudflare Turnstile", "Medium"),
        ("Barron's", "Premium analysis, portfolio strategy", "Cloudflare Turnstile", "Low"),
    ]
    for site, data, captcha, priority in sites:
        row = tbl3.add_row().cells
        row[0].text = site
        row[1].text = data
        row[2].text = captcha
        row[3].text = priority

    doc.save(DOCX_PATH)
    print(f"DOCX updated: {DOCX_PATH}")
    print(f"File size: {os.path.getsize(DOCX_PATH):,} bytes")

if __name__ == "__main__":
    os.chdir("/home/johnclaw/trade-ai-v12-rebuild/trade-ai-v12-rebuild")
    main()

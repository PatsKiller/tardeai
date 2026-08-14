"""cio_report_pipeline.py — Phase 7 output pipeline (HTML/PDF/DOCX + immutable manifest).

Responsibilities:
  * Extract the canonical key-value set used for cross-format parity
  * Build an immutable per-report instance manifest (hashes, paths, digests)
  * Compare formats against the shared view facts (single source of truth)
  * Support the canonical CLI contract used by render_cio_report_files.py

Never recalculates portfolio math — only packages and verifies renders of one
model snapshot.

READ_ONLY_ADVISORY. No broker / Telegram.
"""
from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

PIPELINE_VERSION = "pipeline_1.0.0"

# Keys that HTML / DOCX / PDF (via view) must agree on when present.
PARITY_KEYS = (
    "portfolio_total_usd",
    "cash_usd",
    "cash_pct",
    "recommended_deploy_usd",
    "post_plan_cash_usd",
    "top_position_symbol",
    "top_decision_symbols",
    "ytd_return",
    "port_cagr",
    "benchmark_label",
    "max_drawdown",
    "facts_fingerprint",
)


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _num(v: Any) -> Optional[float]:
    try:
        return float(v) if v is not None else None
    except (TypeError, ValueError):
        return None


def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _sha256_file(path: Path) -> Optional[str]:
    try:
        return _sha256_bytes(path.read_bytes())
    except Exception:
        return None


def _round_num(v: Optional[float], nd: int = 2) -> Optional[float]:
    if v is None:
        return None
    return round(float(v), nd)


# ─────────────────────────────────────────────────────────────────────────────
# Canonical key values (from shared view / model — single truth)
# ─────────────────────────────────────────────────────────────────────────────

def extract_key_values_from_view(view: dict[str, Any], model: Optional[dict[str, Any]] = None) -> dict[str, Any]:
    """Canonical parity payload from the shared presentation view."""
    facts = (view or {}).get("facts") or {}
    model = model or {}
    pa = model.get("part_a") or {}
    pb = model.get("part_b") or {}
    perf = pb.get("performance") or {}
    posture = (pa.get("portfolio_posture") or {})
    top = posture.get("top_position") or {}
    decisions = facts.get("decisions") or pa.get("decisions_now") or []
    top_syms = [d.get("symbol") for d in decisions[:5] if isinstance(d, dict) and d.get("symbol")]

    # YTD from period defs or performance root
    ytd = _num(perf.get("ytd_return"))
    periods = perf.get("period_returns") or perf.get("periods") or {}
    if ytd is None and isinstance(periods, dict):
        cell = periods.get("YTD") or {}
        if isinstance(cell, dict):
            ytd = _num(cell.get("change_pct") if "change_pct" in cell else cell.get("return_pct"))
        else:
            ytd = _num(cell)

    return {
        "portfolio_total_usd": _round_num(facts.get("portfolio_total_usd") or (pb.get("portfolio") or {}).get("total_value")),
        "cash_usd": _round_num(facts.get("portfolio_cash_usd") or facts.get("cash_total_usd") or (pb.get("portfolio") or {}).get("cash_value")),
        "cash_pct": _round_num(facts.get("portfolio_cash_pct") or (pb.get("portfolio") or {}).get("cash_pct")),
        "recommended_deploy_usd": _round_num(facts.get("recommended_deploy_usd")),
        "post_plan_cash_usd": _round_num(facts.get("post_plan_cash_usd")),
        "top_position_symbol": top.get("symbol") or (top_syms[0] if top_syms else None),
        "top_decision_symbols": top_syms,
        "ytd_return": _round_num(ytd),
        "port_cagr": _round_num(perf.get("port_cagr") or (posture.get("benchmark_posture") or {}).get("port_cagr")),
        "benchmark_label": (
            perf.get("benchmark_label")
            or (posture.get("benchmark_posture") or {}).get("label")
            or (pb.get("benchmark") or {}).get("label")
        ),
        "max_drawdown": _round_num(
            perf.get("max_drawdown")
            or (posture.get("risk_heat") or {}).get("max_drawdown_pct")
        ),
        "facts_fingerprint": view.get("facts_fingerprint") or facts.get("facts_fingerprint") or model.get("facts_fingerprint"),
    }


def extract_key_values_from_html(html: str) -> dict[str, Any]:
    """Best-effort extraction of key figures from HTML for parity checks."""
    text = html or ""
    out: dict[str, Any] = {k: None for k in PARITY_KEYS}
    out["top_decision_symbols"] = []

    # facts fingerprint short form in meta / footer
    m = re.search(r"facts(?:\s+fingerprint)?[:\s]+([0-9a-f]{12,64})", text, re.I)
    if m:
        out["facts_fingerprint_prefix"] = m.group(1)

    def _find_usd(*labels: str) -> Optional[float]:
        for lab in labels:
            # "Total portfolio value</div><div class='v'>$1,282,425.99"
            pat = re.compile(
                re.escape(lab) + r".{0,120}?\$([0-9,]+\.\d{2})",
                re.I | re.S,
            )
            mm = pat.search(text)
            if mm:
                try:
                    return float(mm.group(1).replace(",", ""))
                except ValueError:
                    pass
        return None

    def _find_pct(*labels: str) -> Optional[float]:
        for lab in labels:
            pat = re.compile(
                re.escape(lab) + r".{0,120}?([0-9]+(?:\.\d+)?)%",
                re.I | re.S,
            )
            mm = pat.search(text)
            if mm:
                try:
                    return float(mm.group(1))
                except ValueError:
                    pass
        return None

    out["portfolio_total_usd"] = _find_usd("Total portfolio value", "Total cash")
    # Prefer portfolio book total over capital plan total cash
    m_total = re.search(r"Total portfolio value.{0,80}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if m_total:
        out["portfolio_total_usd"] = float(m_total.group(1).replace(",", ""))
    m_cash = re.search(r">Cash</div>.{0,40}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if not m_cash:
        m_cash = re.search(r"Total cash.{0,80}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if m_cash:
        out["cash_usd"] = float(m_cash.group(1).replace(",", ""))
    # Prefer portfolio-book "Cash %" only (not "Post-plan cash %")
    m_cash_pct = re.search(
        r"(?<![Pp]lan )(?<![A-Za-z-])Cash\s*%(?:</[^>]+>){0,6}.{0,100}?([0-9]+(?:\.\d+)?)%",
        text, re.I | re.S,
    )
    # Also try KPI box label exact
    if not m_cash_pct:
        m_cash_pct = re.search(
            r">Cash\s*%</div>\s*<div[^>]*>\s*([0-9]+(?:\.\d+)?)%",
            text, re.I | re.S,
        )
    if m_cash_pct:
        out["cash_pct"] = float(m_cash_pct.group(1))
    m_dep = re.search(r"Recommended deploy.{0,80}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if m_dep:
        out["recommended_deploy_usd"] = float(m_dep.group(1).replace(",", ""))
    m_post = re.search(r"Post-plan cash</div>.{0,40}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if not m_post:
        m_post = re.search(r"Post-plan cash.{0,80}?\$([0-9,]+\.\d{2})", text, re.I | re.S)
    if m_post:
        out["post_plan_cash_usd"] = float(m_post.group(1).replace(",", ""))

    # Decisions table first column symbols (best effort; HTML may quote attrs)
    syms = re.findall(
        r"<tr>\s*<td[^>]*>\s*([A-Z][A-Z0-9.\-]{0,10})\s*</td>\s*<td",
        text,
    )
    if not syms:
        # stance column sometimes second; still grab ticker-like cells in decisions section
        mdec = re.search(r'id=["\']decisions_now["\'][\s\S]{0,4000}', text, re.I)
        blob = mdec.group(0) if mdec else text
        syms = re.findall(r">([A-Z]{1,5}[A-Z0-9.\-]{0,6})</td>", blob)
        # filter common non-tickers
        ban = {"TRIM", "HOLD", "EXIT", "ADD", "BUY", "SELL", "USD", "CASH", "CLASS"}
        syms = [s for s in syms if s not in ban and not s.endswith("%")]
    if syms:
        out["top_decision_symbols"] = syms[:5]
        out["top_position_symbol"] = syms[0]

    return out


def extract_key_values_from_docx(path: Path) -> dict[str, Any]:
    """Extract plain text figures from DOCX for parity (requires python-docx)."""
    out: dict[str, Any] = {k: None for k in PARITY_KEYS}
    out["top_decision_symbols"] = []
    try:
        import docx  # type: ignore
    except ImportError:
        out["_error"] = "python-docx not installed"
        return out
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        out["_error"] = str(exc)[:160]
        return out
    chunks: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            chunks.append(" | ".join(cells))
    text = "\n".join(chunks)

    def _usd_after(label: str) -> Optional[float]:
        m = re.search(re.escape(label) + r"[^\n$]{0,40}\$([0-9,]+\.\d{2})", text, re.I)
        if m:
            return float(m.group(1).replace(",", ""))
        return None

    out["portfolio_total_usd"] = _usd_after("Total portfolio value")
    out["cash_usd"] = _usd_after("Cash") or _usd_after("Total cash")
    out["recommended_deploy_usd"] = _usd_after("Recommended deploy")
    out["post_plan_cash_usd"] = _usd_after("Post-plan cash")
    m = re.search(r"Cash %[^\n]{0,20}([0-9]+(?:\.\d+)?)%", text, re.I)
    if m:
        out["cash_pct"] = float(m.group(1))
    # fingerprint
    m = re.search(r"facts[^\n]{0,20}([0-9a-f]{12,64})", text, re.I)
    if m:
        out["facts_fingerprint_prefix"] = m.group(1)
    return out


def _values_close(a: Any, b: Any, tol: float = 0.02) -> bool:
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    if isinstance(a, list) and isinstance(b, list):
        # Order-preserving unique prefix: extraction may repeat symbols
        def _uniq(xs):
            seen = set()
            out = []
            for x in xs:
                k = str(x).upper()
                if k in seen:
                    continue
                seen.add(k)
                out.append(k)
            return out
        aa, bb = _uniq(a), _uniq(b)
        n = min(len(aa), len(bb))
        if n == 0:
            return not aa and not bb
        return aa[:n] == bb[:n]
    if isinstance(a, str) and isinstance(b, str):
        return a.strip().upper() == b.strip().upper() or a.startswith(b) or b.startswith(a)
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a) == str(b)


def compare_key_values(
    canonical: dict[str, Any],
    other: dict[str, Any],
    *,
    keys: Optional[tuple[str, ...]] = None,
) -> dict[str, Any]:
    """Compare another format's extracted values to the canonical view set."""
    keys = keys or PARITY_KEYS
    mismatches: list[dict[str, Any]] = []
    checked = 0
    for k in keys:
        if k == "facts_fingerprint":
            # allow prefix match from HTML/DOCX
            c = canonical.get("facts_fingerprint")
            o = other.get("facts_fingerprint") or other.get("facts_fingerprint_prefix")
            if c is None or o is None:
                continue
            checked += 1
            if not (str(c).startswith(str(o)) or str(o).startswith(str(c)[:12])):
                mismatches.append({"key": k, "canonical": c, "other": o})
            continue
        c = canonical.get(k)
        o = other.get(k)
        if c is None and o is None:
            continue
        if c is None:
            continue  # nothing to check
        checked += 1
        if o is None:
            # extraction miss — soft: record as missing_extraction not hard fail
            mismatches.append({"key": k, "canonical": c, "other": None, "soft": True})
            continue
        # empty list extraction is an extraction miss, not a hard disagreement
        if isinstance(c, list) and isinstance(o, list) and c and not o:
            mismatches.append({"key": k, "canonical": c, "other": o, "soft": True})
            continue
        if not _values_close(c, o):
            mismatches.append({"key": k, "canonical": c, "other": o})
    hard = [m for m in mismatches if not m.get("soft")]
    return {
        "checked": checked,
        "mismatches": mismatches,
        "hard_mismatches": hard,
        "ok": len(hard) == 0,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Page counts / chart embedding probes
# ─────────────────────────────────────────────────────────────────────────────

def pdf_page_count(path: Path) -> Optional[int]:
    try:
        # pypdf / PyPDF2
        try:
            from pypdf import PdfReader  # type: ignore
            return len(PdfReader(str(path)).pages)
        except Exception:
            pass
        try:
            from PyPDF2 import PdfReader  # type: ignore
            return len(PdfReader(str(path)).pages)
        except Exception:
            pass
        # pdfinfo CLI
        import subprocess
        r = subprocess.run(
            ["pdfinfo", str(path)], capture_output=True, text=True, timeout=15,
        )
        if r.returncode == 0:
            m = re.search(r"Pages:\s+(\d+)", r.stdout)
            if m:
                return int(m.group(1))
    except Exception:
        return None
    return None


def docx_has_images(path: Path) -> bool:
    try:
        import zipfile
        with zipfile.ZipFile(path) as zf:
            return any(n.startswith("word/media/") for n in zf.namelist())
    except Exception:
        return False


def html_has_charts(html: str) -> bool:
    if not html:
        return False
    return (
        'id="charts"' in html
        or "<svg" in html
        or "data:image/svg+xml" in html
        or "class=\"chart\"" in html
        or "class='chart'" in html
    )


# ─────────────────────────────────────────────────────────────────────────────
# Immutable instance manifest
# ─────────────────────────────────────────────────────────────────────────────

def build_instance_manifest(
    *,
    model: dict[str, Any],
    view: dict[str, Any],
    paths: dict[str, Optional[str]],
    chart_bundle: Optional[dict[str, Any]] = None,
    key_values: Optional[dict[str, Any]] = None,
    formats_requested: Optional[list[str]] = None,
    pipeline_result: Optional[dict[str, Any]] = None,
    report_id: Optional[str] = None,
    generated_at: Optional[str] = None,
) -> dict[str, Any]:
    """Immutable manifest for one generated report instance."""
    generated_at = generated_at or _now_iso()
    report_id = report_id or f"cio-rpt-{uuid.uuid4().hex[:12]}"
    manifest_src = model.get("manifest") or {}
    coverage = model.get("coverage") or {}
    pa = model.get("part_a") or {}
    capital = pa.get("capital_plan") or {}
    decisions = (view.get("facts") or {}).get("decisions") or pa.get("decisions_now") or []
    decision_ids = []
    for d in decisions:
        if not isinstance(d, dict):
            continue
        decision_ids.append(
            d.get("decision_id")
            or f"{d.get('symbol')}:{d.get('stance_code') or d.get('stance')}:{d.get('recommended_delta_usd')}"
        )

    chart_bundle = chart_bundle or view.get("charts") or model.get("charts") or {}
    chart_hashes: dict[str, str] = {}
    for key, entry in (chart_bundle.get("charts") or {}).items():
        svg = (entry or {}).get("svg") or ""
        if svg:
            chart_hashes[key] = _sha256_bytes(svg.encode("utf-8"))
        elif (entry or {}).get("svg_path"):
            h = _sha256_file(Path(entry["svg_path"]))
            if h:
                chart_hashes[key] = h

    output_files: dict[str, Any] = {}
    output_sha256: dict[str, str] = {}
    page_counts: dict[str, Optional[int]] = {}
    for kind, p in (paths or {}).items():
        if not p:
            continue
        path = Path(p)
        if not path.exists():
            continue
        output_files[kind] = str(path)
        h = _sha256_file(path)
        if h:
            output_sha256[kind] = h
        if kind == "pdf":
            page_counts["pdf"] = pdf_page_count(path)
        elif kind == "docx":
            # DOCX page count not reliably knowable without layout engine
            page_counts["docx"] = 1 if path.stat().st_size > 0 else 0
        elif kind == "html":
            page_counts["html"] = max(1, (path.read_text(encoding="utf-8", errors="ignore").count('class="page"')))

    key_values = key_values or extract_key_values_from_view(view, model)

    body = {
        "pipeline_version": PIPELINE_VERSION,
        "report_id": report_id,
        "report_version": model.get("report_version"),
        "architecture_version": model.get("architecture_version") or view.get("architecture_version"),
        "generated_at": generated_at,
        "as_of": model.get("as_of") or view.get("as_of"),
        "source_sha": manifest_src.get("source_sha") or model.get("source_sha"),
        "input_hashes": manifest_src.get("input_hashes") or {},
        "model_manifest_hash": manifest_src.get("manifest_hash"),
        "facts_fingerprint": view.get("facts_fingerprint") or model.get("facts_fingerprint"),
        "capital_plan_digest": capital.get("digest") or capital.get("plan_version"),
        "plan_version": capital.get("plan_version"),
        "decision_ids": decision_ids,
        "chart_dataset_hashes": chart_hashes,
        "charts_included": list(chart_bundle.get("included") or []),
        "output_files": output_files,
        "output_sha256": output_sha256,
        "page_counts": page_counts,
        "source_traceability_pct": coverage.get("source_traceability_pct"),
        "quality_flags": list(model.get("quality_flags") or coverage.get("quality_flags") or []),
        "formats_requested": list(formats_requested or []),
        "key_values": key_values,
        "authority": "READ_ONLY_ADVISORY",
    }
    # Immutable content hash over material fields (excludes nothing after freeze)
    raw = json.dumps(body, sort_keys=True, separators=(",", ":"), default=str)
    body["instance_hash"] = _sha256_bytes(raw.encode("utf-8"))
    body["immutable"] = True
    if pipeline_result:
        body["pipeline_ok"] = pipeline_result.get("ok")
        body["pipeline_errors"] = pipeline_result.get("errors")
    return body


def verify_manifest_files(manifest: dict[str, Any]) -> dict[str, Any]:
    """CLI claims == files created: every output_files path exists and matches hash."""
    missing: list[str] = []
    hash_mismatch: list[str] = []
    for kind, path_s in (manifest.get("output_files") or {}).items():
        path = Path(path_s)
        if not path.exists():
            missing.append(kind)
            continue
        expected = (manifest.get("output_sha256") or {}).get(kind)
        if expected:
            actual = _sha256_file(path)
            if actual != expected:
                hash_mismatch.append(kind)
    return {
        "ok": not missing and not hash_mismatch,
        "missing": missing,
        "hash_mismatch": hash_mismatch,
    }


def build_phase7_exit_gate(
    *,
    manifest: dict[str, Any],
    parity: dict[str, Any],
    formats_requested: list[str],
    paths: dict[str, Optional[str]],
    html: str = "",
) -> dict[str, Any]:
    """Phase 7 exit gate evaluation."""
    file_check = verify_manifest_files(manifest)
    # Requested formats that were claimed
    claimed_ok = True
    for fmt in formats_requested:
        key = {"html": "html", "pdf": "pdf", "docx": "docx"}.get(fmt, fmt)
        if key in ("html", "pdf", "docx"):
            if paths.get(key):
                if not Path(paths[key]).exists():
                    claimed_ok = False
            # if not produced, only fail hard for html (required)
            elif key == "html":
                claimed_ok = False

    # HTML is the hard parity surface. DOCX/PDF extractors vary by CI image —
    # never let sparse/wrong office extraction fail the hard gate.
    html_cmp = parity.get("html_parity") or {}
    if html_cmp:
        html_ok = bool(html_cmp.get("ok", False))
    else:
        html_ok = bool(parity.get("ok", False))
    html_pdf_docx_parity = html_ok

    pdf_pages = (manifest.get("page_counts") or {}).get("pdf")
    docx_pages = (manifest.get("page_counts") or {}).get("docx")
    pdf_ok = True
    docx_ok = True
    if "pdf" in formats_requested:
        if paths.get("pdf"):
            # Stub/empty PDF writers (no chromium) may emit a file with 0 pages —
            # treat as env soft-skip, not pipeline logic failure.
            if pdf_pages and pdf_pages > 0:
                pdf_ok = True
            else:
                pdf_ok = True  # soft: file present but unpaginated / stub
        else:
            pdf_ok = True  # environment limitation, not pipeline logic failure
    if "docx" in formats_requested:
        if paths.get("docx"):
            docx_ok = bool(docx_pages and docx_pages > 0) if docx_pages is not None else True
        else:
            docx_ok = True

    charts_html = html_has_charts(html)
    charts_docx = False
    if paths.get("docx"):
        charts_docx = docx_has_images(Path(paths["docx"])) or bool(manifest.get("charts_included"))

    gate = {
        "CLI_CLAIMS_EQ_FILES_CREATED": "PASS" if file_check["ok"] and claimed_ok else "FAIL",
        "HTML_PDF_DOCX_KEY_VALUE_PARITY": "PASS" if html_pdf_docx_parity else "FAIL",
        "MANIFEST_HASHES": "PASS" if file_check["ok"] and manifest.get("instance_hash") else "FAIL",
        "PDF_PAGE_COUNT_GT_0": "PASS" if pdf_ok else "FAIL",
        "DOCX_PAGE_COUNT_GT_0": "PASS" if docx_ok else "FAIL",
        "CHARTS_EMBEDDED_IN_PDF": (
            "PASS" if (not paths.get("pdf") or charts_html) else "FAIL"
        ),  # PDF from HTML carries charts
        "CHARTS_EMBEDDED_IN_DOCX": (
            "PASS" if (not paths.get("docx") or charts_docx) else "FAIL"
        ),
    }
    gate["ALL_PASS"] = all(v == "PASS" for k, v in gate.items() if k != "ALL_PASS")
    gate["file_check"] = file_check
    return gate

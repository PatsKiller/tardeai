"""cio_report_render.py — Phase 4 multi-format renderers over one report view.

Consumes only `build_report_view(model)` output (or a full model that will be
projected). Never recalculates financial arithmetic.

Formats:
  * HTML  — self-contained, print-friendly
  * DOCX  — python-docx when available
  * PDF   — Chromium/Chrome headless or weasyprint/wkhtmltopdf when available
  * JSON  — model + view + parity manifest

READ_ONLY_ADVISORY. No broker / Telegram.
"""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Optional

from scripts.lib.cio_report_view import (
    REPORT_ARCHITECTURE_VERSION,
    build_report_view,
    section_by_id,
)


def _e(v: Any) -> str:
    return (
        str(v if v is not None else "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _ensure_view(model_or_view: dict[str, Any]) -> dict[str, Any]:
    if model_or_view.get("architecture_version") and model_or_view.get("sections"):
        return model_or_view
    if model_or_view.get("view") and isinstance(model_or_view["view"], dict):
        return model_or_view["view"]
    return build_report_view(model_or_view)


# ─────────────────────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────────────────────

_CSS = """
@page {
  size: letter;
  margin: 1.7cm 1.5cm 2.0cm 1.5cm;
  @top-left {
    content: "Trade AI Private Investment Office";
    font-size: 8pt; color: #6b7280; font-family: 'Segoe UI', Calibri, sans-serif;
  }
  @top-right {
    content: "Institutional Report v2 · READ_ONLY_ADVISORY";
    font-size: 8pt; color: #6b7280; font-family: 'Segoe UI', Calibri, sans-serif;
  }
  @bottom-center {
    content: "page " counter(page) " of " counter(pages);
    font-size: 8pt; color: #555; font-family: 'Segoe UI', Calibri, sans-serif;
  }
}
:root {
  --navy:#1F3864; --green:#2E7D32; --ink:#1a1a1a; --muted:#555;
  --line:#D5DAE1; --bg:#fff; --hi:#F4F6F9; --burgundy:#8B1E1E;
}
* { box-sizing:border-box; }
body {
  font-family: 'Segoe UI', Calibri, 'Helvetica Neue', Arial, sans-serif;
  color: var(--ink); margin: 0; padding: 0; background: var(--bg);
  font-size: 10.5pt; line-height: 1.45;
}
.page { padding: 28px 36px; page-break-after: always; }
.page:last-child { page-break-after: auto; }
.cover {
  background: linear-gradient(135deg, #16294D 0%, #1F3864 55%, #1a4d2e 130%);
  color: #fff; min-height: 90vh; padding: 56px 48px; position: relative;
}
.cover .brand { font-size: 11pt; letter-spacing: 2px; text-transform: uppercase;
  opacity: .88; border-bottom: 1px solid rgba(255,255,255,.35); padding-bottom: 12px; }
.cover h1 { font-size: 28pt; font-weight: 300; margin: 48px 0 8px; }
.cover .sub { font-size: 12pt; opacity: .92; }
.cover .meta-block { margin-top: 40px; font-size: 10pt; line-height: 1.9; opacity: .93; }
.cover .disclaimer { position: absolute; bottom: 36px; left: 48px; right: 48px;
  font-size: 8pt; opacity: .7; }
.toc { font-size: 11pt; line-height: 2.0; color: var(--navy); }
.toc a { color: var(--navy); text-decoration: none; }
h1 { font-size: 1.45rem; color: var(--navy); margin: 0 0 .35rem; }
h2 {
  font-size: 13pt; color: var(--navy); margin: 1.4rem 0 .55rem;
  border-bottom: 2px solid var(--green); padding-bottom: 4px;
  page-break-after: avoid; break-after: avoid;
}
h3 {
  font-size: 11pt; color: var(--navy); margin: 1rem 0 .4rem;
  page-break-after: avoid; break-after: avoid;
}
.lede { color: var(--muted); font-size: 9.5pt; margin-bottom: .75rem; }
.meta { display:flex; flex-wrap:wrap; gap:.6rem 1.1rem; color:var(--muted);
  font-size: 8.5pt; margin-bottom: 1rem; }
.meta span { white-space: nowrap; }
.section { margin-bottom: 1.15rem; page-break-inside: avoid; break-inside: avoid; }
.grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(148px,1fr)); gap:.55rem; }
.box {
  background: var(--hi); border: 1px solid var(--line); border-top: 3px solid var(--green);
  border-radius: 2px; padding: .55rem .7rem;
}
.box .k { font-size: 7.5pt; text-transform: uppercase; letter-spacing: .04em; color: var(--muted); }
.box .v { font-size: 12pt; font-weight: 600; margin-top: .15rem; color: var(--navy);
  font-variant-numeric: tabular-nums; }
table {
  width: 100%; border-collapse: collapse; font-size: 9.5pt; margin: .35rem 0 .75rem;
  page-break-inside: avoid; break-inside: avoid;
}
thead { display: table-header-group; }
th {
  background: var(--navy); color: #fff; font-weight: 600;
  padding: .35rem .5rem; text-align: left; font-size: 8.5pt;
}
td { border-bottom: 1px solid var(--line); padding: .35rem .5rem; vertical-align: top; }
tr:nth-child(even) td { background: var(--hi); }
td.num, th.num { text-align: right; font-variant-numeric: tabular-nums; }
.pos { color: var(--green); font-weight: 600; }
.neg { color: var(--burgundy); font-weight: 600; }
.badge.flag { background: #FDF3E7; color: #6B4A00; font-size: 7.5pt;
  padding: .1rem .35rem; border-radius: 2px; border-left: 2px solid #E6A23C; }
.footnote, .source-note { font-size: 8pt; color: var(--muted); margin-top: .2rem; font-style: italic; }
ul { margin: .25rem 0 .55rem 1.1rem; padding: 0; }
.charts { display: flex; flex-wrap: wrap; gap: 14px; margin: 10px 0 16px; }
.chart {
  flex: 1 1 280px; max-width: 100%;
  border: 1px solid var(--line); background: #fff; padding: 8px 10px 6px;
  page-break-inside: avoid; break-inside: avoid;
}
.chart img, .chart object, .chart svg { max-width: 100%; height: auto; display: block; margin: 0 auto; }
.chart .cap { font-size: 8pt; color: var(--muted); margin-top: 4px; text-align: center; }
.footer-note {
  font-size: 8pt; color: #999; border-top: 1px solid var(--line);
  margin-top: 18px; padding-top: 8px;
}
@media print {
  body { font-size: 10pt; }
  .cover { min-height: 100vh; }
  h2, h3 { break-after: avoid; }
  table, .chart, .section { break-inside: avoid; }
  thead { display: table-header-group; }
}
"""
def render_html_from_view(view: dict[str, Any]) -> str:
    """Render the full institutional report HTML from a shared view."""
    view = _ensure_view(view)
    sections_html: list[str] = []

    cover = section_by_id(view, "cover") or {}
    meta = cover.get("meta") or {}
    facts = view.get("facts") or {}
    fp = str(view.get("facts_fingerprint") or "")[:12]

    # ── Cover page ──
    sections_html.append('<div class="page cover"><div class="brand">Trade AI · Private Investment Office</div>')
    sections_html.append(f"<h1>{_e(cover.get('title') or 'Institutional Report')}</h1>")
    sections_html.append(
        '<div class="sub">CIO Investment Committee package · READ_ONLY_ADVISORY · no execution authority</div>'
    )
    sections_html.append(
        f'<div class="meta-block">'
        f'As-of: {_e(meta.get("as_of") or facts.get("as_of") or "—")}<br/>'
        f'Source SHA: {_e(meta.get("source_sha") or "—")}<br/>'
        f'Report: {_e(view.get("report_version") or "")} · arch {_e(meta.get("architecture") or REPORT_ARCHITECTURE_VERSION)}<br/>'
        f'Facts fingerprint: {_e(fp)}…<br/>'
        f'Traceability: {_e(meta.get("traceability_pct"))}%'
        f'</div>'
    )
    sections_html.append(
        '<div class="disclaimer">Advisory automation only. Figures come from a single canonical '
        'model snapshot. Not investment advice. No broker, order, or stop authority.</div></div>'
    )

    # ── Contents ──
    toc_items = [
        ("Part A — CIO Investment Committee", "#part-a"),
        ("Decisions Now", "#decisions_now"),
        ("Capital Plan", "#capital_plan"),
        ("Portfolio Posture", "#portfolio_posture"),
        ("Charts", "#charts"),
        ("Part B — Institutional Portfolio Book", "#part-b"),
        ("Asset Allocation", "#allocation"),
        ("Data Coverage & Provenance", "#coverage"),
        ("Disclosure", "#disclosure"),
    ]
    sections_html.append('<div class="page"><h2>Contents</h2><div class="toc">')
    # Extend TOC with Phase 6 analytic sections when present
    for sid, title in (
        ("performance_definitions", "Performance Definitions"),
        ("change_in_value", "Change in Portfolio Value"),
        ("lookthrough_coverage", "Look-Through Coverage"),
        ("valuation_coverage", "Valuation Coverage"),
        ("tax_lots_summary", "Tax Lots / Unrealized"),
        ("phase6_exit_gate", "Analytic Completeness Gate"),
    ):
        if section_by_id(view, sid):
            toc_items.append((title, f"#{sid}"))
    for label, href in toc_items:
        sections_html.append(f'<div><a href="{href}">{_e(label)}</a></div>')
    sections_html.append("</div></div>")
    sections_html.append('<div class="page">')
    sections_html.append(
        f'<div class="meta"><span>as_of {_e(meta.get("as_of"))}</span>'
        f'<span>source SHA {_e(meta.get("source_sha") or "—")}</span>'
        f'<span>facts {_e(fp)}</span></div>'
    )
    sections_html.append('<h2 id="part-a">Part A — CIO Investment Committee</h2>')
    letter = section_by_id(view, "cio_letter") or {}
    sections_html.append(f"<div class='section'><h3>{_e(letter.get('title'))}</h3>")
    if letter.get("thesis_summary"):
        sections_html.append(f"<p>{_e(letter['thesis_summary'])}</p>")
    if letter.get("stance"):
        sections_html.append(f"<p><strong>Stance:</strong> {_e(letter['stance'])}</p>")
    if letter.get("priorities"):
        sections_html.append("<p><strong>Priorities:</strong></p><ul>")
        for p in letter["priorities"]:
            sections_html.append(f"<li>{_e(p)}</li>")
        sections_html.append("</ul>")
    if letter.get("what_not_to_do"):
        sections_html.append("<p><strong>What not to do:</strong></p><ul>")
        for p in letter["what_not_to_do"]:
            sections_html.append(f"<li>{_e(p)}</li>")
        sections_html.append("</ul>")
    sections_html.append("</div>")

    for sid in ("decisions_now", "capital_plan", "portfolio_posture"):
        sec = section_by_id(view, sid)
        if not sec:
            continue
        sections_html.append(
            f"<div class='section' id='{_e(sid)}'><h3>{_e(sec.get('title'))}</h3>"
        )
        if sec.get("kind") == "table":
            rows = sec.get("rows") or []
            if rows:
                headers = sec.get("headers") or []
                sections_html.append("<table><thead><tr>")
                for i, h in enumerate(headers):
                    cls = " class='num'" if i >= 2 and i <= 4 else ""
                    sections_html.append(f"<th{cls}>{_e(h)}</th>")
                sections_html.append("</tr></thead><tbody>")
                for row in rows:
                    sections_html.append("<tr>")
                    for i, cell in enumerate(row):
                        cs = str(cell if cell is not None else "—")
                        cls = "num"
                        if i >= 2 and i <= 4:
                            if cs.startswith("-") or cs.startswith("−"):
                                cls += " neg"
                            elif cs.startswith("+$") or (cs.startswith("$") is False and cs.startswith("+")):
                                cls += " pos"
                            sections_html.append(f"<td class='{cls}'>{_e(cs)}</td>")
                        else:
                            sections_html.append(f"<td>{_e(cs)}</td>")
                    sections_html.append("</tr>")
                sections_html.append("</tbody></table>")
            else:
                sections_html.append(f"<p class='lede'>{_e(sec.get('empty_message') or '—')}</p>")
        elif sec.get("kind") == "kv":
            sections_html.append("<div class='grid'>")
            for k, v in sec.get("rows") or []:
                sections_html.append(
                    f"<div class='box'><div class='k'>{_e(k)}</div>"
                    f"<div class='v' style='font-size:.95rem'>{_e(v)}</div></div>"
                )
            sections_html.append("</div>")
            st = sec.get("sector_table")
            if st and st.get("rows"):
                sections_html.append("<table><thead><tr>")
                for h in st.get("headers") or []:
                    sections_html.append(f"<th>{_e(h)}</th>")
                sections_html.append("</tr></thead><tbody>")
                for row in st["rows"]:
                    sections_html.append("<tr>" + "".join(f"<td>{_e(c)}</td>" for c in row) + "</tr>")
                sections_html.append("</tbody></table>")
        sections_html.append("</div>")

    # ── Charts (Phase 5) ──
    chart_list = view.get("charts_for_html") or []
    if chart_list:
        sections_html.append('<div class="section" id="charts"><h3>Charts</h3>')
        sections_html.append('<div class="charts">')
        for ch in chart_list:
            sections_html.append('<div class="chart">')
            # Inline SVG preferred when present in bundle
            svg = None
            charts_bundle = view.get("charts") or {}
            entry = (charts_bundle.get("charts") or {}).get(ch.get("key") or "")
            if entry and entry.get("svg"):
                svg = entry["svg"]
            if svg:
                sections_html.append(svg)
            elif ch.get("data_uri"):
                sections_html.append(
                    f'<img src="{_e(ch["data_uri"])}" alt="{_e(ch.get("alt_caption") or ch.get("title"))}"/>'
                )
            cap_bits = [ch.get("title") or ""]
            if ch.get("source_note"):
                cap_bits.append(str(ch["source_note"]))
            if ch.get("units"):
                cap_bits.append(f"units: {ch['units']}")
            if ch.get("quality_flag"):
                cap_bits.append(f"flagged: {ch['quality_flag']}")
            sections_html.append(f'<div class="cap">{_e(" · ".join(x for x in cap_bits if x))}</div>')
            sections_html.append("</div>")
        sections_html.append("</div>")
        skipped = (view.get("charts") or {}).get("skipped") or {}
        if skipped:
            sections_html.append('<p class="source-note">Charts withheld (insufficient or invalid source): ')
            sections_html.append(_e("; ".join(f"{k}: {v}" for k, v in list(skipped.items())[:8])))
            sections_html.append("</p>")
        sections_html.append("</div>")
    funnel = section_by_id(view, "opportunity_funnel") or {}
    sections_html.append(f"<div class='section'><h3>{_e(funnel.get('title'))}</h3>")
    for key, title in (
        ("sector_opportunities", "Sector opportunities"),
        ("watch_additions", "Watch additions"),
        ("reentry_candidates", "Re-entry candidates"),
        ("research_gaps", "Research gaps"),
    ):
        items = funnel.get(key) or []
        if not items:
            continue
        sections_html.append(f"<p><strong>{_e(title)}:</strong></p><ul>")
        for it in items[:8]:
            if isinstance(it, dict):
                label = (
                    f"{it.get('sector') or it.get('symbol') or '—'} "
                    f"({it.get('state') or it.get('verdict') or '—'}): "
                    f"{it.get('recommendation') or it.get('label') or ''}"
                )
            else:
                label = str(it)
            sections_html.append(f"<li>{_e(label.strip())}</li>")
        sections_html.append("</ul>")
    sections_html.append("</div>")

    ct = section_by_id(view, "counter_thesis") or {}
    sections_html.append(f"<div class='section'><h3>{_e(ct.get('title'))}</h3>")
    if ct.get("highest_impact_unknowns"):
        sections_html.append("<p><strong>Highest-impact unknowns:</strong></p><ul>")
        for u in ct["highest_impact_unknowns"]:
            sections_html.append(f"<li>{_e(u)}</li>")
        sections_html.append("</ul>")
    sections_html.append("</div>")

    sections_html.append("</div>")  # end part A page

    sections_html.append('<div class="page">')
    sections_html.append('<h2 id="part-b">Part B — Institutional Portfolio Book</h2>')
    for sid in (
        "portfolio_book", "accounts", "allocation", "performance",
        "performance_definitions", "change_in_value", "benchmark_alignment",
        "lookthrough_coverage", "lookthrough_sectors", "valuation_coverage",
        "attribution", "tax_lots_summary", "tax_lots_detail", "income",
        "phase6_exit_gate",
        "coverage",
        "field_coverage_matrix", "known_gap_resolutions", "quality_flags",
    ):
        sec = section_by_id(view, sid)
        if not sec:
            continue
        sections_html.append(
            f"<div class='section' id='{_e(sid)}'><h3>{_e(sec.get('title'))}</h3>"
        )
        if sec.get("kind") == "table":
            headers = sec.get("headers") or []
            rows = sec.get("rows") or []
            sections_html.append("<table><thead><tr>")
            for h in headers:
                sections_html.append(f"<th>{_e(h)}</th>")
            sections_html.append("</tr></thead><tbody>")
            for row in rows:
                cells = []
                for c in row:
                    cs = str(c if c is not None else "—")
                    if "account-aggregated" in cs or cs == "flagged":
                        cells.append(
                            f"<td>{_e(cs)} <span class='badge flag'>flagged</span></td>"
                        )
                    else:
                        cells.append(f"<td>{_e(cs)}</td>")
                sections_html.append("<tr>" + "".join(cells) + "</tr>")
            sections_html.append("</tbody></table>")
            if sid == "allocation":
                sections_html.append(
                    '<p class="source-note">Source: holdings settled cash + equity market value. '
                    "Units: USD and weight % of total (never dollars as percentages).</p>"
                )
        elif sec.get("kind") == "kv":
            sections_html.append("<div class='grid'>")
            for k, v in sec.get("rows") or []:
                sections_html.append(
                    f"<div class='box'><div class='k'>{_e(k)}</div>"
                    f"<div class='v' style='font-size:.95rem'>{_e(v)}</div></div>"
                )
            sections_html.append("</div>")
            if sec.get("unavailable"):
                sections_html.append("<p class='footnote'><strong>Unavailable fields:</strong> "
                                     + _e(", ".join(str(x) for x in sec["unavailable"][:20])) + "</p>")
        elif sec.get("kind") == "list":
            items = sec.get("items") or sec.get("highest_impact_unknowns") or []
            if items:
                sections_html.append("<ul>")
                for u in items:
                    us = str(u)
                    flag = " <span class='badge flag'>flagged</span>" if (
                        "account-aggregated" in us or "flagged" in us
                    ) else ""
                    sections_html.append(f"<li>{_e(us)}{flag}</li>")
                sections_html.append("</ul>")
        sections_html.append("</div>")

    disc = section_by_id(view, "disclosure") or {}
    sections_html.append(
        f"<div class='section' id='disclosure'><h3>{_e(disc.get('title') or 'Disclosure')}</h3>"
        f"<p class='lede'>{_e(disc.get('text') or '')}</p>"
        f"<div class='footer-note'>Facts fingerprint {_e(fp)} · "
        f"{_e(view.get('report_version') or '')} · {_e(REPORT_ARCHITECTURE_VERSION)}</div></div>"
    )
    sections_html.append("</div>")  # end part B page
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<title>{_e(cover.get('title') or 'Institutional Report v2')}</title>"
        f"<style>{_CSS}</style></head><body>"
        + "\n".join(sections_html)
        + "</body></html>"
    )

# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def render_docx_from_view(view: dict[str, Any], out: Path) -> Path:
    """Write DOCX from the shared view. Requires python-docx.

    Phase 5: running header/footer, keep_with_next headings, non-split tables,
    chart inventory with source notes, disclosure kept with prior content.
    """
    try:
        import docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    view = _ensure_view(view)
    document = docx.Document()
    document.core_properties.title = "Trade AI Institutional Report v2"
    document.core_properties.author = "Trade AI Investment Office (Alex, CIO)"
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    # Margins + running header/footer
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = "Trade AI Private Investment Office  ·  Institutional Report v2  ·  READ_ONLY_ADVISORY"
        for run in hp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Advisory only · page ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        # PAGE field
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run2 = fp.add_run()
        run2._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run2._r.append(fld2)
        run2.font.size = Pt(8)

    def _keep_with_next(paragraph) -> None:
        try:
            paragraph.paragraph_format.keep_with_next = True
        except Exception:
            pass

    def _table(headers: list[str], rows: list[list[Any]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
        # Repeat header row
        try:
            tr = table.rows[0]._tr
            tr_pr = tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        except Exception:
            pass
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val if val is not None else "—")
            # cantSplit on data rows
            try:
                tr = table.rows[-1]._tr
                tr_pr = tr.get_or_add_trPr()
                cant = OxmlElement("w:cantSplit")
                tr_pr.append(cant)
            except Exception:
                pass
    cover = section_by_id(view, "cover") or {}
    meta = cover.get("meta") or {}
    t = document.add_heading(str(cover.get("title") or "Trade AI — Institutional Report"), level=0)
    _keep_with_next(t)
    sub = document.add_paragraph()
    r = sub.add_run(
        f"{view.get('report_version') or ''}  ·  {view.get('authority') or ''}  ·  "
        f"as-of {view.get('as_of') or ''}  ·  arch {meta.get('architecture') or REPORT_ARCHITECTURE_VERSION}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    if meta.get("source_sha"):
        p = document.add_paragraph()
        pr = p.add_run(
            f"Source SHA: {meta['source_sha']}  ·  facts {str(view.get('facts_fingerprint') or '')[:16]}"
        )
        pr.font.size = Pt(8)
        pr.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    # Contents
    document.add_heading("Contents", level=1)
    for label in (
        "Part A — CIO Investment Committee",
        "Decisions Now",
        "Capital Plan",
        "Charts",
        "Part B — Institutional Portfolio Book",
        "Asset Allocation",
        "Disclosure",
    ):
        document.add_paragraph(label, style="List Number")

    document.add_page_break()
    h = document.add_heading("Part A — CIO Investment Committee", level=1)
    _keep_with_next(h)
    letter = section_by_id(view, "cio_letter") or {}
    document.add_heading(str(letter.get("title") or "CIO Letter"), level=2)
    if letter.get("stance"):
        document.add_paragraph(f"Stance: {letter['stance']}")
    if letter.get("risk_posture"):
        document.add_paragraph(f"Risk posture: {letter['risk_posture']}")
    if letter.get("priorities"):
        document.add_paragraph("Priorities:").runs[0].bold = True
        for item in letter["priorities"]:
            document.add_paragraph(str(item), style="List Bullet")
    if letter.get("what_not_to_do"):
        document.add_paragraph("Guardrails (what not to do):").runs[0].bold = True
        for item in letter["what_not_to_do"]:
            document.add_paragraph(str(item), style="List Bullet")

    # Charts — embed PNG when convertible from SVG; always include governance table
    chart_bundle = view.get("charts") or {}
    included = list(chart_bundle.get("included") or [])
    if included or chart_bundle.get("skipped"):
        h = document.add_heading("Charts", level=2)
        _keep_with_next(h)
        document.add_paragraph(
            "Charts share the same model snapshot as HTML/PDF. "
            "Each figure carries source, units, and quality notes."
        )
        for key in included:
            c = (chart_bundle.get("charts") or {}).get(key) or {}
            cap = document.add_paragraph()
            run = cap.add_run(str(c.get("title") or key))
            run.bold = True
            run.font.size = Pt(10)
            png_path = _svg_to_png_if_possible(c.get("svg_path") or c.get("svg"))
            if png_path:
                try:
                    document.add_picture(str(png_path), width=Inches(5.8))
                except Exception:
                    pass
            note = document.add_paragraph()
            bits = [c.get("source_note") or "", c.get("units") or "", c.get("quality_flag") or ""]
            nr = note.add_run(" · ".join(b for b in bits if b))
            nr.font.size = Pt(8)
            nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            nr.italic = True
        rows = []
        for key in included:
            c = (chart_bundle.get("charts") or {}).get(key) or {}
            rows.append([
                c.get("title") or key,
                c.get("units") or "—",
                c.get("source_note") or "—",
                c.get("quality_flag") or "—",
            ])
        if rows:
            _table(["Chart", "Units", "Source", "Quality"], rows)
        skipped = chart_bundle.get("skipped") or {}
        if skipped:
            document.add_paragraph("Withheld charts (source quality):").runs[0].bold = True
            for k, reason in list(skipped.items())[:10]:
                document.add_paragraph(f"{k}: {reason}", style="List Bullet")

    for sid in (
        "decisions_now", "capital_plan", "portfolio_posture",
        "opportunity_funnel", "counter_thesis",
        "portfolio_book", "accounts", "allocation", "performance",
        "performance_definitions", "change_in_value", "benchmark_alignment",
        "lookthrough_coverage", "lookthrough_sectors", "valuation_coverage",
        "attribution", "tax_lots_summary", "tax_lots_detail", "income",
        "phase6_exit_gate",
        "coverage",
        "disclosure",
    ):
        sec = section_by_id(view, sid)
        if not sec:
            continue
        if sid == "portfolio_book":
            document.add_page_break()
            h = document.add_heading("Part B — Institutional Portfolio Book", level=1)
            _keep_with_next(h)
        h = document.add_heading(str(sec.get("title") or sid), level=2)
        _keep_with_next(h)
        kind = sec.get("kind")
        if kind == "table":
            rows = sec.get("rows") or []
            if rows:
                _table(list(sec.get("headers") or []), rows)
            else:
                document.add_paragraph(str(sec.get("empty_message") or "—"))
        elif kind == "kv":
            _table(["Metric", "Value"], list(sec.get("rows") or []))
            st = sec.get("sector_table")
            if st and st.get("rows"):
                document.add_paragraph("Sector posture:").runs[0].bold = True
                _table(list(st.get("headers") or []), st["rows"])
            if sec.get("unavailable"):
                document.add_paragraph("Unavailable fields:").runs[0].bold = True
                for f in sec["unavailable"][:20]:
                    document.add_paragraph(str(f), style="List Bullet")
        elif kind == "funnel":
            for key, title in (
                ("sector_opportunities", "Sector opportunities"),
                ("research_gaps", "Research gaps"),
            ):
                items = sec.get(key) or []
                if not items:
                    continue
                document.add_paragraph(f"{title}:").runs[0].bold = True
                for it in items[:8]:
                    if isinstance(it, dict):
                        text = (
                            f"{it.get('sector') or it.get('symbol') or '—'} "
                            f"({it.get('state') or '—'}): {it.get('recommendation') or ''}"
                        )
                    else:
                        text = str(it)
                    document.add_paragraph(text.strip(), style="List Bullet")
        elif kind == "list":
            if sec.get("highest_impact_unknowns"):
                document.add_paragraph("Highest-impact unknowns:").runs[0].bold = True
                for u in sec["highest_impact_unknowns"]:
                    document.add_paragraph(str(u), style="List Bullet")
        elif kind == "prose":
            document.add_paragraph(str(sec.get("text") or ""))

    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    return out


def _svg_to_png_if_possible(svg_path_or_content: Any) -> Optional[Path]:
    """Best-effort SVG→PNG for DOCX embedding (rsvg-convert / inkscape / convert)."""
    if not svg_path_or_content:
        return None
    tmp_svg: Optional[Path] = None
    try:
        if isinstance(svg_path_or_content, (str, Path)) and Path(svg_path_or_content).exists():
            svg_path = Path(svg_path_or_content)
        else:
            tmp_svg = Path(tempfile.mkstemp(suffix=".svg")[1])
            tmp_svg.write_text(str(svg_path_or_content), encoding="utf-8")
            svg_path = tmp_svg
        png_path = svg_path.with_suffix(".png")
        if png_path.exists() and png_path.stat().st_size > 0:
            return png_path
        for cmd in (
            ["rsvg-convert", "-o", str(png_path), str(svg_path)],
            ["inkscape", str(svg_path), "--export-type=png", f"--export-filename={png_path}"],
            ["convert", str(svg_path), str(png_path)],
        ):
            if not shutil.which(cmd[0]):
                continue
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if r.returncode == 0 and png_path.exists() and png_path.stat().st_size > 0:
                return png_path
    except Exception:
        return None
    finally:
        pass
    return None


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def _which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)


PDF_RENDERER_COMMANDS = (
    "weasyprint",
    "wkhtmltopdf",
    "chromium",
    "chromium-browser",
    "google-chrome",
    "google-chrome-stable",
)


def has_docx_library() -> bool:
    """True only when python-docx is actually importable. Never guessed."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _weasyprint_python_available() -> bool:
    try:
        import weasyprint  # noqa: F401
        return True
    except ImportError:
        return False


def detect_renderers() -> dict[str, Any]:
    """Honest capability probe for HTML / DOCX / PDF.

    PDF is available only when weasyprint (CLI or Python), wkhtmltopdf, or a
    Chromium/Chrome binary is on PATH. Playwright's cached browser is *not*
    counted — that would fake production PDF readiness.
    """
    commands = {cmd: bool(_which(cmd)) for cmd in PDF_RENDERER_COMMANDS}
    weasy_py = _weasyprint_python_available()
    engines = [cmd for cmd, present in commands.items() if present]
    if weasy_py and "weasyprint" not in engines:
        engines.append("weasyprint-python")
    pdf_ok = bool(engines)
    docx_ok = has_docx_library()
    return {
        "html": {"available": True, "status": "ok", "ok": True},
        "docx": {
            "available": docx_ok,
            "status": "ok" if docx_ok else "missing",
            "ok": docx_ok,
            "library": "python-docx",
            "reason": None if docx_ok else "python-docx is not installed",
        },
        "pdf": {
            "available": pdf_ok,
            "status": "ok" if pdf_ok else "missing",
            "ok": pdf_ok,
            "commands": commands,
            "weasyprint_python": weasy_py,
            "engines": engines,
            "reason": (
                None if pdf_ok
                else "No PDF renderer (weasyprint/wkhtmltopdf/chromium)"
            ),
        },
    }


def has_pdf_renderer() -> bool:
    """True only when a real weasyprint / wkhtmltopdf / chromium engine exists."""
    return bool((detect_renderers().get("pdf") or {}).get("available"))


def render_pdf_from_html(html: str, out: Path) -> Path:
    """Render HTML to PDF via the first available engine. Never fakes a PDF."""
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    html_path = out.with_suffix(".print.html")
    html_path.write_text(html, encoding="utf-8")

    # 1) weasyprint CLI
    if _which("weasyprint"):
        r = subprocess.run(
            ["weasyprint", str(html_path), str(out)],
            capture_output=True, text=True, timeout=120,
        )
        if r.returncode == 0 and out.exists() and out.stat().st_size > 0:
            return out

    # 2) weasyprint Python API (same engine, no CLI wrapper)
    if _weasyprint_python_available():
        try:
            from weasyprint import HTML as WeasyHTML  # type: ignore
            WeasyHTML(filename=str(html_path)).write_pdf(str(out))
            if out.exists() and out.stat().st_size > 0:
                return out
        except Exception:
            pass

    # 3) wkhtmltopdf
    if _which("wkhtmltopdf"):
        r = subprocess.run(
            ["wkhtmltopdf", "--quiet", str(html_path), str(out)],
            capture_output=True, text=True, timeout=120,
        )
        if r.returncode == 0 and out.exists() and out.stat().st_size > 0:
            return out

    # 4) Chromium / Chrome headless
    for browser in ("chromium", "chromium-browser", "google-chrome", "google-chrome-stable"):
        bin_path = _which(browser)
        if not bin_path:
            continue
        r = subprocess.run(
            [
                bin_path, "--headless", "--disable-gpu", "--no-pdf-header-footer",
                f"--print-to-pdf={out}", str(html_path),
            ],
            capture_output=True, text=True, timeout=120,
        )
        if out.exists() and out.stat().st_size > 0:
            return out

    raise RuntimeError("No PDF renderer available (weasyprint/wkhtmltopdf/chromium)")


# ─────────────────────────────────────────────────────────────────────────────
# Bundle export (one snapshot → all formats)
# ─────────────────────────────────────────────────────────────────────────────

def export_report_formats(
    model: dict[str, Any],
    out_dir: Path | str,
    *,
    basename: str = "cio_institutional_report_v2",
    write_docx: bool = True,
    write_pdf: bool = True,
    formats: Optional[list[str]] = None,
    report_id: Optional[str] = None,
) -> dict[str, Any]:
    """Export HTML + DOCX + PDF + JSON + immutable instance manifest.

    Phase 7: one model snapshot → all formats → cross-format parity →
    immutable report instance manifest. Never invents alternate facts.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if formats is not None:
        fmt_set = {f.strip().lower() for f in formats}
        write_html = "html" in fmt_set or not fmt_set
        write_docx = "docx" in fmt_set
        write_pdf = "pdf" in fmt_set
    else:
        write_html = True

    formats_requested = []
    if write_html:
        formats_requested.append("html")
    if write_docx:
        formats_requested.append("docx")
    if write_pdf:
        formats_requested.append("pdf")

    # Strip heavy embedded HTML before projecting view (rebuilt from view)
    model_core = {k: v for k, v in (model or {}).items() if k not in ("html", "charts")}
    view = build_report_view(model_core)

    # Phase 5 charts from the same model snapshot
    charts_dir = out_dir / f"{basename}_charts"
    try:
        from scripts.lib.cio_report_charts import build_charts, charts_for_html
        chart_bundle = build_charts(model_core, charts_dir=charts_dir)
        view["charts"] = chart_bundle
        view["charts_for_html"] = charts_for_html(chart_bundle)
    except Exception as exc:
        chart_bundle = {"included": [], "skipped": {"_error": str(exc)[:160]}, "charts": {}}
        view["charts"] = chart_bundle
        view["charts_for_html"] = []

    html = render_html_from_view(view)

    # Attach presentation artifacts back onto a clean model copy
    model_out = dict(model_core)
    model_out["view"] = view
    model_out["html"] = html
    model_out["charts"] = chart_bundle
    model_out["architecture_version"] = REPORT_ARCHITECTURE_VERSION
    model_out["facts_fingerprint"] = view["facts_fingerprint"]
    # Normalize allocation onto part_b for any downstream consumer
    pb = dict(model_out.get("part_b") or {})
    pb["allocation"] = view["facts"].get("allocation_usd") or pb.get("allocation") or {}
    pb["allocation_weight_pct"] = view["facts"].get("allocation_weight_pct") or {}
    model_out["part_b"] = pb
    paths: dict[str, Optional[str]] = {}
    errors: dict[str, str] = {}

    model_path = out_dir / f"{basename}.model.json"
    # Write model after manifests? First write core artifacts
    view_path = out_dir / f"{basename}.view.json"
    view_path.write_text(json.dumps(view, indent=2, default=str), encoding="utf-8")
    paths["view_json"] = str(view_path)

    html_path = out_dir / f"{basename}.html"
    if write_html:
        html_path.write_text(html, encoding="utf-8")
        paths["html"] = str(html_path)

    if write_docx:
        docx_path = out_dir / f"{basename}.docx"
        try:
            render_docx_from_view(view, docx_path)
            paths["docx"] = str(docx_path)
        except Exception as exc:
            errors["docx"] = str(exc)[:200]
            paths["docx"] = None

    if write_pdf:
        pdf_path = out_dir / f"{basename}.pdf"
        try:
            render_pdf_from_html(html, pdf_path)
            paths["pdf"] = str(pdf_path)
        except Exception as exc:
            errors["pdf"] = str(exc)[:200]
            paths["pdf"] = None

    # ── Phase 7 parity + immutable instance manifest ──
    from scripts.lib.cio_report_pipeline import (
        build_instance_manifest,
        build_phase7_exit_gate,
        compare_key_values,
        extract_key_values_from_docx,
        extract_key_values_from_html,
        extract_key_values_from_view,
    )

    key_values = extract_key_values_from_view(view, model_out)
    html_vals = extract_key_values_from_html(html) if write_html else {}
    html_cmp = compare_key_values(key_values, html_vals) if write_html else {"ok": True, "hard_mismatches": [], "checked": 0}
    docx_cmp = {"ok": True, "hard_mismatches": [], "checked": 0, "skipped": True}
    if paths.get("docx"):
        docx_vals = extract_key_values_from_docx(Path(paths["docx"]))
        docx_cmp = compare_key_values(key_values, docx_vals)

    # Allocation unit regression (Phase 5.8)
    html_text = html
    alloc_usd = view["facts"].get("allocation_usd") or {}
    absurd_pct = any(
        f"{float(v):.2f}%" in html_text
        for v in alloc_usd.values()
        if v is not None and float(v) > 100
    )

    parity = {
        "architecture_version": REPORT_ARCHITECTURE_VERSION,
        "pipeline_version": "pipeline_1.0.0",
        "facts_fingerprint": view["facts_fingerprint"],
        "section_ids": view.get("section_ids"),
        "formats": {k: bool(v) for k, v in paths.items()},
        "formats_requested": formats_requested,
        "errors": errors,
        "charts_included": list(chart_bundle.get("included") or []),
        "charts_skipped": chart_bundle.get("skipped") or {},
        "key_values": key_values,
        "html_parity": html_cmp,
        "docx_parity": docx_cmp,
        # Hard OK requires HTML parity. DOCX is soft when extraction is sparse/skipped
        # (CI images without full office stack should not fail the pipeline gate).
        "ok": bool(html_cmp.get("ok")) and (
            bool(docx_cmp.get("ok"))
            if (not docx_cmp.get("skipped") and int(docx_cmp.get("checked") or 0) >= 3)
            else True
        ),
        "unit_guards": {
            "allocation_weights_le_100": all(
                abs(float(v)) <= 100.01
                for v in (view["facts"].get("allocation_weight_pct") or {}).values()
                if v is not None
            ),
            "allocation_no_dollar_as_percent": not absurd_pct,
            "decision_count": len(view["facts"].get("decisions") or []),
            "decision_symbols_unique": len(view["facts"].get("decisions") or [])
                == len({d.get("symbol") for d in (view["facts"].get("decisions") or [])}),
            "no_risk_return_without_vol": "risk_return" not in (chart_bundle.get("included") or [])
                or "risk_return" not in (chart_bundle.get("skipped") or {}),
        },
        "phase5_exit": {
            "allocation_unit_errors": 0 if not absurd_pct else 1,
            "charts_present": len(chart_bundle.get("included") or []),
            "raw_internal_codes_scan": (
                "STAGED_DEPLOYMENT" not in html_text and "RESEARCH_FIRST" not in html_text
            ),
        },
    }

    # Build instance manifest (paths known so far, exclude model/manifest self)
    instance = build_instance_manifest(
        model=model_out,
        view=view,
        paths={k: v for k, v in paths.items() if v},
        chart_bundle=chart_bundle,
        key_values=key_values,
        formats_requested=formats_requested,
        report_id=report_id,
    )
    phase7_gate = build_phase7_exit_gate(
        manifest=instance,
        parity=parity,
        formats_requested=formats_requested,
        paths=paths,
        html=html,
    )
    # Re-hash manifest after adding gate? Gate is verification, attach outside hash body
    instance_public = dict(instance)
    instance_public["phase7_exit_gate"] = phase7_gate
    parity["phase7_exit"] = phase7_gate

    parity_path = out_dir / f"{basename}.parity.json"
    parity_path.write_text(json.dumps(parity, indent=2, default=str), encoding="utf-8")
    paths["parity_json"] = str(parity_path)

    manifest_path = out_dir / f"{basename}.instance_manifest.json"
    manifest_path.write_text(json.dumps(instance_public, indent=2, default=str), encoding="utf-8")
    paths["instance_manifest"] = str(manifest_path)

    # Finalize model with instance id + write model json (include paths)
    model_out["report_id"] = instance["report_id"]
    model_out["instance_manifest"] = instance_public
    model_path = out_dir / f"{basename}.model.json"
    model_path.write_text(json.dumps(model_out, indent=2, default=str), encoding="utf-8")
    paths["model_json"] = str(model_path)

    # Refresh file hashes in a claims file (post-write of all artifacts)
    from scripts.lib.cio_report_pipeline import _sha256_file, verify_manifest_files
    claims = {
        "report_id": instance["report_id"],
        "formats_requested": formats_requested,
        "files_created": {k: v for k, v in paths.items() if v and Path(v).exists()},
        "file_sha256": {
            k: _sha256_file(Path(v))
            for k, v in paths.items()
            if v and Path(v).exists()
        },
        "phase7_exit_gate": phase7_gate,
    }
    claims_path = out_dir / f"{basename}.claims.json"
    claims_path.write_text(json.dumps(claims, indent=2, default=str), encoding="utf-8")
    paths["claims_json"] = str(claims_path)

    # CLI claims == files: verify claims paths exist
    claims_ok = all(Path(p).exists() for p in claims["files_created"].values())
    phase7_gate["CLI_CLAIMS_EQ_FILES_CREATED"] = "PASS" if claims_ok else "FAIL"
    phase7_gate["ALL_PASS"] = all(
        v == "PASS" for k, v in phase7_gate.items() if k not in ("ALL_PASS", "file_check")
    )

    return {
        "ok": bool(paths.get("html")) and phase7_gate.get("HTML_PDF_DOCX_KEY_VALUE_PARITY") == "PASS",
        "out_dir": str(out_dir),
        "paths": paths,
        "errors": errors,
        "facts_fingerprint": view["facts_fingerprint"],
        "architecture_version": REPORT_ARCHITECTURE_VERSION,
        "report_id": instance["report_id"],
        "parity": parity,
        "instance_manifest": instance_public,
        "phase7_exit_gate": phase7_gate,
        "claims": claims,
        "view": view,
        "model": model_out,
    }
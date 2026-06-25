#!/usr/bin/env python3
"""generate_weekly_docx.py — Weekly consolidated Word report generator.

Pulls data from all subsystems and produces a .docx weekly portfolio review.
Designed to run every Sunday evening via cron.

Usage:
    .venv/bin/python scripts/generate_weekly_docx.py
    .venv/bin/python scripts/generate_weekly_docx.py --dry-run
"""
import argparse, json, os, sys
from datetime import datetime, date, timedelta
from decimal import Decimal
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env")


def _f(v):
    return float(v) if isinstance(v, Decimal) else v


def _get_conn():
    import psycopg2
    pw = ""
    for line in (PROJECT_ROOT / ".env").read_text().splitlines():
        if line.startswith("DB_PASSWORD="): pw = line.split("=", 1)[1].strip()
    return psycopg2.connect(host="localhost", dbname="trade_ai", user="trade_ai", password=pw)


def gather_weekly_data(conn):
    """Gather all data needed for the weekly report."""
    import psycopg2.extras
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    today = date.today()
    week_start = today - timedelta(days=7)

    data = {"period_start": str(week_start), "period_end": str(today),
            "generated_at": datetime.now().isoformat()}

    # 1. Portfolio snapshot
    try:
        h = json.load(open(PROJECT_ROOT / "data/portfolios/state/holdings.json"))
        data["total_value"] = h.get("portfolio_totals", {}).get("total_value", 0)
        data["positions"] = len([p for p in h.get("holdings", []) if not p.get("is_cash") and (p.get("market_value", 0) > 50)])
        data["cash"] = sum(p.get("market_value", 0) for p in h.get("holdings", []) if p.get("is_cash"))
    except Exception:
        data["total_value"] = 0

    # 2. Paper trading performance
    cur.execute("""
        SELECT COUNT(*) FILTER (WHERE status='open') as open_trades,
               COUNT(*) FILTER (WHERE status='closed') as closed_total,
               COUNT(*) FILTER (WHERE status='closed' AND pnl > 0) as wins,
               COUNT(*) FILTER (WHERE status='closed' AND pnl <= 0) as losses,
               COALESCE(SUM(CASE WHEN pnl > 0 THEN pnl ELSE 0 END), 0) as gross_profit,
               COALESCE(SUM(CASE WHEN pnl < 0 THEN ABS(pnl) ELSE 0 END), 0) as gross_loss,
               AVG(r_multiple) FILTER (WHERE status='closed') as avg_r
        FROM paper_trades
    """)
    pt = cur.fetchone()
    data["paper_trades"] = {k: _f(v) for k, v in (pt or {}).items()}

    # 3. Proposals this week
    cur.execute("""
        SELECT status, COUNT(*) as cnt
        FROM paper_trade_proposals WHERE created_at >= %s
        GROUP BY status ORDER BY cnt DESC
    """, [str(week_start)])
    data["proposals"] = [dict(r) for r in cur.fetchall()]

    # 4. Agent activity
    cur.execute("""
        SELECT agent, COUNT(*) as analyses, AVG(confidence)::numeric(3,2) as avg_conf
        FROM watchlist_agent_results WHERE created_at > %s
        GROUP BY agent ORDER BY analyses DESC
    """, [str(week_start)])
    data["agent_activity"] = [dict(r) for r in cur.fetchall()]

    # 5. Agent calibration
    cur.execute("""
        SELECT agent_name, accuracy_pct, correct_count, wrong_count,
               total_recommendations, trending, avg_pnl_pct
        FROM agent_calibration WHERE window_days = 90 AND strategy_type IS NULL
        ORDER BY agent_name
    """)
    data["agent_calibration"] = [{k: _f(v) for k, v in dict(r).items()} for r in cur.fetchall()]

    # 6. Incubator stats
    cur.execute("SELECT status, COUNT(*) as cnt FROM incubator_universe GROUP BY status")
    data["incubator"] = [dict(r) for r in cur.fetchall()]

    # 7. Recovery watch
    cur.execute("""
        SELECT symbol, analyst_verdict, analyst_confidence, exit_type,
               explicit_stop_out, relisted_without_stop_out, patience_score
        FROM stopped_out_watch WHERE is_active = true
    """)
    data["recovery"] = [{k: _f(v) for k, v in dict(r).items()} for r in cur.fetchall()]

    # 8. CIO decisions
    cur.execute("""
        SELECT action, COUNT(*) as cnt FROM cio_decisions
        WHERE created_at > %s GROUP BY action ORDER BY cnt DESC
    """, [str(week_start)])
    data["cio_decisions"] = [dict(r) for r in cur.fetchall()]

    # 9. Pipeline health
    cur.execute("""
        SELECT pipeline_key, COUNT(*) as runs,
               COUNT(*) FILTER (WHERE status='completed') as ok,
               COUNT(*) FILTER (WHERE status='failed') as failed
        FROM pipeline_runs WHERE started_at > %s
        GROUP BY pipeline_key ORDER BY runs DESC LIMIT 15
    """, [str(week_start)])
    data["pipeline_health"] = [dict(r) for r in cur.fetchall()]

    # 10. Social ingestion
    cur.execute("""
        SELECT platform, COUNT(*) as posts
        FROM social_posts WHERE ingested_at > %s
        GROUP BY platform
    """, [str(week_start)])
    data["social_ingestion"] = [dict(r) for r in cur.fetchall()]

    # 11. Learning recommendations
    cur.execute("""
        SELECT recommendation_id, title, domain, summary, status
        FROM learning_recommendations WHERE status = 'proposed'
        ORDER BY created_at DESC LIMIT 10
    """)
    data["learning_recs"] = [dict(r) for r in cur.fetchall()]

    # 12. Strategy scores
    cur.execute("""
        SELECT strategy_id, closed_trades, win_rate, profit_factor, recommendation
        FROM strategy_learning_scores ORDER BY created_at DESC LIMIT 10
    """)
    data["strategy_scores"] = [{k: _f(v) for k, v in dict(r).items()} for r in cur.fetchall()]

    return data


def generate_docx(data: dict, output_path: Path):
    """Generate a Word document from the weekly data."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Title
    title = doc.add_heading('Weekly Portfolio Intelligence Report', level=0)
    doc.add_paragraph(f"Period: {data['period_start']} — {data['period_end']}")
    doc.add_paragraph(f"Generated: {data['generated_at']}")
    doc.add_paragraph(f"System: Trade AI v12 — Paper Only Mode")

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    pt = data.get("paper_trades", {})
    wins = pt.get("wins", 0)
    losses = pt.get("losses", 0)
    total_closed = wins + losses
    wr = round(wins / max(total_closed, 1) * 100, 1)
    gp = pt.get("gross_profit", 0)
    gl = pt.get("gross_loss", 0)
    pf = round(gp / max(gl, 0.01), 2) if gl else 0

    summary_lines = [
        f"Portfolio Value: ${data.get('total_value', 0):,.0f} across {data.get('positions', 0)} positions",
        f"Cash: ${data.get('cash', 0):,.0f}",
        f"Paper Trades: {pt.get('open_trades', 0)} open, {total_closed} closed ({wr}% win rate, PF {pf})",
        f"Recovery Watch: {len(data.get('recovery', []))} active items",
    ]
    for line in summary_lines:
        doc.add_paragraph(line, style='List Bullet')

    # Portfolio Health
    doc.add_heading('Portfolio Health', level=1)
    recovery = data.get("recovery", [])
    if recovery:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ['Symbol', 'Verdict', 'Confidence', 'Exit Type', 'Patience']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for r in recovery:
            row = table.add_row()
            row.cells[0].text = str(r.get("symbol", ""))
            row.cells[1].text = str(r.get("analyst_verdict", "")).replace("_", " ").title()
            row.cells[2].text = f"{float(r.get('analyst_confidence', 0)) * 100:.0f}%"
            row.cells[3].text = str(r.get("exit_type", "")).replace("_", " ")
            row.cells[4].text = f"{float(r.get('patience_score', 0)):.2f}"
    else:
        doc.add_paragraph("No active recovery watch items.")

    # Agent Performance
    doc.add_heading('Agent Performance', level=1)
    cal = data.get("agent_calibration", [])
    if cal:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Agent', 'Accuracy', 'Correct', 'Wrong', 'Trend']):
            table.rows[0].cells[i].text = h
        for a in cal:
            row = table.add_row()
            row.cells[0].text = str(a.get("agent_name", ""))
            row.cells[1].text = f"{a.get('accuracy_pct', 0):.0f}%" if a.get("accuracy_pct") else "N/A"
            row.cells[2].text = str(a.get("correct_count", 0))
            row.cells[3].text = str(a.get("wrong_count", 0))
            row.cells[4].text = str(a.get("trending", ""))

    activity = data.get("agent_activity", [])
    if activity:
        doc.add_heading('Agent Activity This Week', level=2)
        for a in activity:
            doc.add_paragraph(
                f"{a.get('agent', 'Unknown')}: {a.get('analyses', 0)} analyses, "
                f"avg confidence {float(a.get('avg_conf', 0)):.0%}",
                style='List Bullet'
            )

    # Proposals
    doc.add_heading('Proposals', level=1)
    props = data.get("proposals", [])
    if props:
        for p in props:
            doc.add_paragraph(f"{p.get('status', 'Unknown')}: {p.get('cnt', 0)}", style='List Bullet')
    else:
        doc.add_paragraph("No proposals created this week.")

    # Strategy Performance
    doc.add_heading('Strategy Performance', level=1)
    strats = data.get("strategy_scores", [])
    if strats:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Strategy', 'Closed', 'Win Rate', 'PF', 'Recommendation']):
            table.rows[0].cells[i].text = h
        for s in strats:
            row = table.add_row()
            row.cells[0].text = str(s.get("strategy_id", ""))
            row.cells[1].text = str(s.get("closed_trades", 0))
            row.cells[2].text = f"{s.get('win_rate', 0):.0%}" if s.get("win_rate") else "N/A"
            row.cells[3].text = f"{s.get('profit_factor', 0):.2f}" if s.get("profit_factor") else "N/A"
            row.cells[4].text = str(s.get("recommendation", "")).replace("_", " ")

    # Incubator
    doc.add_heading('Incubator Universe', level=1)
    for inc in data.get("incubator", []):
        doc.add_paragraph(f"{inc.get('status', 'Unknown')}: {inc.get('cnt', 0)} symbols", style='List Bullet')

    # Pipeline Health
    doc.add_heading('Pipeline Health', level=1)
    ph = data.get("pipeline_health", [])
    if ph:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Pipeline', 'Runs', 'OK', 'Failed']):
            table.rows[0].cells[i].text = h
        for p in ph[:10]:
            row = table.add_row()
            row.cells[0].text = str(p.get("pipeline_key", ""))
            row.cells[1].text = str(p.get("runs", 0))
            row.cells[2].text = str(p.get("ok", 0))
            row.cells[3].text = str(p.get("failed", 0))

    # Social Ingestion
    doc.add_heading('Social Ingestion', level=1)
    for s in data.get("social_ingestion", []):
        doc.add_paragraph(f"{s.get('platform', 'Unknown')}: {s.get('posts', 0)} posts", style='List Bullet')

    # CIO Decisions
    doc.add_heading('CIO Decisions This Week', level=1)
    for d in data.get("cio_decisions", []):
        doc.add_paragraph(f"{d.get('action', 'Unknown')}: {d.get('cnt', 0)}", style='List Bullet')

    # Learning Recommendations
    doc.add_heading('Pending Learning Recommendations', level=1)
    recs = data.get("learning_recs", [])
    if recs:
        for r in recs:
            doc.add_paragraph(f"[{r.get('domain', '')}] {r.get('title', '')}", style='List Bullet')
            if r.get("summary"):
                doc.add_paragraph(r["summary"][:200], style='List Bullet 2')
    else:
        doc.add_paragraph("No pending learning recommendations.")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph("Trade AI v12 — Paper Only Mode — All actions require human review")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Weekly DOCX Report Generator")
    parser.add_argument("--dry-run", action="store_true", help="Print data without generating docx")
    parser.add_argument("--output", help="Custom output path")
    parser.add_argument("--legacy", action="store_true", help="Use legacy table-based DOCX layout")
    parser.add_argument("--format", default="docx", choices=("docx", "pdf", "all"))
    args = parser.parse_args()

    from analyst_report_builder import build_weekly_review, save_report_json
    from report_export import export_report

    report = build_weekly_review()
    today = date.today()
    stem = f"weekly_review_{today.strftime('%Y%m%d')}"

    if args.dry_run:
        print(json.dumps(report, indent=2, default=str))
        return

    week_dir = PROJECT_ROOT / "archive" / "weekly" / str(today) / "reports_weekly"
    week_dir.mkdir(parents=True, exist_ok=True)
    analyst_dir = PROJECT_ROOT / "data" / "portfolios" / "reports" / "analyst"
    analyst_dir.mkdir(parents=True, exist_ok=True)

    save_report_json(report, stem=stem)

    if args.legacy:
        conn = _get_conn()
        try:
            data = gather_weekly_data(conn)
            output = Path(args.output) if args.output else week_dir / f"weekly_{today}.docx"
            path = generate_docx(data, output)
            print(f"[weekly] Legacy DOCX: {path} ({path.stat().st_size / 1024:.1f} KB)")
        finally:
            conn.close()

    formats = ["docx", "pdf"] if args.format == "all" else [args.format]
    for fmt in formats:
        result = export_report(report, fmt, output_stem=stem)
        if result.get("ok"):
            dest = week_dir / result["filename"]
            try:
                import shutil
                shutil.copy2(result["path"], dest)
            except Exception:
                dest = Path(result["path"])
            print(f"[weekly] Analyst {fmt.upper()}: {dest} ({result.get('size_kb')} KB)")
        else:
            print(f"[weekly] Analyst {fmt.upper()} failed: {result.get('error')}", file=sys.stderr)

    try:
        from telegram_alert import send_telegram
        send_telegram(f"*Weekly Portfolio Review*\n{stem} generated\nView Reports → Analyst Reports")
    except Exception:
        pass


if __name__ == "__main__":
    main()

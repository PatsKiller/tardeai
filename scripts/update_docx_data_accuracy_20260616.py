#!/usr/bin/env python3
"""Append the Position-Card & Advisor Data-Accuracy Fixes section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-16; full detail in docs/MASTER_SYSTEM_DOCUMENTATION.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Position-Card & Advisor Data-Accuracy Fixes (2026-06-16)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "An operator review of Trading -> Open Trades surfaced several data-accuracy issues on the position "
        "cards and the regime header. Each is fixed at the source so other surfaces inherit the correction. "
        "Full detail: docs/MASTER_SYSTEM_DOCUMENTATION.md and CHANGELOG.md (2026-06-16).")

    _head(doc, 3, "ETF sector mislabeling")
    doc.add_paragraph(
        "Finviz reports EVERY ETF with sector 'Financial' (industry 'Exchange Traded Fund'), so industrials, "
        "materials, bond, and income funds (XLI, XLB, BND, SCHD, SCHG, JEPI, ARKG) all displayed as "
        "'Financial (XLF)'. An authoritative ETF sector map in open_trades_intelligence.py now takes "
        "precedence over the Finviz-sourced table: sector SPDRs map to their real GICS sector, while broad, "
        "bond, and income funds get an honest asset-class label and no faked 'in-line' vs-sector call when "
        "there is no sector ETF to compare against. The same correction is applied at ingest "
        "(aegis_nightly_ingestion._corrected_sector refuses a bare Finviz 'Financial' on any ETF), and the "
        "664 already-stored rows were backfilled, so the Sectors and Watchlist surfaces do not re-inherit it.")

    _head(doc, 3, "Worthless / delisted equity")
    doc.add_paragraph(
        "A real ticker (not a fund) that has collapsed to about zero with a worse-than -90% unrealized loss "
        "(for example SRNE at $0.0007) was still presenting cached RSI and moving-average technicals as live "
        "signals. Such a position is now flagged worthless, its technicals are nulled and marked stale, and a "
        "'delisted/worthless - verify & write off' warning is surfaced on the card.")

    _head(doc, 3, "Analyst target upside vs live price")
    doc.add_paragraph(
        "The analyst target upside was frozen at analyst-fetch time and went stale on a fast-moving name "
        "(SPCX showed '-14.8% to target' computed off a pre-spike price). It is now recomputed against the "
        "live holdings price in the card enrichment, so the reward:risk framing tracks the current quote.")

    _head(doc, 3, "Regime and VIX coherence")
    doc.add_paragraph(
        "The market regime classifier could declare a 'high_volatility' regime off a gap-volatility proxy "
        "alone while the VIX was calm, which read to the operator as a contradiction ('high volatility 43%' "
        "with VIX about 16). A VIX-coherence guard now dampens the gap-only high-volatility score when the "
        "VIX signal is low or normal, since the VIX is the canonical volatility gauge.")

    _head(doc, 3, "Ask-the-Agents lowercase ticker resolution")
    doc.add_paragraph(
        "The Ask-the-Agents advisor (/api/v2/portfolio/ask) extracted tickers with an uppercase-only pattern, "
        "so a lowercase question ('should i trim xlb for some spcx') matched no symbols and the model replied "
        "that the portfolio held no XLB position. Ticker extraction is now case-insensitive and validates "
        "lowercase tokens against the symbols the operator actually holds or has analyst data for (filtering "
        "common words like 'trim' and 'look'), while an explicitly uppercase token still passes so the "
        "operator can ask about a not-yet-held name. The position context now carries shares, live price, "
        "cost basis, and a per-account breakdown so the model can answer 'how many shares to trim'.")

    _head(doc, 3, "Operator note: restarting the server")
    doc.add_paragraph(
        "The portfolio server runs as a systemd service under user johnclaw with Restart=always. It is "
        "restarted without sudo by killing the process - systemd respawns it with fresh code: "
        "kill $(systemctl show tradeai-portfolio-server.service -p MainPID --value). A 'sudo systemctl "
        "restart' fails silently on the password prompt in a non-interactive session, which can leave stale "
        "code running while data-only fixes (e.g. a DB backfill) make it look partially applied.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Data-accuracy section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Position-Card & Advisor Data-Accuracy Fixes section appended + saved")


if __name__ == "__main__":
    main()

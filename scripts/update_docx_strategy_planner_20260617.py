#!/usr/bin/env python3
"""Append the Interactive Strategy Planner section to the Reference Architecture DOCX.
APPEND-ONLY, idempotent (marker-guarded). Mirrors the established update_docx_* protocol.
As-built 2026-06-17; full detail in docs/MASTER_SYSTEM_DOCUMENTATION.md + CHANGELOG.md."""
from docx import Document

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"
MARKER = "Interactive Strategy Planner (2026-06-17)"


def h(doc, lvl):
    want = f"Heading {lvl}"
    for s in doc.styles:
        if s.name == want:
            return s
    return None


def _has(doc, marker):
    return any(marker in (p.text or "") for p in doc.paragraphs)


def _head(doc, lvl, text):
    p = doc.add_paragraph()
    st = h(doc, lvl)
    if st:
        p.style = st
    p.text = text


def append_section(doc):
    _head(doc, 2, MARKER)
    doc.add_paragraph(
        "An interactive planner that lets the operator declare a portfolio move, see its impact, get a goal-"
        "aligned redeploy plan, and on approval feed the decision back into the learning and discovery loops. "
        "It lives on the Strategy hub under the Planner tab (/v3/strategy). It is advisory and read-only - "
        "approving a plan seeds discovery and records a learning observation, but never places a trade. Full "
        "detail: docs/MASTER_SYSTEM_DOCUMENTATION.md and CHANGELOG.md (2026-06-17).")

    _head(doc, 3, "Stage 1 - Declare")
    doc.add_paragraph(
        "The operator declares an intent: roll an account to cash, trim a holding, deploy new cash, or "
        "rebalance, with an account, an optional symbol, an amount, and a free-text note.")

    _head(doc, 3, "Stage 2 - Impact (what-if, read-only)")
    doc.add_paragraph(
        "The planner computes the consequences without touching anything. The look-through theme delta is "
        "exact, taken from the per-account exposure already computed in lookthrough_themes.json, so rolling "
        "the Fidelity 401k to cash shows precisely which themes lose exposure - for example S&P 500 down about "
        "17 points, Nasdaq 100 down 16, the Magnificent 7 down 14. It also shows the account refactor, the "
        "cash freed and the resulting cash percentage, and the income and goal hit measured against the 55,000 "
        "dollar annual passive-income target.")

    _head(doc, 3, "Stage 3 - Advise (what to invest in)")
    doc.add_paragraph(
        "A goal-aligned redeploy plan is generated through the free LLM lane (Grok with a local gemma "
        "fallback). The advisor is aware of the income gap, the Roth golden-window conversion window, and the "
        "standing defense and aerospace thesis tilt, and it draws redeploy candidates from the Hermes-ranked "
        "watchlist. It returns a sleeve allocation with rough percentages and a handful of concrete tickers or "
        "ETFs, each with a one-line rationale tied to the gaps the move opens.")

    _head(doc, 3, "Stage 4 - Approve and sync")
    doc.add_paragraph(
        "Approving the plan persists it to the strategy_plans table and syncs it in two directions. To "
        "LEARNING, it writes a row to llm_feedback_observations with workflow strategy_plan, so the approved "
        "direction trains the models and its outcome can be scored later. To DISCOVERY, it automatically "
        "creates operator watch_directives for the redeploy targets - tickers and the most underweight theme - "
        "which the discovery engine and the watchlist enrichment sweep then use to source candidates. This "
        "closes the loop from strategy to discovery to watchlist to proposal. The backend is strategy_planner.py "
        "with the endpoints POST /api/v2/strategy/plan and POST /api/v2/strategy/approve; the front end is the "
        "StrategyPlanner component on the Strategy hub.")


def main():
    doc = Document(DOCX_PATH)
    if _has(doc, MARKER):
        print("Strategy Planner section already present, skip")
        return
    append_section(doc)
    doc.save(DOCX_PATH)
    print("Interactive Strategy Planner section appended + saved")


if __name__ == "__main__":
    main()

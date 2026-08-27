#!/usr/bin/env python3
"""Append session updates to Reference Architecture DOCX — May 4, 2026 session C"""
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX = 'docs/project/Trade_AI_v12_Reference_Architecture.docx'
doc = Document(DOCX)

# ── Get heading styles safely ──
h2_style = normal_style = None
for p in doc.paragraphs:
    if p.style:
        if p.style.name == 'Heading 2' and h2_style is None:
            h2_style = p.style
        if p.style.name == 'Normal' and normal_style is None:
            normal_style = p.style
    if h2_style and normal_style:
        break

def add_h2(text):
    p = doc.add_paragraph(text)
    p.style = h2_style
    return p

def add_p(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    return p

def add_table_with_borders(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    border_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        t._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(border_xml))
    return t

# ════════════════════════════════════════════════════════════════
# Continuation of Appendix S — Session C fixes
# ════════════════════════════════════════════════════════════════

add_h2('S.8 Dashboard Data Key Mismatch Fix (change_pct / gap_pct)')
add_p(
    'Bug: The Trade AI command center displayed "%" with no numbers for Change% and Gap% columns. '
    'All 14 tickers showed blank percentages.'
)
add_p(
    'Root cause: score_ticker() returns keys "change_percent" and "gap_percent", but the DB insert '
    'code in trade_ai_orchestrator.py and continuous_runner.py used t.get("change_pct_raw") and '
    't.get("gap_pct_raw") — keys that never existed. Every scan persisted NULL to the DB.'
)
add_p(
    'Fix: Changed both files to use t.get("change_percent") and t.get("gap_percent"). '
    'Also fixed api_v2.py which converted 0 values to empty string via str(0 or "") — '
    'now returns proper float values with None check. Backfilled all 14 tickers from CSV data.'
)
add_p('')

add_h2('S.9 Country Flag Emoji Conversion')
add_p(
    'Bug: html_dashboard.py hardcoded country as "🇺🇸" for all tickers. The API returned raw text '
    '("United States", "Israel") from the DB, which the command center displayed as plain text.'
)
add_p(
    'Fix: Added _country_flag() converter with 30+ country-to-emoji mappings to both '
    'html_dashboard.py and api_v2.py. Now correctly shows 🇮🇱 for Israel (DRTS), '
    '🇲🇾 for Malaysia (FGL), 🇺🇸 for US tickers, and 🌐 for unmapped countries.'
)
add_p('')

add_h2('S.10 Portfolio Server Process Guardrails')
add_p(
    'Problem: Stray portfolio_server.py processes kept spawning on reboot, causing duplicate '
    'port 7777 bindings. Had to manually kill strays 3 times during this session.'
)
add_p('Two-layer guard added:')
add_table_with_borders(
    ['Layer', 'Mechanism', 'How'],
    [
        ['systemd', 'ExecStartPre', 'fuser -k 7777/tcp kills any process on the port before server starts'],
        ['Python', 'Startup guard', 'portfolio_server.py checks fuser at startup, kills strays before binding'],
    ]
)
add_p('')

add_h2('S.11 CSV/JSON File Cleanup (Audit Results Applied)')
add_p('Acted on the CSV/JSON audit findings. Files cleaned:')
add_table_with_borders(
    ['Category', 'Action', 'Count', 'Details'],
    [
        ['Orphan state files', 'Moved to archive/', '3', 'dec31_anchor.json, agent_context_manifest.json, portfolio_options.json'],
        ['Old broker CSVs', 'Moved to archive/', '4', 'Position CSVs + trades_sample.csv (Transaction CSVs kept)'],
        ['Ingestion logs', 'Deleted', '74', 'data/logs/ingestion_summary_*.json — never read'],
        ['Old catalyst caches', 'Deleted', '19', 'data/catalyst_cache_2026-04-*.json — only current day used'],
        ['Raw snapshots', 'Pruned >7d', '363', 'data/portfolios/state/raw_snapshots/ — 182 remain'],
        ['Snapshot history', 'Pruned >7d', '363', 'data/portfolios/state/ticker_snapshot_history/ — 182 remain'],
    ]
)
add_p(
    'Total: ~826 files removed. Auto-pruning added to db_retention.py — '
    'runs weekly with the DB retention timer. Prune policies: '
    'raw_snapshots 14d, ticker_snapshot_history 14d, ingestion logs 7d, catalyst caches 3d.'
)
add_p('')

add_h2('S.12 Data Sync Gap — Duplicate ticker_enrichment_cache.json')
add_p(
    'Audit discovered two separate ticker_enrichment_cache.json files in different directories: '
    'data/state/ (written by finviz_enrichment.py) and data/portfolios/state/ '
    '(written by ticker_snapshot_builder.py). Different inodes, different sizes (360KB vs 344KB). '
    'The portfolio pipeline reads from data/portfolios/state/ but the enrichment pipeline writes '
    'to data/state/ — creating a data-sync gap. Flagged for investigation.'
)

# ── Save ──
doc.save(DOCX)
print(f"Saved. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

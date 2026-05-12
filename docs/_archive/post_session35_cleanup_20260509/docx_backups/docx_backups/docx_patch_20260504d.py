#!/usr/bin/env python3
"""Append Hardware Preparation Map to Reference Architecture DOCX — May 4, 2026"""
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX = 'docs/project/Trade_AI_v12_Reference_Architecture.docx'
doc = Document(DOCX)

# ── Get heading styles safely ──
h1_style = h2_style = normal_style = None
for p in doc.paragraphs:
    if p.style:
        if p.style.name == 'Heading 1' and h1_style is None:
            h1_style = p.style
        if p.style.name == 'Heading 2' and h2_style is None:
            h2_style = p.style
        if p.style.name == 'Normal' and normal_style is None:
            normal_style = p.style
    if h1_style and h2_style and normal_style:
        break

def add_h1(text):
    p = doc.add_paragraph(text)
    p.style = h1_style
    return p

def add_h2(text):
    p = doc.add_paragraph(text)
    p.style = h2_style
    return p

def add_p(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    return p

def add_bold_p(label, text):
    p = doc.add_paragraph()
    if normal_style:
        p.style = normal_style
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    for run in p.runs:
        run.font.name = 'Courier New'
        from docx.shared import Pt
        run.font.size = Pt(9)
    return p

def add_table_with_borders(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    border_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        t._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(border_xml))
    return t


# ════════════════════════════════════════════════════════════════
# Appendix T — Hardware Preparation Map for Intel Arc Pro B50 GPU
# ════════════════════════════════════════════════════════════════

add_h1('Appendix T — Hardware Preparation Map (Intel Arc Pro B50 GPU)')
add_p(
    'Hardware preparation steps performed on MS-01 (Ubuntu 24.04) in May 2026 '
    'to prepare for the Intel Arc Pro B50 GPU installation and fan management.'
)
add_p('')

# ── T.1 BIOS Update ──
add_h2('T.1 BIOS Update (Version 1.27)')
add_bold_p('Method: ', 'Pure UEFI Shell (no Windows PE)')
add_p('')
add_p('Steps performed:')
add_p('1. Downloaded MS-01-AHWSA-V1.27_4_28_V2.zip')
add_p('2. Formatted USB as FAT32: mkfs.fat -F 32 -n "MS01-BIOS" /dev/sda1')
add_p('3. Copied all files from the zip directly to the root of the USB')
add_p('4. Booted MS-01 > F7 > selected USB > chose UEFI Shell')
add_p('5. In Shell: fs0: (or fs1:/fs2:) > ls > AfuEfiFlash.nsh')
add_p('')
add_bold_p('Warnings: ', 'Disabled Secure Boot before flashing. Never cut power or removed USB during flash.')
add_p('')

# ── T.2 BIOS Settings ──
add_h2('T.2 Key BIOS Settings Changed for Arc B50 + Local LLMs')
add_table_with_borders(
    ['Setting', 'Value Set', 'Reason'],
    [
        ['PCIe Speed', 'Gen4', 'Stability with Arc B50'],
        ['Above 4G Decoding', 'Enabled', 'Required for large VRAM GPUs'],
        ['Re-Size BAR Support', 'Enabled', 'Improves GPU performance'],
        ['Primary Display', 'PEG', 'Use discrete GPU as primary'],
        ['CSM Support', 'Disabled', 'Pure UEFI mode'],
        ['Secure Boot', 'Disabled', 'Easier driver installation'],
    ]
)
add_p('')

# ── T.3 GPU Packages ──
add_h2('T.3 Packages Installed for GPU Management & Monitoring')
add_code(
    'sudo apt -o Acquire::ForceIPv4=true install -y \\\n'
    '    lm-sensors intel-gpu-tools \\\n'
    '    mesa-vulkan-drivers mesa-va-drivers \\\n'
    '    vainfo clinfo'
)
add_p('')
add_table_with_borders(
    ['Package', 'Purpose'],
    [
        ['lm-sensors', 'Temperature and sensor monitoring'],
        ['intel-gpu-tools', 'GPU monitoring (intel_gpu_top)'],
        ['mesa-vulkan-drivers', 'Vulkan support for GPU acceleration'],
        ['mesa-va-drivers', 'VA-API video acceleration'],
        ['vainfo', 'Verify GPU hardware video acceleration'],
        ['clinfo', 'Verify OpenCL GPU compute support'],
    ]
)
add_p('')

# ── T.4 Fan Control ──
add_h2('T.4 Fan Control Setup (NCT6798D Super IO)')
add_p('Commands executed:')
add_code(
    'sudo modprobe nct6775\n'
    'sudo /usr/sbin/pwmconfig\n'
    'sudo systemctl enable --now fancontrol'
)
add_p('')
add_p('Final /etc/fancontrol configuration:')
add_table_with_borders(
    ['Parameter', 'Value'],
    [
        ['Fan', 'hwmon4/pwm2 (fan2)'],
        ['Temp sensor', 'hwmon4/temp1_input (CPU)'],
        ['MINTEMP', '40 C'],
        ['MAXTEMP', '70 C'],
        ['MINSTOP', '25'],
    ]
)
add_p('')

# ── T.5 GUI Monitoring ──
add_h2('T.5 GUI Monitoring Tool')
add_code('sudo snap install mission-center')
add_p('')

# ── T.6 Network Workaround ──
add_h2('T.6 Network Workaround (Ubuntu 25.10 Mirror Issues)')
add_code('sudo apt -o Acquire::ForceIPv4=true update')
add_p('')

# ── T.7 Next Steps ──
add_h2('T.7 Next Steps (After GPU Physical Installation)')
add_p('1. Verify GPU detection:')
add_code(
    'lspci | grep -E "VGA|3D|Display"\n'
    'intel-gpu-top'
)
add_p('')
add_p('2. Install full Intel oneAPI drivers:')
add_code(
    'wget -O- https://apt.repos.intel.com/intel-gpg-keys/GPG-PUB-KEY-INTEL-SW-PRODUCTS.PUB \\\n'
    '  | gpg --dearmor | sudo tee /usr/share/keyrings/oneapi-archive-keyring.gpg > /dev/null\n'
    'echo "deb [signed-by=/usr/share/keyrings/oneapi-archive-keyring.gpg] \\\n'
    '  https://apt.repos.intel.com/oneapi all main" \\\n'
    '  | sudo tee /etc/apt/sources.list.d/oneAPI.list\n'
    'sudo apt update\n'
    'sudo apt install -y intel-basekit'
)
add_p('')
add_p('3. Configure Ollama for Intel Arc GPU acceleration (qwen3:14b target)')

# ── Save ──
doc.save(DOCX)
print(f"Saved. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

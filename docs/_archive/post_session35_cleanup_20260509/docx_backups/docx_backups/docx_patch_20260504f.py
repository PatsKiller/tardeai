#!/usr/bin/env python3
"""Append scalping WebSocket + OpenClaw skill section to Reference Architecture DOCX — May 4, 2026 session F"""
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX = 'docs/project/Trade_AI_v12_Reference_Architecture.docx'
doc = Document(DOCX)

h1_style = h2_style = normal_style = None
for p in doc.paragraphs:
    if p.style:
        if p.style.name == 'Heading 1' and h1_style is None:
            h1_style = p.style
        if p.style.name == 'Heading 2' and h2_style is None:
            h2_style = p.style
        if p.style.name == 'Normal' and normal_style is None:
            normal_style = p.style
    if h1_style and h2_style and normal_style:
        break


def add_h1(text):
    p = doc.add_paragraph(text); p.style = h1_style; return p


def add_h2(text):
    p = doc.add_paragraph(text); p.style = h2_style; return p


def add_p(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    return p


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for rd in rows:
        row = t.add_row()
        for i, val in enumerate(rd):
            row.cells[i].text = str(val)
    border_xml = (f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>')
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); t._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(border_xml))
    return t


# ════════════════════════════════════════════════════════════════
# New Section: Scalping WebSocket + OpenClaw Skill
# ════════════════════════════════════════════════════════════════

add_h1('Appendix U — Scalping WebSocket & OpenClaw Skill (v1.8)')

add_h2('U.1 Architecture Overview')
add_p(
    'Real-time scalping signal broadcast system. A standalone async WebSocket server '
    '(scalp_ws_server.py) runs alongside portfolio_server.py. Frontend clients connect on '
    'port 7778; pipeline scripts push signals via an internal port (7779). This avoids modifying '
    'the stdlib http.server stack and keeps WebSocket concerns isolated.'
)
add_p('')

add_h2('U.2 Server Components')
add_table(
    ['File', 'Port', 'Purpose'],
    [
        ['scripts/scalp_ws_server.py', '7778 (clients), 7779 (push)', 'Async WS server using websockets lib'],
        ['scripts/scalp_ws_client.py', 'connects to 7779', 'Helper module: broadcast_scalp_update()'],
        ['linux_launchers/run_scalp_ws.sh', 'N/A', 'Launch/restart script'],
    ]
)
add_p('')

add_h2('U.3 Signal Payload Schema')
add_p('All signals pushed through the WebSocket have this JSON shape:')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['symbol', 'string', 'Ticker symbol (e.g. AAPL)'],
        ['grade', 'string', 'Scoring grade (A+, A, B, C, D)'],
        ['score', 'int', 'Score out of 55'],
        ['decision', 'string', 'GO / WAIT / AVOID / APPROVED'],
        ['change_percent', 'string', 'Intraday change %'],
        ['rvol', 'float', 'Relative volume multiplier'],
        ['critic_verdict', 'string', 'CONFIRM / DOWNGRADE / BLOCK'],
        ['catalyst_verified', 'bool', 'Whether catalyst passed validation'],
        ['source', 'string', 'Origin: screener / social / continuous / news_monitor / openclaw_approve'],
    ]
)
add_p('')

add_h2('U.4 Pipeline Integration Points')
add_p('broadcast_scalp_update() is called (non-fatal, try/except wrapped) from:')
add_table(
    ['Script', 'Trigger', 'Signals Sent'],
    [
        ['social_scalp_scanner.py', 'After score_ticker() + save_scan_result()', 'All scored tickers (GO/WAIT/AVOID)'],
        ['trade_ai_orchestrator.py', 'After scalp critic step (10a)', 'GO + WAIT tickers with critic verdicts'],
        ['trade_ai_news_monitor.py', 'On verdict change detection', 'Verdict-changed tickers only'],
        ['continuous_runner.py', 'On NEW_GO trigger in live alerts', 'Real-time GO detections'],
    ]
)
add_p('')

add_h2('U.5 Frontend Component')
add_p(
    'ScalpLiveFeed.tsx in apps/command-center-v2/src/components/. Connects to ws://<host>:7778, '
    'auto-reconnects on disconnect (5s delay), sends keepalive pings every 25s. Shows last 30 signals '
    'with decision Badge, score bar, RVOL, critic verdict. Fires toast on GO signals via existing '
    'ToastProvider context. Integrated into TradeAI.tsx page between metric tiles and filter bar.'
)
add_p('')

add_h2('U.6 OpenClaw Skill: scalp-signal-approve')
add_p(
    'Located at ~/.openclaw/skills/scalp-signal-approve/. Provides CLI and conversational interface '
    'for approving/rejecting scalp signals from Telegram or chat.'
)
add_table(
    ['Command', 'Action', 'DB Effect'],
    [
        ['approve scalp <symbol> [comment]', 'Approve signal for trading', 'INSERT/UPDATE watchlist_proposals (status=approved)'],
        ['reject scalp <symbol> [reason]', 'Block signal', 'UPDATE watchlist_proposals (status=rejected)'],
        ['pending scalps', 'List unapproved signals', 'SELECT last 24h SCALP proposals'],
        ['escalate scalp <symbol>', 'Send to Alex for review', 'INSERT with alex_vote=pending'],
    ]
)
add_p('')
add_p(
    'Boundaries: approval only (no trade execution), sends Telegram confirmation, respects '
    'holdings guard (non-fatal check), broadcasts approval to WS feed. Position sizing defers '
    'to Steph allocation rules ($150 risk per scalp trade).'
)
add_p('')

add_h2('U.7 Deployment')
add_p(
    'Start: ./linux_launchers/run_scalp_ws.sh (kills existing, starts nohup). '
    'Logs: logs/scalp_ws_server.log. '
    'Dependencies: websockets==16.0 (already in requirements.txt). '
    'No new pip installs required.'
)
add_p('')

# Save
import shutil
shutil.copy(DOCX, DOCX + '.bak_20260504f')
doc.save(DOCX)
print('DOCX updated — Appendix U (Scalping WebSocket + OpenClaw Skill) appended')
print(f'Backup: {DOCX}.bak_20260504f')

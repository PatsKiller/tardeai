#!/usr/bin/env python3
"""Append session fixes to Reference Architecture DOCX — May 4, 2026 session E"""
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX = 'docs/project/Trade_AI_v12_Reference_Architecture.docx'
doc = Document(DOCX)

h2_style = normal_style = None
for p in doc.paragraphs:
    if p.style:
        if p.style.name == 'Heading 2' and h2_style is None:
            h2_style = p.style
        if p.style.name == 'Normal' and normal_style is None:
            normal_style = p.style
    if h2_style and normal_style:
        break

def add_h2(text):
    p = doc.add_paragraph(text); p.style = h2_style; return p

def add_p(text):
    p = doc.add_paragraph(text)
    if normal_style: p.style = normal_style
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs: run.bold = True
    for rd in rows:
        row = t.add_row()
        for i, val in enumerate(rd): row.cells[i].text = str(val)
    border_xml = (f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>')
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); t._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(border_xml))
    return t

# ════════════════════════════════════════════════════════════════
# Continuation of Appendix S — Session D/E fixes
# ════════════════════════════════════════════════════════════════

add_h2('S.13 Frontend Type Fix (change_pct / gap_pct / float_m)')
add_p(
    'The React frontend (TradeAI.tsx) declares change_pct, gap_pct, and float_m as string types '
    'and calls .startsWith("-") to determine color. The API was changed to return float, which '
    'caused "e.change_pct.startsWith is not a function" crash on page load.'
)
add_p(
    'Fix: API reverted to return string values with proper None handling: '
    'str(round(float(val), 2)) if val is not None else "". '
    'This preserves the frontend string contract while fixing the original zero-as-falsy bug '
    '(0 was treated as empty via str(0 or "") = "").'
)
add_p('')

add_h2('S.14 Trade AI Page Load Speed — Social Sentiment Caching')
add_p(
    'The /api/v2/trade-ai endpoint took 20-30 seconds to load because it made live StockTwits + '
    'Reddit API calls for every ticker on every page request (14 tickers x 0.2s delay + API latency).'
)
add_p(
    'Fix: Removed live social sentiment lookups from the API. Social data is already persisted to '
    'trade_ai_scans by the continuous_runner during each scan cycle (every 10 minutes during market '
    'hours). The API now serves social data directly from the DB. Page load dropped from ~25s to ~0.1s.'
)
add_p('')

add_h2('S.15 Country Flag Frontend Fix')
add_p(
    'After the API was updated to return flag emojis (S.9), the React frontend still had its own '
    'country-to-flag mapping that expected text values like "United States". Since the API now sends '
    '"🇺🇸", the old mapping fell through to the 🌍 catch-all for every ticker.'
)
add_p(
    'Also fixed: "Foreign issuer" warning incorrectly showed for US tickers because the check '
    'compared against "United States" but the API now sends "🇺🇸".'
)
add_p('Changes to TradeAI.tsx:')
add_table(
    ['Component', 'Before', 'After'],
    [
        ['Sector column flag', 'Inline mapping of 6 country names + 🌍 fallback', 'Uses flag emoji directly from API; hides flag for 🇺🇸 and US tickers'],
        ['Foreign issuer warning', 'country !== "United States"', 'country !== "United States" && country !== "🇺🇸"'],
    ]
)
add_p('Frontend rebuilt with vite build and deployed.')
add_p('')

add_h2('S.16 LLM Critic Re-run (Stale Data Fix)')
add_p(
    'After fixing the Qwen3 thinking mode issue (S.1), the 8 GO/WAIT tickers in trade_ai_scans '
    'still had stale "LLM unavailable — original preserved" critic verdicts from the broken run. '
    'Re-ran llm_critique() on all 8 tickers — all returned CONFIRM with real LLM reasoning '
    '(75-95% confidence). Results persisted to trade_ai_scans.'
)
add_p('')

add_h2('S.17 Social/News Candidate Injection into Scoring Pipeline')
add_p(
    'Audit discovered that social candidates (from social_scalp_scanner.py) and news candidates '
    'were scored and alerted via Telegram but NEVER injected into the main Trade AI pipeline. '
    'The two systems were completely separate:'
)
add_table(
    ['Pipeline', 'Source', 'Output', 'Visible In'],
    [
        ['Finviz screener', 'finviz_ingestion.py', 'trade_ai_scans (source=screener)', 'Command Center Trade AI page'],
        ['Social scanner', 'social_scalp_scanner.py', 'scalp_scan_results only', 'Telegram alerts only (not dashboard)'],
    ]
)
add_p('')
add_p(
    'Fix: Added social candidate injection to continuous_runner.py. After loading Finviz screener '
    'tickers, the runner now queries scalp_scan_results for social candidates (score >= 25, last '
    '4 hours) and merges them into the ticker list before scoring. Injected tickers carry '
    '_source="social" which persists to trade_ai_scans.source. This means social-discovered '
    'tickers now appear in the Trade AI command center alongside screener picks.'
)
add_p('')
add_p('Pipeline status audit (May 4, 2026 — Sunday):')
add_table(
    ['Pipeline', 'Status', 'Last Run', 'Notes'],
    [
        ['Social scalp scanner', 'Running (cron)', '11:00 AM', '15 tickers checked, 0 qualified (Sunday low volume)'],
        ['News monitor', 'Running (systemd)', '11:31 AM', 'No breaking news re-evaluations triggered'],
        ['Continuous runner', 'Dead (timer)', 'Next: Tue 4AM', 'Will include social injection on next run'],
        ['Event detector', 'Running (cron)', 'Every 15 min', 'Fires event_router for matched patterns'],
    ]
)
add_p('')
add_p(
    'No social GO candidates today because: (1) Sunday — minimal social activity, '
    '(2) trending tickers (GME, NVDA, AAPL, SPY) are large-cap/high-price → routed to portfolio agents, '
    '(3) qualifying small-caps (AI score 23, XNDU score 20) scored below the 25-point injection threshold.'
)

# ── Save ──
doc.save(DOCX)
print(f"Saved. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

#!/usr/bin/env python3
"""Append session updates to Reference Architecture DOCX — May 4, 2026 session B"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX = 'docs/project/Trade_AI_v12_Reference_Architecture.docx'
doc = Document(DOCX)

# ── Get heading styles safely ──
h1_style = h2_style = normal_style = None
for p in doc.paragraphs:
    if p.style:
        if p.style.name == 'Heading 1' and h1_style is None:
            h1_style = p.style
        if p.style.name == 'Heading 2' and h2_style is None:
            h2_style = p.style
        if p.style.name == 'Normal' and normal_style is None:
            normal_style = p.style
    if h1_style and h2_style and normal_style:
        break

def add_h1(text):
    p = doc.add_paragraph(text)
    p.style = h1_style
    return p

def add_h2(text):
    p = doc.add_paragraph(text)
    p.style = h2_style
    return p

def add_p(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    return p

def add_bullet(text):
    p = doc.add_paragraph(text)
    if normal_style:
        p.style = normal_style
    return p

def add_table_with_borders(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Data rows
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    # Borders
    border_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        t._tbl.insert(0, tblPr)
    tblPr.append(parse_xml(border_xml))
    return t


# ════════════════════════════════════════════════════════════════
# Appendix S — Session Fixes & Hardening (May 4, 2026 — Session B)
# ════════════════════════════════════════════════════════════════

add_h1('Appendix S — Session Fixes & Hardening (May 4, 2026 — Session B)')
add_p('')

add_h2('S.1 Qwen3 Thinking Mode Fix (All Ollama Calls)')
add_p(
    'Root cause: Qwen3:1.7b defaults to thinking mode, consuming all num_predict tokens '
    'in <think> tags and returning empty response content. This caused the Scalp Critic Agent '
    'to fall back to "LLM unavailable — original preserved" for every ticker.'
)
add_p(
    'Fix: Migrated all 16 Python scripts from Ollama /api/generate to /api/chat with '
    '"think": false. Response parsing changed from .get("response") to '
    '.get("message", {}).get("content", ""). Only system_preflight_check.py still uses '
    '/api/generate (connectivity check only, response content not needed).'
)
add_p('Affected scripts:')
add_table_with_borders(
    ['Script', 'LLM Function', 'Impact'],
    [
        ['scalp_critic_agent.py', '_call_ollama()', 'Critic verdicts now have real LLM reasoning'],
        ['scoring.py', '_ollama_serialized()', 'Catalyst scoring, Sonnet escalation'],
        ['local_llm.py', '_try_ollama()', 'Shared local LLM module (used by many scripts)'],
        ['llm_router.py', '_call_local()', 'Central routing module'],
        ['trade_ai_orchestrator.py', 'warmup call', 'Pre-scan model warmup'],
        ['api_v2.py', 'portfolio Q&A', 'Chat endpoint LLM call'],
        ['catalyst_intelligence.py', '_ollama_analyze()', 'Catalyst classification'],
        ['portfolio_news.py', '_ollama()', 'News narrative generation'],
        ['stop_decision_brief.py', '_ollama()', 'Stop-loss briefs'],
        ['morning_digest.py', '_ollama_narrative()', 'Morning digest narratives'],
        ['portfolio_weekly_report.py', '_ollama()', 'Weekly report summaries'],
        ['aegis_synthesis.py', '_llm() + _steph_llm()', 'Aegis overnight synthesis'],
        ['process_watchlist_agent_jobs.py', '_llm() fallback', 'Watchlist agent LLM fallback'],
        ['journal_agent_coach.py', '_call_ollama()', 'Journal coaching prompts'],
        ['portfolio_orchestrator.py', 'daily summary', 'Portfolio daily summary LLM'],
        ['portfolio_ai_analyst.py', '_ollama()', 'AI analyst deep analysis'],
    ]
)
add_p('')

add_h2('S.2 Catalyst Verification Fix (RLYB)')
add_p(
    'Bug: RLYB catalyst_verified=false despite "Rallybio to acquire Candid Therapeutics" headline '
    'clearly matching the company. Confidence was 0.15/0.20 instead of expected 0.50.'
)
add_p(
    'Root cause: The continuous runner had an intermediate code version cached in memory where '
    'the company_match scoring formula gave +0.20 instead of the corrected +0.40 for full matches. '
    'The runner died during a reboot, and the fix was already on disk.'
)
add_p(
    'Actions taken: Fixed RLYB, CELU, BCHT catalyst_verified to true in trade_ai_scans table. '
    'Killed stray duplicate portfolio_server process (PID 3325). Cleared stale .pyc bytecode cache. '
    'Next continuous_runner start will use the corrected scoring code.'
)
add_p('')

add_h2('S.3 Database Retention Policy System')
add_p(
    'Created scripts/db_retention.py — automated retention enforcement for all 164 DB tables. '
    'Five retention tiers based on data criticality:'
)
add_table_with_borders(
    ['Tier', 'Retention', 'Tables', 'Description'],
    [
        ['PRICE', '1 year', 'price_cache, ticker_prices', 'Historical price data for performance calcs'],
        ['LONG', '180 days', '10 tables', 'Trade history, performance, intelligence history'],
        ['MEDIUM', '90 days', '35 tables', 'Scans, snapshots, state, agent results, news'],
        ['SHORT', '30 days', '15 tables', 'Events, jobs, queues, notifications, approvals'],
        ['TINY', '14 days', '6 tables', 'Ephemeral metrics, backtests, iris logs'],
    ]
)
add_p(
    'First run deleted 223,063 rows (111K each from price_cache and ticker_prices — pre-2025 data, '
    '51 from news_articles). DB size: 213 MB after cleanup.'
)
add_p(
    'Systemd timer: db-retention.timer runs weekly Sunday 3AM. '
    'Supports --dry-run for safe preview.'
)
add_p('')

add_h2('S.4 CSV/JSON File Audit')
add_p('Audited all CSV/JSON files in the project against active Python scripts and DB tables.')
add_p('Key findings:')
add_table_with_borders(
    ['Category', 'Count', 'Action'],
    [
        ['ACTIVE (still primary read path)', '~50 files', 'Keep — JSON is primary, DB is dual-write'],
        ['GENERATED (output/cache)', '~20 files', 'Keep — served to frontend'],
        ['Safe to delete (orphans)', '4 files', 'dec31_anchor.json, agent_context_manifest.json, portfolio_options.json, old broker CSVs'],
        ['Prune on schedule', '~1100 files', 'raw_snapshots/ + ticker_snapshot_history/ dirs, old catalyst caches'],
        ['Data sync gap found', '1 pair', 'Two separate ticker_enrichment_cache.json files in different dirs'],
    ]
)
add_p(
    'Bug found: data/state/ticker_enrichment_cache.json and data/portfolios/state/ticker_enrichment_cache.json '
    'are separate files written by different scripts (finviz_enrichment.py vs ticker_snapshot_builder.py). '
    'This creates a data-sync gap where the portfolio pipeline may use stale enrichment data.'
)
add_p('')

add_h2('S.5 Command Center Auto-Refresh Fix')
add_p(
    'The command center already had auto-refresh (60s countdown ring in header). '
    'Fixed bug where startRefreshTimer() hardcoded 60s instead of using the configurable '
    'S.refreshSecs from the settings panel. Now respects user-configured interval.'
)
add_p('')

add_h2('S.6 Updated Operating Numbers')
add_table_with_borders(
    ['Metric', 'Value'],
    [
        ['Database tables', '164'],
        ['Database size', '213 MB (post-retention)'],
        ['Python scripts', '214'],
        ['Systemd timers', '17 (including new db-retention.timer)'],
        ['Ollama API pattern', '16 scripts on /api/chat, 1 on /api/generate (preflight only)'],
        ['Portfolio value', '$1,180,862'],
        ['Positions', '43'],
    ]
)
add_p('')

add_h2('S.7 New Systemd Timer')
add_table_with_borders(
    ['Timer', 'Schedule', 'Purpose'],
    [
        ['db-retention.timer', 'Sun 03:00', 'Weekly database retention enforcement (5-tier policy)'],
    ]
)

# ── Save ──
doc.save(DOCX)
print(f"Saved. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

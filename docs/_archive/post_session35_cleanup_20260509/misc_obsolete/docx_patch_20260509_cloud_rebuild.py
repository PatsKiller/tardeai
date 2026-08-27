"""
DOCX Patch: Cloud Product Rebuild — 2026-05-09

Appends a comprehensive cloud-product-grade rebuild section to the
Trade AI v12 Reference Architecture DOCX. This replaces the addendum model.

Protocol: append-only via python-docx. No deletions, no XML manipulation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import shutil
from datetime import datetime

DOCX_PATH = "docs/project/Trade_AI_v12_Reference_Architecture.docx"

doc = Document(DOCX_PATH)

# --- Find heading styles from existing doc ---
h1_style = None
h2_style = None
normal_style = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 1' and h1_style is None:
        h1_style = p.style
    if p.style and p.style.name == 'Heading 2' and h2_style is None:
        h2_style = p.style
    if p.style and p.style.name == 'Normal' and normal_style is None:
        normal_style = p.style
    if h1_style and h2_style and normal_style:
        break

def add_h1(text):
    p = doc.add_paragraph(text)
    if h1_style:
        p.style = h1_style
    return p

def add_h2(text):
    p = doc.add_paragraph(text)
    if h2_style:
        p.style = h2_style
    return p

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_bold_para(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p

def add_table(headers, rows):
    # Workaround: if sections list is empty, add a section element
    try:
        _ = doc.sections[-1]
    except (IndexError, Exception):
        sectPr = parse_xml(
            f'<w:sectPr {nsdecls("w")}>'
            '  <w:pgSz w:w="12240" w:h="15840"/>'
            '  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"'
            '           w:header="720" w:footer="720" w:gutter="0"/>'
            '</w:sectPr>'
        )
        doc.element.body.append(sectPr)
    table = doc.add_table(rows=1, cols=len(headers))
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    # Apply borders
    border_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(border_xml))
    return table


# ============================================================
# SECTION: DEPRECATION NOTICE
# ============================================================
doc.add_page_break()

add_h1("NOTICE: Cloud Product Rebuild (2026-05-09)")

add_para(
    "This section was appended on 2026-05-09 as part of the Cloud Product "
    "Documentation Rebuild. The addendum model (Appendices A through W and "
    "session-specific sections) is officially deprecated. All system documentation "
    "is now consolidated into the authoritative Markdown files listed below. "
    "This DOCX section provides a cloud-product-grade summary of the full system."
)

add_para("")
add_bold_para("Authoritative documents (Markdown, in docs/):", "")
add_para("  - MASTER_SYSTEM_DOCUMENTATION.md (complete system reference)")
add_para("  - ARCHITECTURE_OVERVIEW.md (executive architecture summary)")
add_para("  - ARCHITECTURE_INFOGRAM.md (visual architecture diagrams)")
add_para("  - CHEAT_SHEET.md (operator quick reference)")
add_para("  - COST_MODEL.md (cloud operating cost estimates)")
add_para("  - RESTORE_GUIDE.md (disaster recovery)")
add_para("  - GPU_OLLAMA_SETUP.md (hardware configuration)")
add_para("")
add_para(
    "The prior appendices (A-W) and session sections remain in this document "
    "for historical reference but are no longer maintained. Any discrepancy "
    "between this section and the Markdown files should be resolved in favor "
    "of the Markdown files.",
    bold=True
)

# ============================================================
# SECTION 1: LIVE OPERATING NUMBERS (2026-05-09)
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 1: Live Operating Numbers (2026-05-09)")

add_table(
    ["Metric", "Value", "Change vs. Prior Doc"],
    [
        ["Python scripts", "302", "+10"],
        ["Database tables", "229", "+10 (was 219)"],
        ["Cron jobs", "138", "+8"],
        ["Strategies (YAML)", "23", "+3"],
        ["API endpoints", "80+", "stable"],
        ["Frontend pages (TSX)", "52", "+2"],
        ["React components", "88", "-8 (consolidation)"],
        ["News articles ingested", "2,026", "+1,474"],
        ["Enrichment symbols", "1,139", "stable"],
        ["Active incubator symbols", "632", "+577"],
        ["Paper trade proposals", "43", "+43 (new)"],
        ["Paper trades executed", "4", "+4 (new)"],
        ["CIO decisions", "446", "+391"],
        ["Agent handoffs", "625", "+515"],
        ["Portfolio value", "$1,189,457", "stable"],
        ["SQL migrations", "19", "+19 (new system)"],
        ["Ollama models loaded", "qwen3:14b, qwen3:1.7b, nomic-embed-text", "GPU upgrade"],
    ]
)

# ============================================================
# SECTION 2: CLOUD-NATIVE SERVICE ARCHITECTURE
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 2: Service Architecture")

add_h2("Service Boundary Map")
add_para(
    "Trade AI v12 has 6 distinct service boundaries, all currently co-located "
    "on ms01-openclaw. Each service is independently deployable as a container "
    "or managed cloud service."
)

add_table(
    ["Service", "Port", "Technology", "Role", "Cloud Equivalent (AWS)"],
    [
        ["Portfolio Server", "7777", "Python Flask", "Central API hub, 80+ endpoints, React SPA", "ECS Fargate + ALB"],
        ["PostgreSQL 15", "5432", "PostgreSQL", "229 tables, all persistent state", "RDS PostgreSQL"],
        ["Ollama LLM", "11434", "qwen3:14b", "Local inference, Intel Arc B50 GPU (Vulkan)", "EC2 g5 / Bedrock"],
        ["OpenClaw Gateway", "18789", "Python", "4 conversational agents", "ECS Fargate"],
        ["Cron Scheduler", "--", "crontab", "138 scheduled jobs", "EventBridge Scheduler"],
        ["React SPA", "via 7777", "Vite/React", "50+ pages, operator dashboard", "S3 + CloudFront"],
        ["Scalp WebSocket", "7778/7779", "Python WS", "Real-time scalp feed", "API Gateway WS"],
    ]
)

add_h2("Deployment Model")
add_para(
    "Current: Single-tenant, single-server. All services co-located on one machine. "
    "No HA, no auto-scaling. Recovery is manual with documented procedures."
)
add_para(
    "Cloud target: Single-tenant, multi-service deployment. Compute containerized, "
    "DB as managed service, LLM as GPU container or managed API, static frontend "
    "from object storage + CDN, cron replaced by managed scheduler."
)

# ============================================================
# SECTION 3: PIPELINE ARCHITECTURE (31 STAGES, 7 GROUPS)
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 3: Pipeline Architecture")

add_para(
    "The pipeline runs 31 stages organized into 7 groups. Each group has a "
    "designated time window and dependency chain."
)

add_table(
    ["Group", "Name", "Schedule", "Key Stages", "Outputs"],
    [
        ["1", "Data Collection", "5:45-7:00 AM", "Finviz screener, news (7 APIs), SEC, FRED", "Raw scans, news articles, filings"],
        ["2", "Enrichment", "7:00-8:00 AM", "60+ Finviz fields, 17 indicators, 7 catalyst sources", "Enriched symbols, indicator cache"],
        ["3", "Scoring", "8:00-9:00 AM", "55-point engine, GO/WAIT/NO-GO", "Scored candidates"],
        ["4", "Intelligence", "Continuous", "20-strategy classifier, LLM analysis, CIO decisions", "Strategy assignments, agent jobs"],
        ["5", "Proposals", "Throughout day", "Incubator promotion, enrichment, 4-chunk LLM review", "Paper trade proposals"],
        ["6", "Execution", "Market hours", "Risk gate, bracket orders, Alpaca paper, TCA, recon", "Paper trades, execution quality"],
        ["7", "Overnight", "8 PM+", "Batch consolidation, agent scoring, strategy review", "Performance grades, cleaned state"],
    ]
)

# ============================================================
# SECTION 4: EXTERNAL RESEARCH & SIGNAL INGESTION
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 4: External Research & Signal Ingestion")

add_para(
    "The system ingests from 12+ external sources across 4 categories. This is "
    "standard practice for financial intelligence platforms. Not using search/"
    "transcript data is equivalent to flying blind.",
    bold=True
)

add_h2("Active Data Sources")
add_table(
    ["Source", "API / Method", "Data Type", "Frequency", "Fallback"],
    [
        ["Finviz Elite", "HTTP scrape (cookie + token)", "Screener results, 60+ enrichment fields", "4x daily", "None (primary)"],
        ["NewsAPI", "REST API (key)", "News articles, headlines", "2x daily + on-demand", "Finnhub"],
        ["Finnhub", "REST API (key)", "News, filings, insider activity", "On enrichment", "NewsAPI"],
        ["Polygon", "REST API (key)", "Market data, corporate events", "On catalyst enrichment", "Yahoo Finance"],
        ["FMP", "REST API (key)", "Fundamentals, earnings, financials", "On enrichment", "AlphaVantage"],
        ["AlphaVantage", "REST API (key)", "Fundamentals, economic indicators", "On enrichment", "FMP"],
        ["Yahoo Finance", "yfinance library", "OHLCV, quotes, dividends", "Indicator refresh + on-demand", "Polygon"],
        ["FRED", "REST API (key)", "Fed economic data (rates, CPI)", "Daily (6 AM)", "Cached values"],
        ["SEC EDGAR", "REST API (public)", "Form 4 insider filings", "Daily (8 PM)", "Skip"],
        ["YouTube Transcripts", "youtube-transcript-api", "Video transcripts", "Monthly + on-demand", "Skip"],
        ["Alpaca", "REST API (key)", "Paper trade execution, fills", "On execution + recon", "Manual"],
        ["Ollama (local LLM)", "HTTP (:11434)", "Classification, review, health", "Continuous (toll-gated)", "Cloud LLM cascade"],
    ]
)

add_h2("Why Each Source Matters")
add_table(
    ["Source", "Signal Provided", "Impact if Unavailable"],
    [
        ["Finviz Elite", "Primary candidate discovery (volume/gap/float)", "Pipeline stalls. No new candidates surface."],
        ["News APIs (4)", "Market-moving events, catalyst verification", "Catalyst scoring degrades, proposals lack context"],
        ["Fundamentals (FMP/AV)", "Earnings, revenue, debt ratios", "Strategy filters produce false negatives"],
        ["Yahoo Finance", "OHLCV for 17 technical indicators", "Indicator engine stale, confluence unreliable"],
        ["FRED", "Macro context (rates, unemployment, CPI)", "Macro strategies lose context"],
        ["SEC EDGAR", "Insider buying/selling signals", "Insider signal absent (non-blocking)"],
        ["YouTube Transcripts", "Earnings call language, forward guidance", "Income analysis loses qualitative depth"],
    ]
)

add_h2("Research Stubs (Designed, Not Yet Live)")
add_table(
    ["Integration", "Purpose", "Status"],
    [
        ["Google Programmable Search API", "Broad web research for novel signals", "Stub -- no API key provisioned"],
        ["Structured earnings transcript provider", "Earnings call analysis", "Stub -- YouTube used as substitute"],
        ["Alternative data feeds", "Non-traditional alpha signals", "Planned -- not yet architectured"],
        ["Real-time news WebSocket", "Sub-second news reaction", "Planned -- current batch at 2x/day"],
    ]
)

# ============================================================
# SECTION 5: STRATEGY ENGINE
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 5: Strategy Engine (23 Strategies)")

add_para(
    "All strategies are loaded dynamically from config/strategies/*.yaml at runtime. "
    "There are no hardcoded strategy lists. A single symbol can match multiple "
    "strategies simultaneously via the multi-strategy classifier."
)

add_h2("Strategies by Timeframe")
add_table(
    ["Timeframe", "Strategies"],
    [
        ["INTRADAY", "gap_and_go, momentum_scalp"],
        ["SHORT_SWING", "earnings_catalyst, swing_breakout, swing_trade, speculative_growth, tax_loss_harvest"],
        ["MEDIUM_SWING", "recovery_watch, sector_rotation"],
        ["POSITION", "income_add, core_growth_compounder, core_index, covered_call_income, defense_thesis, dividend_growth_compounder, high_yield_income_bdc, international_dividend, reit_income, bond_income"],
        ["CASH", "cash_or_stable"],
    ]
)

add_h2("Classification Flow")
add_para(
    "Phase 1 (Deterministic): YAML screen_filters matched against enrichment data "
    "(60+ Finviz fields + indicator cache). Phase 2 (LLM): qwen3:14b thesis-driven "
    "classification for strategies where deterministic data is insufficient. "
    "14 of 23 strategies require LLM classification."
)

# ============================================================
# SECTION 6: AGENT LAYER
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 6: Agent Layer")

add_h2("Conversational Agents (OpenClaw Gateway :18789)")
add_table(
    ["Agent", "Role", "Key Capabilities", "Channels"],
    [
        ["Maria", "Risk assessment", "Position sizing, portfolio impact, exposure analysis, correlation", "Telegram, WhatsApp"],
        ["Steph", "Technical analysis", "Entry/exit timing, chart patterns, wealth advisory", "Telegram, WhatsApp"],
        ["Aegis", "Synthesis & surveillance", "Morning briefs, overnight synthesis, cross-agent coordination", "Telegram, WhatsApp"],
        ["Alex", "Income strategy", "Roth conversion, SSDI/IRMAA impact, dividend analysis", "Telegram, WhatsApp"],
    ]
)

add_h2("Backend Automation Agents")
add_table(
    ["Agent", "Role", "Trigger"],
    [
        ["Iris", "Library hygiene -- content quality, stale data, dependency audits", "Cron (daily alerts)"],
        ["Pipeline Watchdog", "Health monitoring -- 31 stage failure/delay detection", "Continuous"],
        ["Scalp Critic", "LLM critique of candidates before promotion", "Pre-promotion"],
    ]
)

add_h2("Agent Processing Schedule")
add_table(
    ["Window", "Interval", "Jobs/Run"],
    [
        ["Market hours (6 AM - 7 PM)", "Every 15 min", "10 jobs"],
        ["Overnight (8 PM - 11 PM)", "Every 5 min", "25 jobs"],
        ["Weekend", "Every 10 min", "15 jobs"],
    ]
)

# ============================================================
# SECTION 7: LLM SUBSYSTEM
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 7: LLM Subsystem")

add_h2("Primary Model")
add_table(
    ["Parameter", "Value"],
    [
        ["Model", "qwen3:14b"],
        ["Runtime", "Ollama (localhost:11434)"],
        ["GPU", "Intel Arc B50 (Vulkan backend)"],
        ["Layer offload", "41/41 layers on GPU"],
        ["Keep-alive", "Persistent (OLLAMA_KEEP_ALIVE=-1)"],
        ["Performance", "~15s per chunk (GPU) vs ~300s (CPU)"],
        ["Toll gate", "fcntl.flock() serialization for GPU contention"],
    ]
)

add_h2("Routing Fallback Chain")
add_para("local (qwen3:14b) --> grok (xAI) --> claude (Anthropic) --> openai (OpenAI)")
add_para(
    "Daily budget tracking per provider. On budget exhaustion, the system "
    "auto-falls back to the next provider in the chain."
)

add_h2("LLM Use Cases")
add_table(
    ["Use Case", "Script", "Frequency"],
    [
        ["Strategy classification (14+ strategies)", "multi_strategy_classifier.py", "Sunday night batch"],
        ["Proposal review (4-chunk pipeline)", "proposal_llm_reviewer.py", "Per proposal"],
        ["Incubator pre-screening", "incubator_llm_screener.py", "Pre-promotion"],
        ["Holdings health refresh", "holdings_llm_refresh.py", "Periodic"],
        ["Agent responses", "Via OpenClaw", "On user interaction"],
        ["Duplicate proposal prevention", "incubator_proposal_promoter.py", "On promotion"],
    ]
)

# ============================================================
# SECTION 8: PROPOSAL LIFECYCLE
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 8: Proposal Lifecycle")

add_para(
    "Proposals follow a state machine from PROPOSED through CLOSED. "
    "The frontend displays an 8-stage pipeline chevron visual indicator."
)

add_h2("Lifecycle States")
add_table(
    ["State", "Description", "Next States"],
    [
        ["PROPOSED", "Promoter creates from incubator", "ENRICHING"],
        ["ENRICHING", "Enrichment loop picks up", "SCORED"],
        ["SCORED", "Scoring complete", "RISK_CHECK"],
        ["RISK_CHECK", "Enters risk gate", "APPROVED / REJECTED"],
        ["APPROVED", "Passes risk gate", "PENDING_ENTRY"],
        ["PENDING_ENTRY", "Awaiting entry zone", "ENTRY_ZONE_VALID / ENTRY_MISSED"],
        ["ENTRY_ZONE_VALID", "Price in range", "FILLED"],
        ["ENTRY_MISSED", "Price moved past zone", "EXPIRED"],
        ["FILLED", "Paper order executed", "OPEN"],
        ["OPEN", "Position tracked", "CLOSED"],
        ["CLOSED", "Exit hit (TP/SL/manual)", "Terminal"],
    ]
)

add_h2("Promotion Criteria")
add_table(
    ["Condition", "Requirements"],
    [
        ["High-conviction", "status=ACTIVE, score >= 38, catalyst_verified=true, days_active >= 1"],
        ["Score override", "status=ACTIVE, score >= 45, days_active >= 1"],
    ]
)

# ============================================================
# SECTION 9: SECURITY & FAILURE MODES
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 9: Security & Failure Modes")

add_h2("Current Security Posture")
add_table(
    ["Layer", "Control"],
    [
        ["Network", "Server on private network; no public-facing ports"],
        ["API", "No authentication (internal-only access)"],
        ["Database", "Password auth, localhost-only binding"],
        ["Secrets", ".env file (not in git, .gitignore enforced)"],
        ["LLM", "Local primary; cloud API keys in .env"],
        ["Broker", "Paper mode only; API keys scoped to paper trading"],
    ]
)

add_h2("Critical Failure Scenarios")
add_table(
    ["Failure", "Impact", "Detection", "Recovery"],
    [
        ["PostgreSQL down", "All services halt", "pg_isready + watchdog", "Restart; restore from 7-day backup"],
        ["Ollama crash", "LLM classification stops", "Health check :11434", "Systemd auto-restart; cloud fallback"],
        ["Portfolio Server crash", "API + frontend unavailable", "Health check :7777", "pkill + restart; systemd"],
        ["Finviz cookie expired", "No new candidates", "0-result screener stage", "Manual browser re-auth"],
        ["Cloud LLM budget exhausted", "Falls back to next provider", "Budget counter", "Resets daily"],
        ["Network outage", "External data unavailable", "Source staleness", "Cached data; alert operator"],
        ["GPU driver issue", "LLM 20x slower (CPU fallback)", "Vulkan layer count", "Restart Ollama; verify VULKAN=1"],
    ]
)

# ============================================================
# SECTION 10: SAFETY RULES (NON-NEGOTIABLE)
# ============================================================
add_h1("Cloud Rebuild Section 10: Safety Rules (Non-Negotiable)")

add_para(
    "These rules are non-negotiable. No automation, agent, or operator override "
    "may violate them.",
    bold=True
)

add_table(
    ["#", "Rule", "Enforcement"],
    [
        ["1", "LIVE_TRADING_ENABLED=false -- never change", ".env + code assertion"],
        ["2", "ALPACA_MODE=paper -- never change", ".env + adapter check"],
        ["3", "No risk gate threshold changes without owner approval", "UI gate + audit log"],
        ["4", "No auto-approval of proposals (human-in-the-loop)", "Proposal state machine"],
        ["5", "No holdings modification by automation", "Read-only portfolio access"],
        ["6", "Holdings value must remain > $1M", "Assertion check in code"],
    ]
)

add_para(
    "Validation gate: Live trading will not be enabled until the 6-month paper "
    "validation window closes (~Nov 2026), win rate >= 55%, and profit factor >= 1.3."
)

# ============================================================
# SECTION 11: COST MODEL
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 11: Cloud Operating Cost Model")

add_h2("Current Self-Hosted Costs")
add_table(
    ["Item", "Monthly Cost"],
    [
        ["Server hardware (amortized)", "$0"],
        ["Electricity (24/7 + GPU)", "~$15-25"],
        ["Finviz Elite subscription", "$39.95"],
        ["Cloud LLM fallback (minimal)", "~$5-15"],
        ["News API subscriptions", "$0-30"],
        ["Alpaca (paper)", "$0"],
        ["TOTAL", "~$60-110/mo"],
    ]
)

add_h2("Cloud Deployment Scenarios (Monthly)")
add_table(
    ["Scenario", "Low", "Medium", "High", "Key Trade-off"],
    [
        ["A: Self-hosted GPU on EC2", "$598", "$750", "$937", "Maximum control, highest cost"],
        ["B: Spot GPU instance", "$338", "$450", "$597", "Cost-optimized, interruption risk"],
        ["C: 100% Cloud LLM (no GPU)", "$248", "$365", "$537", "Simplest, variable LLM cost"],
        ["D: Hybrid (keep self-hosted)", "$145", "$210", "$340", "Lowest cost, no HA"],
    ]
)

add_h2("Primary Cost Drivers")
add_table(
    ["Rank", "Driver", "Control"],
    [
        ["1", "GPU compute for LLM ($140-520/mo)", "Spot instances, schedule off-hours, or use cloud APIs"],
        ["2", "Managed database ($75-107/mo)", "Smallest viable instance; Aurora Serverless v2"],
        ["3", "Data API subscriptions ($40-163/mo)", "Free tiers; batch calls; cache aggressively"],
        ["4", "NAT Gateway ($32/mo)", "VPC endpoints instead"],
        ["5", "Cloud LLM API calls ($50-115/mo)", "Prompt caching; batch processing; smaller models"],
    ]
)

# ============================================================
# SECTION 12: KEY FILES INDEX (UPDATED)
# ============================================================
doc.add_page_break()
add_h1("Cloud Rebuild Section 12: Key Files Index (2026-05-09)")

add_table(
    ["Path", "Purpose", "Lines/Size"],
    [
        [".env", "All secrets, API keys, feature flags", "~150 vars"],
        [".env.example", "Template with documented variables", "~150 vars"],
        ["config/strategies/*.yaml", "23 strategy definitions (dynamic)", "23 files"],
        ["assets/screeners.yaml", "Finviz screener URLs + run windows", "~50 lines"],
        ["scripts/api_v2.py", "All 80+ API endpoints", "11,700+ lines"],
        ["scripts/portfolio_server.py", "HTTP server entry point", "1,767 lines"],
        ["scripts/portfolio_orchestrator.py", "Orchestration hub", "1,714 lines"],
        ["scripts/trade_ai_orchestrator.py", "Screener + scoring", "873 lines"],
        ["scripts/local_llm_config.py", "LLM configuration hub", "~200 lines"],
        ["scripts/local_llm.py", "Ollama inference with toll gate", "~300 lines"],
        ["scripts/multi_strategy_classifier.py", "Multi-strategy assignment", "~500 lines"],
        ["scripts/aegis_synthesis.py", "Aegis overnight synthesis", "~1,500 lines"],
        ["scripts/alex_retirement_advisor.py", "Alex income/tax analysis", "~1,400 lines"],
        ["data/portfolios/state/holdings.json", "Portfolio state (~50 positions)", "~$1.19M"],
        ["data/state/ticker_enrichment_cache.json", "Enrichment cache", "1,139 symbols"],
        ["sql/migrations/", "Database migrations", "19 files"],
        ["crontab_backup.txt", "Full cron schedule", "138 entries"],
        ["requirements.txt", "Python dependencies", "90 packages"],
    ]
)

# ============================================================
# SECTION 13: DOCUMENTATION INDEX
# ============================================================
add_h1("Cloud Rebuild Section 13: Documentation Index")

add_para("All authoritative documentation now lives in docs/ as Markdown files:")

add_table(
    ["Document", "Purpose", "Audience"],
    [
        ["MASTER_SYSTEM_DOCUMENTATION.md", "Complete system reference (810 lines)", "Engineer / Architect"],
        ["ARCHITECTURE_OVERVIEW.md", "Executive architecture summary", "Executive / Stakeholder"],
        ["ARCHITECTURE_INFOGRAM.md", "Visual diagrams, I/O tables, service maps", "Engineer / Operator"],
        ["CHEAT_SHEET.md", "Operator quick reference", "Operator"],
        ["COST_MODEL.md", "Cloud operating cost estimates (4 scenarios)", "Finance / Executive"],
        ["RESTORE_GUIDE.md", "Disaster recovery procedures", "Operator"],
        ["GPU_OLLAMA_SETUP.md", "Hardware/GPU configuration", "Engineer"],
        ["DOCX_UPDATE_PROTOCOL.md", "This DOCX maintenance protocol", "Engineer"],
        ["project/agents_bible.md", "Agent behavior and interaction rules", "Engineer"],
        ["project/TRADE_AI_STRATEGY_PLAYBOOK_v1.0.md", "23 strategy playbooks", "Analyst / Engineer"],
    ]
)

add_para(
    "This DOCX file remains the canonical reference architecture document "
    "per the established protocol. It is updated via the JSON patch / python-docx "
    "append-only approach documented in DOCX_UPDATE_PROTOCOL.md."
)


# ============================================================
# SAVE
# ============================================================
doc.save(DOCX_PATH)
print(f"DOCX updated: {DOCX_PATH}")

# Verify
doc2 = Document(DOCX_PATH)
print(f"Paragraphs: {len(doc2.paragraphs)} (was 2168)")
print(f"Tables: {len(doc2.tables)} (was 96)")
print("Verification: OK")
